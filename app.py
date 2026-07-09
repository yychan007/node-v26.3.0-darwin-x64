# -*- coding: utf-8 -*-
"""
Upgraded Technical Standards Search Portal
Features:
- Flask web app
- Login + role-based access
- SQLite database
- Upload PDF/TXT/CSV/XLSX/DOCX
- Deduplication by SHA256
- Requirement-level parsing
- Better Dutch-English search
- OCR fallback for scanned PDFs
- PDF preview and page jump (with inline page image)
- Table preview for CSV/XLSX
- Requirement full block display

Run:
    python app.py

Default URL:
    [127.0.0.1](http://127.0.0.1:5000)

Initial admin:
    set ADMIN_USERNAME and ADMIN_PASSWORD before first start

For production, also set:
    APP_ENV=production
    APP_SECRET_KEY=<strong-random-secret>
"""

import os
import re
import io
import csv
import math
import json
import time
import uuid
import hashlib
import html
import secrets
import zipfile
import tempfile
import mimetypes
import xml.etree.ElementTree as ET
from datetime import datetime
from collections import defaultdict, Counter
from pathlib import Path
from typing import Any, cast

import pdfplumber
import pandas as pd
from openpyxl import Workbook
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Border, Font, Side
from openpyxl.utils import get_column_letter

import storage
import translation

from flask import (
    Flask, request, render_template_string, redirect, url_for,
    send_file, flash, abort, session, jsonify
)
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash

from flask_sqlalchemy import SQLAlchemy  # pyright: ignore[reportMissingImports]
from flask_login import (  # pyright: ignore[reportMissingImports]
    LoginManager, login_user, logout_user, login_required,
    current_user, UserMixin
)
from sqlalchemy.exc import OperationalError, PendingRollbackError  # pyright: ignore[reportMissingImports]
from werkzeug.exceptions import HTTPException

from PIL import Image

def extract_text_from_image(path):
    if not TESSERACT_AVAILABLE or pytesseract is None:
        return "", [(0, 0, 1)], []

    try:
        image = Image.open(path)
        text = pytesseract.image_to_string(image)
        return text, [(0, len(text), 1)], []
    except Exception as e:
        print(f"Image OCR error: {path} -> {e}")
        return "", [(0, 0, 1)], []


# Optional packages
DOCX_AVAILABLE = True
DocxDocument = None
try:
    from docx import Document as DocxDocument
except Exception:
    DOCX_AVAILABLE = False

OCR_AVAILABLE = True
fitz = None
try:
    import fitz  # PyMuPDF
except Exception:
    OCR_AVAILABLE = False

TESSERACT_AVAILABLE = True
PDF2IMAGE_AVAILABLE = False
pytesseract = None
try:
    import pytesseract
except Exception:
    TESSERACT_AVAILABLE = False

convert_from_path = None
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except Exception:
    PDF2IMAGE_AVAILABLE = False

SEMANTIC_AVAILABLE = True
SentenceTransformer = None
try:
    from sentence_transformers import SentenceTransformer
except Exception:
    SEMANTIC_AVAILABLE = False


# =========================================================
# Base paths
# =========================================================
try:
    BASE_DIR = Path(__file__).resolve().parent
except NameError:
    BASE_DIR = Path.cwd()

APP_ENV = os.environ.get("APP_ENV", "development").strip().lower() or "development"
IS_PRODUCTION = APP_ENV == "production"


def env_flag(name, default=False):
    value = os.environ.get(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def env_int(name, default):
    raw = os.environ.get(name)
    if raw is None or str(raw).strip() == "":
        return default
    return int(raw)


SEMANTIC_SEARCH_ENABLED = env_flag("ENABLE_SEMANTIC_SEARCH", default=False)
OCR_ENABLED = env_flag("ENABLE_OCR", default=False)
INDEX_BATCH_SIZE = int(os.environ.get("INDEX_BATCH_SIZE", "50"))
MAX_PDF_PAGES = env_int("MAX_PDF_PAGES", 100 if IS_PRODUCTION else 0)
MAX_INDEX_FILE_MB = env_int("MAX_INDEX_FILE_MB", 40 if IS_PRODUCTION else 0)
TEXT_PREVIEW_LIMIT = env_int("TEXT_PREVIEW_LIMIT", 150000)
TRANSLATE_MAX_CELLS = env_int("TRANSLATE_MAX_CELLS", 500)
STORAGE_QUOTA_MB = env_int("STORAGE_QUOTA_MB", 1024)


DEFAULT_DATA_DIR = BASE_DIR / "data"
DATA_DIR = Path(os.environ.get("DATA_DIR", str(DEFAULT_DATA_DIR))).expanduser()
DOC_FOLDER = DATA_DIR / "documents"
DICT_FOLDER = DATA_DIR / "dictionaries"
PREVIEW_FOLDER = DATA_DIR / "previews"
DATABASE_PATH = Path(
    os.environ.get("DATABASE_PATH", str(DATA_DIR / "search_portal.db"))
).expanduser()

for path in (DATA_DIR, DOC_FOLDER, DICT_FOLDER, PREVIEW_FOLDER, DATABASE_PATH.parent):
    path.mkdir(parents=True, exist_ok=True)

REQUIREMENT_MASTER_SCAN_CACHE_TTL_SECONDS = 300
_requirement_master_scan_cache = {}
REQUIREMENT_LOOKUP_MAX_SCAN_SECONDS = 8.0
REQUIREMENT_LOOKUP_MAX_MASTER_DOCS = 1

REQUIREMENT_MASTER_RESULT_CACHE_TTL_SECONDS = 300
_requirement_master_result_cache = {}


def build_database_uri():
    database_url = os.environ.get("DATABASE_URL", "").strip()
    if database_url:
        if database_url.startswith("postgres://"):
            database_url = database_url.replace(
                "postgres://", "postgresql+psycopg://", 1
            )
        elif database_url.startswith("postgresql://"):
            database_url = database_url.replace(
                "postgresql://", "postgresql+psycopg://", 1
            )
        return database_url
    return f"sqlite:///{DATABASE_PATH}"


APP_SECRET_KEY = os.environ.get("APP_SECRET_KEY", "").strip()
if not APP_SECRET_KEY:
    if IS_PRODUCTION:
        raise RuntimeError("APP_SECRET_KEY must be set when APP_ENV=production.")
    APP_SECRET_KEY = "dev-secret-key-change-me"


# =========================================================
# Flask config
# =========================================================
app = Flask(__name__)
app.config["SECRET_KEY"] = APP_SECRET_KEY
app.config["SQLALCHEMY_DATABASE_URI"] = build_database_uri()
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    # Render/Postgres connections may drop; pre_ping avoids stale pooled conns.
    "pool_pre_ping": True,
    # Recycle periodically to reduce SSL EOF / idle disconnect issues.
    "pool_recycle": 180,
}
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024  # 100 MB
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["REMEMBER_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
if IS_PRODUCTION:
    app.config["SESSION_COOKIE_SECURE"] = True
    app.config["REMEMBER_COOKIE_SECURE"] = True

db = SQLAlchemy(app)
login_manager = LoginManager(app)
# Avoid static type issues with Flask-Login's stubs.
setattr(login_manager, "login_view", "login")
login_manager.login_message_category = "warning"

def _db_recover_and_retry(fn, *, retries=1, label="db_op"):
    """
    Best-effort recovery for transient Postgres/SSL disconnects on Render.
    Rolls back the session, disposes the engine, then retries.
    """
    last_exc = None
    for attempt in range(retries + 1):
        try:
            return fn()
        except (OperationalError, PendingRollbackError) as exc:
            last_exc = exc
            try:
                db.session.rollback()
            except Exception:
                pass
            try:
                db.engine.dispose()
            except Exception:
                pass
            if attempt >= retries:
                raise
            time.sleep(0.15)
        except Exception as exc:
            # Other exceptions should not be retried here.
            raise
    raise RuntimeError(f"{label} failed")  # pragma: no cover


@app.errorhandler(Exception)
def _handle_unexpected_exception(exc):
    if isinstance(exc, HTTPException):
        return exc
    try:
        db.session.rollback()
    except Exception:
        pass
    try:
        db.engine.dispose()
    except Exception:
        pass
    message = str(exc) or repr(exc)
    return (
        render_template_string(
            """
<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>Server error</title>
  """ + BASE_CSS + """
</head>
<body>
  <div class="container">
    <h2>Internal Server Error</h2>
    <div class="warning" style="margin-top:12px;">
      <strong>Error:</strong> {{ message }}
    </div>
    <div class="notice" style="margin-top:12px;">
      Try going back to <a href="{{ url_for('home') }}">Home</a>.
    </div>
  </div>
</body>
</html>
""",
            message=message,
        ),
        500,
    )


@app.context_processor
def inject_translation_context():
    return {
        "translation_enabled": translation.translation_enabled(),
        "translation_provider_info": get_translation_status(),
    }


def get_translation_status():
    status_fn = getattr(translation, "provider_status", None)
    if callable(status_fn):
        return status_fn()
    provider_name = get_translation_provider()
    return {
        "enabled": translation.translation_enabled(),
        "provider": provider_name,
        "configured": True,
        "model": "google-translate",
    }


def get_translation_provider():
    provider_fn = getattr(translation, "translation_provider", None)
    return provider_fn() if callable(provider_fn) else "google"

# =========================================================
# Semantic model lazy load
# =========================================================
semantic_model = None

def get_semantic_model():
    global semantic_model
    if not SEMANTIC_SEARCH_ENABLED:
        return None
    if semantic_model is None and SEMANTIC_AVAILABLE and SentenceTransformer is not None:
        try:
            semantic_model = SentenceTransformer("paraphrase-multilingual-MiniLM-L12-v2")
        except Exception:
            semantic_model = None
    return semantic_model


# =========================================================
# Settings
# =========================================================
ALLOWED_EXTENSIONS = {
    "pdf", "txt", "csv", "xlsx", "docx",
    "png", "jpg", "jpeg"
}
if DOCX_AVAILABLE:
    ALLOWED_EXTENSIONS.add("docx")

RESULTS_PER_PAGE = 10

ADMIN_ROLE = "admin"
USER_ROLE = "user"

STOPWORDS = {
    "a","an","and","are","as","at","be","by","for","from","has","he","in","is","it","its",
    "of","on","that","the","to","was","were","will","with","i","you","we","they","this",
    "these","those","or","but","so","not","can","could","would","should","may","might",
    "do","does","did","done","have","had","having","been","being","if","then","than",
    "there","here","about","into","over","under","de","het","een","en","van","op","te",
    "met","voor","dat","die","is","zijn","wordt","moet","dient","bij","of","in","aan",
    "tot","als","uit","door","niet","welke","welk"
}

MULTI_LANG_SYNONYMS = {
    "aarding": ["aarding", "aard", "geaard", "aardrail", "aardverbinding", "grounding", "earthing", "earth"],
    "grounding": ["grounding", "earthing", "earth", "aarding", "aard", "geaard", "aardrail", "aardverbinding"],
    "earthing": ["earthing", "grounding", "earth", "aarding", "aard", "geaard", "aardrail", "aardverbinding"],
    "earth": ["earth", "earthing", "grounding", "aarding", "aard", "geaard"],
    "resistance": ["resistance", "weerstand"],
    "weerstand": ["weerstand", "resistance"],
    "voltage": ["voltage", "spanning"],
    "spanning": ["spanning", "voltage"],
    "transformer": ["transformer", "transformator"],
    "transformator": ["transformator", "transformer"],
    "blast": ["blast", "scherf", "scherfwand", "blastwall", "blast-wall"],
    "wall": ["wall", "wand", "scherfwand", "blast wall", "blastwall"],
    "blastwall": ["blastwall", "blast wall", "scherfwand"],
    "blast-wall": ["blast-wall", "blast wall", "scherfwand"],
    "scherfwand": ["scherfwand", "blast wall", "blastwall", "blast-wall"],
    "requirement": ["requirement", "requirements", "req", "eis", "eisen"],
    "req": ["req", "requirement", "eis"],
    "eis": ["eis", "requirement", "req"],
    "kabel": ["kabel", "cable", "cables"],
    "cable": ["cable", "cables", "kabel", "kabels"],
    "cables": ["cables", "cable", "kabel", "kabels"],
    "bescherming": ["bescherming", "protection"],
    "protection": ["protection", "bescherming"],
    "installatie": ["installatie", "installation"],
    "installation": ["installation", "installatie"],
    "stroom": ["stroom", "current"],
    "current": ["current", "stroom"],
    "frequentie": ["frequentie", "frequency"],
    "frequency": ["frequency", "frequentie"],
}

DRAWING_MARKERS = ("drawing", "tekening")
CONTENTS_MARKERS = (
    "table of contents",
    "contents",
    "inhoudsopgave",
    "inhoud",
)
DOCX_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
DOCX_R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
DOCX_A_NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"

REQ_ID_REGEX = re.compile(r'\b([A-Z]{1,10}-Req-\d+(?:\.\d+)?)\b', re.IGNORECASE)

DOCUMENT_TYPE_STANDARD = "standard"
DOCUMENT_TYPE_REQUIREMENT_MASTER = "requirement_master"

DOCUMENT_TYPE_OPTIONS = [
    {"id": DOCUMENT_TYPE_STANDARD, "label": "Standard document"},
    {"id": DOCUMENT_TYPE_REQUIREMENT_MASTER, "label": "Requirement master table (總表)"},
]
DOCUMENT_TYPE_IDS = {option["id"] for option in DOCUMENT_TYPE_OPTIONS}
DOCUMENT_TYPE_SPREADSHEET_ONLY = {DOCUMENT_TYPE_REQUIREMENT_MASTER}

# =========================================================
# Database models
# =========================================================
class User(db.Model, UserMixin):
    __tablename__ = "users"

    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(120), unique=True, nullable=False, index=True)
    password_hash = db.Column(db.String(255), nullable=False)
    role = db.Column(db.String(20), nullable=False, default=USER_ROLE)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    @property
    def is_admin(self):
        return self.role == ADMIN_ROLE


class DocumentRecord(db.Model):
    __tablename__ = "documents"

    id = db.Column(db.Integer, primary_key=True)
    original_filename = db.Column(db.String(500), nullable=False)
    stored_filename = db.Column(db.String(500), nullable=False, unique=True)
    extension = db.Column(db.String(20), nullable=False)
    file_hash = db.Column(db.String(64), nullable=False, unique=True, index=True)
    size_bytes = db.Column(db.Integer, nullable=False)
    uploaded_by = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    uploaded_at = db.Column(db.DateTime, default=datetime.utcnow)
    is_ocr = db.Column(db.Boolean, default=False)
    text_preview = db.Column(db.Text, default="")
    page_offsets_json = db.Column(db.Text, default="")
    status = db.Column(db.String(50), default="indexed")
    document_type = db.Column(db.String(50), default=DOCUMENT_TYPE_STANDARD, nullable=False)

    requirements = db.relationship("RequirementBlock", backref="document", cascade="all, delete-orphan")
    tables = db.relationship("TablePreview", backref="document", cascade="all, delete-orphan")


def normalize_document_type(value):
    text = (value or "").strip().lower()
    return text if text in DOCUMENT_TYPE_IDS else DOCUMENT_TYPE_STANDARD


def document_type_label(document_type):
    for option in DOCUMENT_TYPE_OPTIONS:
        if option["id"] == document_type:
            return option["label"]
    return DOCUMENT_TYPE_OPTIONS[0]["label"]


def is_requirement_master_document(doc):
    return normalize_document_type(getattr(doc, "document_type", None)) == DOCUMENT_TYPE_REQUIREMENT_MASTER


def is_tennet_requirements_document(doc):
    filename = (getattr(doc, "original_filename", "") or "").lower()
    compact = re.sub(r"[^a-z0-9]+", "", filename)
    return "tennetrequirements" in compact


def should_skip_table_translation(doc):
    return is_tennet_requirements_document(doc)


def is_requirement_master_table_dataframe(df):
    if df is None or df.empty:
        return False
    normalized_columns = {
        re.sub(r"[^a-z0-9]+", "", str(col).strip().lower())
        for col in df.columns
        if str(col).strip()
    }
    return "speccode" in normalized_columns and "sourcedocument" in normalized_columns


class RequirementBlock(db.Model):
    __tablename__ = "requirement_blocks"

    id = db.Column(db.Integer, primary_key=True)
    document_id = db.Column(db.Integer, db.ForeignKey("documents.id"), nullable=False, index=True)

    requirement_id = db.Column(db.String(120), index=True)
    title = db.Column(db.String(500), default="")
    section = db.Column(db.String(300), default="")
    major_section = db.Column(db.String(300), default="")
    page = db.Column(db.Integer, default=1)
    char_start = db.Column(db.Integer, default=0)
    category = db.Column(db.String(120), default="General")
    definition = db.Column(db.Text, default="")
    summary = db.Column(db.Text, default="")
    full_text = db.Column(db.Text, nullable=False)

    token_blob = db.Column(db.Text, default="")
    text_hash = db.Column(db.String(64), index=True)
    semantic_text = db.Column(db.Text, default="")


class TablePreview(db.Model):
    __tablename__ = "table_previews"

    id = db.Column(db.Integer, primary_key=True)
    document_id = db.Column(db.Integer, db.ForeignKey("documents.id"), nullable=False, index=True)

    sheet_name = db.Column(db.String(255), default="")
    page = db.Column(db.Integer, default=1)
    table_format = db.Column(db.String(20), default="csv")
    html_table = db.Column(db.Text, default="")
    html_table_en = db.Column(db.Text, default="")
    html_table_es = db.Column(db.Text, default="")
    csv_text = db.Column(db.Text, default="")
    csv_text_en = db.Column(db.Text, default="")
    csv_text_es = db.Column(db.Text, default="")
    preview_title = db.Column(db.String(255), default="")


class TranslationCache(db.Model):
    __tablename__ = "translation_cache"

    text_hash = db.Column(db.String(64), primary_key=True)
    source_text = db.Column(db.Text, default="")
    translated_text = db.Column(db.Text, default="")
    provider = db.Column(db.String(32), default="")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class DocumentPageTranslation(db.Model):
    __tablename__ = "document_page_translations"
    __table_args__ = (
        db.UniqueConstraint("document_id", "page", name="uq_document_page_translation"),
    )

    id = db.Column(db.Integer, primary_key=True)
    document_id = db.Column(
        db.Integer, db.ForeignKey("documents.id", ondelete="CASCADE"), nullable=False
    )
    page = db.Column(db.Integer, nullable=False)
    source_text = db.Column(db.Text, default="")
    translated_text = db.Column(db.Text, default="")
    translated_text_es = db.Column(db.Text, default="")
    provider = db.Column(db.String(32), default="")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class DictionarySource(db.Model):
    __tablename__ = "dictionary_sources"

    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(255), nullable=False)
    original_filename = db.Column(db.String(500), nullable=False, index=True)
    stored_filename = db.Column(db.String(500), nullable=False, unique=True)
    sheet_name = db.Column(db.String(255), default="")
    entry_kind = db.Column(db.String(32), default="abbreviation")
    entry_count = db.Column(db.Integer, default=0)
    size_bytes = db.Column(db.Integer, default=0)
    uploaded_by = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    uploaded_at = db.Column(db.DateTime, default=datetime.utcnow)

    entries = db.relationship(
        "DictionaryEntry", backref="source", cascade="all, delete-orphan"
    )


class DictionaryEntry(db.Model):
    __tablename__ = "dictionary_entries"

    id = db.Column(db.Integer, primary_key=True)
    source_id = db.Column(
        db.Integer,
        db.ForeignKey("dictionary_sources.id", ondelete="CASCADE"),
        nullable=False,
        index=True,
    )
    term = db.Column(db.String(500), nullable=False, index=True)
    content = db.Column(db.Text, default="")
    content_nl = db.Column(db.Text, default="")
    content_en = db.Column(db.Text, default="")
    search_keys = db.Column(db.Text, default="")
    entry_kind = db.Column(db.String(32), default="abbreviation")
    row_number = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))


# =========================================================
# Helpers
# =========================================================
def is_admin():
    return True

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def sha256_of_file(path, chunk_size=1024 * 1024):
    h = hashlib.sha256()
    with open(path, "rb") as f:
        while True:
            chunk = f.read(chunk_size)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()

def preprocess(text):
    text = (text or "").lower()
    tokens = re.findall(r"[a-zA-ZÀ-ÿ0-9\-_\.]+", text)
    return [t for t in tokens if t not in STOPWORDS and len(t) > 1]

def expand_query_tokens(tokens):
    out = []
    for t in tokens:
        variants = list(MULTI_LANG_SYNONYMS.get(t, [t]))
        if len(t) > 3 and t.endswith("s"):
            variants.append(t[:-1])
        elif len(t) > 2 and not t.endswith("s"):
            variants.append(t + "s")
        out.extend(variants)
    return list(dict.fromkeys(out))


def get_all_expanded_terms(query):
    q = (query or "").strip().lower()
    terms = []
    if q:
        terms.append(q)
    for token in preprocess(query):
        terms.extend(expand_query_tokens([token]))
    return list(dict.fromkeys(terms))


def default_exact_search_terms(query):
    tokens = preprocess(query)
    if tokens:
        return tokens
    q = (query or "").strip().lower()
    return [q] if q else []


def resolve_active_search_terms(query, selected_terms=None):
    all_terms = get_all_expanded_terms(query)
    if not selected_terms:
        return default_exact_search_terms(query)

    allowed = {term.lower() for term in all_terms}
    active = []
    for raw in selected_terms:
        value = (raw or "").strip().lower()
        if value not in allowed:
            continue
        for candidate in all_terms:
            if candidate.lower() == value:
                active.append(candidate)
                break
    return list(dict.fromkeys(active)) or default_exact_search_terms(query)


def term_matches_text(term, text):
    value = (term or "").strip().lower()
    haystack = (text or "").lower()
    if not value or not haystack:
        return False
    if " " in value:
        return value in haystack
    return re.search(r"\b" + re.escape(value) + r"\b", haystack) is not None


def count_exact_term_hits(block, query, exact_terms):
    combined = " ".join(
        [
            block.title or "",
            block.summary or "",
            block.definition or "",
            block.full_text or "",
            block.requirement_id or "",
        ]
    )
    hits = 0
    q = (query or "").strip().lower()
    if q and q in combined.lower():
        hits += 3
    for term in exact_terms:
        if term_matches_text(term, combined):
            hits += 1
    return hits

def detect_category(text, title=""):
    combined = f"{title} {text}".lower()
    if "emc" in combined and "aard" in combined:
        return "EMC grounding"
    if "aardrail" in combined:
        return "Earth rail"
    if "aardverbinding" in combined:
        return "Earth connection"
    if "aarding" in combined or "geaard" in combined or "grounding" in combined or "earthing" in combined:
        return "Grounding"
    if "klem" in combined and "aard" in combined:
        return "Grounding terminal"
    if "kast" in combined and ("aard" in combined or "geaard" in combined):
        return "Cabinet grounding"
    if "definitie" in combined or "definition" in combined:
        return "Definition"
    if "cable" in combined or "kabel" in combined:
        return "Cables"
    return "General"

def strip_html_legacy(text):
    return re.sub(r"<[^>]+>", "", text or "")


def get_cached_translation(text_hash):
    row = TranslationCache.query.get(text_hash)
    return row.translated_text if row else None


def store_cached_translation(text_hash, source_text, translated_text, provider=None):
    if not translated_text:
        return
    row = TranslationCache.query.get(text_hash)
    if row:
        return
    cache_cls = cast(Any, TranslationCache)
    db.session.add(
        cache_cls(
            text_hash=text_hash,
            source_text=source_text[:5000],
            translated_text=translated_text,
            provider=provider or get_translation_provider(),
        )
    )
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()


def get_document_full_text(doc):
    text = (doc.text_preview or "").strip()
    if text:
        return text

    ext = doc.extension.lower()
    if ext not in {"pdf", "docx", "txt"}:
        return ""
    if not storage.document_exists(doc.stored_filename, DOC_FOLDER):
        return ""

    try:
        with storage.open_document_local_path(doc.stored_filename, DOC_FOLDER) as file_path:
            if ext == "docx":
                text, _ = extract_text_from_docx(file_path)
            elif ext == "txt":
                text, _ = extract_text_from_txt(file_path)
            elif ext == "pdf" and OCR_AVAILABLE and fitz is not None:
                with fitz.open(file_path) as pdf_doc:
                    # Convert to str explicitly to keep the type checker happy.
                    text_parts = []
                    for i in range(len(pdf_doc)):
                        page_text = pdf_doc.load_page(i).get_text("text") or ""
                        text_parts.append(str(page_text))
                    text = "\n".join(text_parts)
            else:
                text = ""
    except Exception as exc:
        print(f"Full text extract error for doc {doc.id}: {exc}")
        text = ""
    return (text or "").strip()


def split_text_into_virtual_pages(text, max_chars=3500):
    text = (text or "").strip()
    if not text:
        return [""]
    if len(text) <= max_chars:
        return [text]

    parts = []
    start = 0
    while start < len(text):
        end = min(len(text), start + max_chars)
        if end < len(text):
            split_at = text.rfind("\n\n", start, end)
            if split_at <= start + 400:
                split_at = text.rfind("\n", start, end)
            if split_at > start + 400:
                end = split_at + 1
        parts.append(text[start:end].strip())
        start = end
    return [part for part in parts if part] or [""]


def block_index_to_virtual_page(block_index, total_blocks, total_pages):
    if total_pages <= 1 or total_blocks <= 0:
        return 1
    return max(
        1,
        min(total_pages, int(((block_index + 1) * total_pages - 1) / total_blocks) + 1),
    )


def get_docx_viewer_image_indexes(
    document_path, page=1, total_pages=1, highlight_index=None, max_images=8, page_tolerance=2
):
    candidates = collect_docx_drawing_image_candidates(document_path)
    if not candidates:
        return []

    indexes = []
    seen = set()

    def add_index(image_index):
        if image_index is None or image_index in seen:
            return
        seen.add(image_index)
        indexes.append(image_index)

    if highlight_index is not None:
        add_index(highlight_index)

    page_matches = []
    nearby_matches = []
    fallback_matches = []
    for candidate in candidates:
        image_index = candidate["image_index"]
        if image_index in seen:
            continue
        est_page = block_index_to_virtual_page(
            candidate.get("block_index", 0),
            candidate.get("total_blocks", 1),
            total_pages,
        )
        if est_page == page:
            page_matches.append(image_index)
        elif abs(est_page - page) <= page_tolerance:
            nearby_matches.append(image_index)
        elif candidate.get("in_drawing_zone"):
            fallback_matches.append(image_index)

    for image_index in page_matches + nearby_matches + fallback_matches:
        add_index(image_index)
        if len(indexes) >= max_images:
            break

    return indexes[:max_images]


def extract_page_text_from_preview(doc, page_num):
    if page_num < 1:
        return ""

    ext = doc.extension.lower()
    if ext in {"docx", "txt"}:
        chunks = split_text_into_virtual_pages(get_document_full_text(doc))
        if page_num > len(chunks):
            return ""
        return chunks[page_num - 1]

    offsets = load_page_offsets(doc)
    text = doc.text_preview or ""
    if offsets and text:
        for start, end, page in offsets:
            if page == page_num:
                return text[start:end].strip()
    return ""


def page_has_drawing_section_label(page_text):
    for line in (page_text or "").splitlines():
        normalized = re.sub(r"\s+", " ", line.strip().lower())
        if not normalized:
            continue
        if normalized in DRAWING_MARKERS:
            return True
        if (normalized.startswith("drawing ") or normalized.startswith("tekening ")) and len(normalized) < 140:
            return True
    return False


def _pdf_page_image_coverage(page):
    page_area = max(0.0, float(page.rect.width) * float(page.rect.height))
    if page_area <= 0:
        return 0.0, 0

    covered = 0.0
    image_count = 0
    try:
        image_infos = page.get_image_info(xrefs=True)
    except Exception:
        image_infos = []

    for info in image_infos or []:
        bbox = info.get("bbox")
        if not bbox or len(bbox) < 4:
            continue
        width = max(0.0, float(bbox[2]) - float(bbox[0]))
        height = max(0.0, float(bbox[3]) - float(bbox[1]))
        if width * height > 2500:
            covered += width * height
            image_count += 1

    return min(1.0, covered / page_area), image_count


def _is_visual_drawing_page(coverage, image_count, text_len):
    if image_count <= 0:
        return False
    if text_len > 1800 and coverage < 0.22:
        return False
    if coverage >= 0.24:
        return True
    if coverage >= 0.10 and text_len < 700:
        return True
    if coverage >= 0.16 and text_len < 1200:
        return True
    return False


def pdf_page_is_visual_drawing_page(doc, page_num, page_text="", pdf_doc=None):
    if not OCR_AVAILABLE or fitz is None:
        return False

    text = (page_text or extract_page_text_from_preview(doc, page_num) or "").strip()
    text_len = len(text)

    try:
        if pdf_doc is not None:
            if page_num < 1 or page_num > len(pdf_doc):
                return False
            page = pdf_doc.load_page(page_num - 1)
            coverage, image_count = _pdf_page_image_coverage(page)
            return _is_visual_drawing_page(coverage, image_count, text_len)

        if not storage.document_exists(doc.stored_filename, DOC_FOLDER):
            return False
        with storage.open_document_local_path(doc.stored_filename, DOC_FOLDER) as file_path:
            with fitz.open(file_path) as opened:
                if page_num < 1 or page_num > len(opened):
                    return False
                page = opened.load_page(page_num - 1)
                coverage, image_count = _pdf_page_image_coverage(page)
                return _is_visual_drawing_page(coverage, image_count, text_len)
    except Exception as exc:
        print(f"PDF visual drawing check error doc {doc.id} page {page_num}: {exc}")
        return False


def select_pdf_drawing_pages(doc, query, active_terms, center_page, max_images=4):
    offsets = load_page_offsets(doc)
    if not offsets:
        return []
    if not storage.document_exists(doc.stored_filename, DOC_FOLDER) or fitz is None:
        return []

    q = (query or "").lower().strip()
    query_terms = [t for t in re.findall(r"[a-z0-9\-]+", q) if len(t) > 1]
    expanded_terms = [(t or "").lower().strip() for t in (active_terms or []) if t]
    expanded_terms = [t for t in expanded_terms if len(t) > 1]
    total_pages = max((p for _s, _e, p in offsets), default=1)

    def page_matches_terms(page_text):
        page_text = (page_text or "").lower()
        if q and q in page_text:
            return True
        if any(term in page_text for term in expanded_terms):
            return True
        if any(token in page_text for token in query_terms):
            return True
        return False

    scored_pages = []
    seen_pages = set()
    try:
        with storage.open_document_local_path(doc.stored_filename, DOC_FOLDER) as file_path:
            with fitz.open(file_path) as pdf_doc:
                for _start, _end, page_num in offsets:
                    if page_num < 1 or page_num > total_pages:
                        continue

                    page_text = extract_page_text_from_preview(doc, page_num)
                    page_text_lower = (page_text or "").lower()
                    if any(marker in page_text_lower for marker in CONTENTS_MARKERS):
                        continue
                    if not pdf_page_is_visual_drawing_page(doc, page_num, page_text, pdf_doc=pdf_doc):
                        continue

                    prev_text = ""
                    if page_num > 1:
                        prev_text = extract_page_text_from_preview(doc, page_num - 1) or ""
                    prev_text_lower = prev_text.lower()
                    prev_has_drawing_label = page_has_drawing_section_label(prev_text)
                    prev_matches = page_matches_terms(prev_text_lower)
                    current_matches = page_matches_terms(page_text_lower)

                    if not (prev_matches or current_matches):
                        continue

                    score = 8.0
                    if prev_has_drawing_label and prev_matches:
                        score += 5.0
                    if current_matches:
                        score += 2.0
                    if page_has_drawing_section_label(page_text) and current_matches:
                        score += 1.0

                    term_hits = sum(1 for term in expanded_terms if term in prev_text_lower or term in page_text_lower)
                    score += min(term_hits, 5) * 1.5
                    query_hits = sum(1 for token in query_terms if token in prev_text_lower or token in page_text_lower)
                    score += min(query_hits, 5) * 1.2

                    if isinstance(center_page, int) and center_page > 0:
                        distance = abs(page_num - center_page)
                        score += max(0.0, 2.5 - (distance * 0.35))

                    if page_num not in seen_pages:
                        scored_pages.append((page_num, score))
                        seen_pages.add(page_num)
    except Exception as exc:
        print(f"PDF drawing page select error for doc {doc.id}: {exc}")
        return []

    if not scored_pages:
        return []

    scored_pages.sort(key=lambda item: (item[1], -abs(item[0] - (center_page or item[0]))), reverse=True)
    return [page for page, _score in scored_pages[:max_images]]


def pdf_source_file_missing(doc):
    return (
        doc.extension.lower() == "pdf"
        and not storage.document_exists(doc.stored_filename, DOC_FOLDER)
    )


def extract_single_page_text(doc, page_num):
    ext = doc.extension.lower()
    preview_text = extract_page_text_from_preview(doc, page_num)
    if preview_text:
        return preview_text

    if ext in {"docx", "txt"}:
        return ""

    if ext != "pdf" or not storage.document_exists(doc.stored_filename, DOC_FOLDER):
        return ""

    if not OCR_AVAILABLE or fitz is None:
        return ""

    try:
        with storage.open_document_local_path(doc.stored_filename, DOC_FOLDER) as file_path:
            with fitz.open(file_path) as pdf_doc:
                if page_num < 1 or page_num > len(pdf_doc):
                    return ""
                page_text = pdf_doc.load_page(page_num - 1).get_text("text") or ""
                return str(page_text).strip()
    except Exception as exc:
        print(f"Page text extract error for doc {doc.id} page {page_num}: {exc}")
        return ""


def _group_pdf_blocks_into_rows(blocks, page_height):
    if not blocks:
        return []

    heights = [max(4.0, b["y1"] - b["y0"]) for b in blocks]
    median_h = sorted(heights)[len(heights) // 2] if heights else 12.0
    y_tolerance = max(8.0, min(20.0, median_h * 0.75))

    for block in blocks:
        block["y_center"] = (block["y0"] + block["y1"]) / 2.0

    blocks.sort(key=lambda b: (b["y_center"], b["x0"]))
    rows = []
    for block in blocks:
        if not rows:
            rows.append({"y0": block["y0"], "cells": [block]})
            continue
        last_row = rows[-1]
        last_center = sum(c["y_center"] for c in last_row["cells"]) / len(last_row["cells"])
        if abs(last_center - block["y_center"]) <= y_tolerance:
            last_row["cells"].append(block)
        else:
            rows.append({"y0": block["y0"], "cells": [block]})

    for row in rows:
        row["cells"].sort(key=lambda c: c["x0"])
    return rows


def _translation_cell_class(cell, row_cells, page_width):
    text = (cell.get("translation") or "").strip()
    text_len = len(text)
    cell_count = len(row_cells)
    width_ratio = max(cell["x1"] - cell["x0"], 1.0) / max(page_width, 1.0)

    if cell_count == 1:
        return "translation-flow-cell translation-flow-cell-full"
    if cell_count >= 4 or (text_len <= 24 and width_ratio < 0.2):
        return "translation-flow-cell translation-flow-cell-table"
    if text_len <= 36 and (text.endswith(":") or width_ratio < 0.22):
        return "translation-flow-cell translation-flow-cell-label"
    if text_len <= 80 and width_ratio < 0.3:
        return "translation-flow-cell translation-flow-cell-narrow"
    return "translation-flow-cell translation-flow-cell-body"


def extract_pdf_page_layout_rows(doc, page_num, y_tolerance=12):
    if doc.extension.lower() != "pdf" or not OCR_AVAILABLE or fitz is None:
        return None
    if not storage.document_exists(doc.stored_filename, DOC_FOLDER):
        return None

    try:
        with storage.open_document_local_path(doc.stored_filename, DOC_FOLDER) as file_path:
            with fitz.open(file_path) as pdf_doc:
                if page_num < 1 or page_num > len(pdf_doc):
                    return None
                page = pdf_doc.load_page(page_num - 1)
                page_width = float(page.rect.width) or 1.0
                page_height = float(page.rect.height) or 1.0
                blocks = []
                for item in page.get_text("blocks"):
                    if len(item) < 7:
                        continue
                    x0, y0, x1, y1, text, _block_no, block_type = item[:7]
                    if int(block_type) != 0:
                        continue
                    text = (text or "").strip()
                    if not text:
                        continue
                    blocks.append(
                        {
                            "x0": float(x0),
                            "y0": float(y0),
                            "x1": float(x1),
                            "y1": float(y1),
                            "text": text,
                        }
                    )

                if not blocks:
                    return None

                blocks.sort(key=lambda b: (b["y0"], b["x0"]))
                rows = _group_pdf_blocks_into_rows(blocks, page_height)

                return {
                    "page_width": page_width,
                    "page_height": page_height,
                    "rows": rows,
                }
    except Exception as exc:
        print(f"PDF layout extract error for doc {doc.id} page {page_num}: {exc}")
        return None


def translate_layout_rows(layout_rows, target_lang="en"):
    translated_rows = []
    for row in layout_rows:
        translated_cells = []
        for cell in row["cells"]:
            source_text = cell.get("text") or ""
            translated = ""
            if source_text:
                normalized_source = source_text.replace("\r\n", "\n")
                source_lines = normalized_source.split("\n")
                if len(source_lines) > 1:
                    translated_lines = []
                    for line in source_lines:
                        line_text = line.strip()
                        if not line_text:
                            translated_lines.append("")
                            continue
                        line_translated = translate_to_language(line_text, target_lang)
                        translated_lines.append(line_translated or line_text)
                    translated = "\n".join(translated_lines).strip()
                else:
                    translated = translate_to_language(source_text, target_lang) or ""
            translated_cells.append(
                {
                    **cell,
                    "translation": translated or source_text,
                }
            )
        translated_rows.append({"y0": row["y0"], "cells": translated_cells})
    return translated_rows


def _normalize_translation_cell_text(text):
    raw = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = []
    previous_blank = False
    for line in raw.split("\n"):
        normalized = re.sub(r"\s+", " ", line).strip()
        if not normalized:
            if previous_blank:
                continue
            lines.append("")
            previous_blank = True
            continue
        lines.append(normalized)
        previous_blank = False
    return "\n".join(lines).strip()


def build_translation_flow_html(translated_rows, page_width=1.0):
    if not translated_rows:
        return ""

    page_width = max(float(page_width or 1.0), 1.0)
    parts = ['<div class="translation-layout-flow">']
    for row in translated_rows:
        cells = row["cells"]
        if len(cells) >= 4:
            parts.append('<table class="translation-mini-table"><tr>')
            for cell in cells:
                text = html.escape(_normalize_translation_cell_text(cell.get("translation") or ""))
                parts.append(f"<td>{text}</td>")
            parts.append("</tr></table>")
            continue

        parts.append('<div class="translation-flow-row">')
        for cell in cells:
            text = html.escape(_normalize_translation_cell_text(cell.get("translation") or ""))
            css_class = _translation_cell_class(cell, cells, page_width)
            parts.append(f'<div class="{css_class}">{text}</div>')
        parts.append("</div>")
    parts.append("</div>")
    return "".join(parts)


def build_layout_translation_payload(doc, page_num, target_lang="en"):
    layout = extract_pdf_page_layout_rows(doc, page_num)
    if not layout or not layout.get("rows"):
        return None

    translated_rows = translate_layout_rows(layout["rows"], target_lang=target_lang)
    source = "\n".join(
        " | ".join(cell.get("text") or "" for cell in row["cells"])
        for row in layout["rows"]
    ).strip()
    translation = "\n".join(
        " | ".join(cell.get("translation") or "" for cell in row["cells"])
        for row in translated_rows
    ).strip()
    translation_html = build_translation_flow_html(
        translated_rows, page_width=layout.get("page_width") or 1.0
    )
    return {
        "source": source,
        "translation": translation,
        "translation_html": translation_html,
    }


def normalize_translation_lang(lang):
    return "en"


def get_or_translate_page(doc, page_num, target_lang="en"):
    target_lang = normalize_translation_lang(target_lang)
    row = DocumentPageTranslation.query.filter_by(
        document_id=doc.id, page=page_num
    ).first()

    layout_payload = None
    if doc.extension.lower() == "pdf":
        layout_payload = build_layout_translation_payload(doc, page_num, target_lang=target_lang)

    cached_translation = ""
    if row:
        cached_translation = (
            row.translated_text_es if target_lang == "es" else row.translated_text
        ) or ""
    if cached_translation:
        cached_source = row.source_text if row else ""
        cached_provider = row.provider if row else get_translation_provider()
        response = {
            "page": page_num,
            "lang": target_lang,
            "source": cached_source,
            "translation": cached_translation,
            "provider": cached_provider,
            "cached": True,
        }
        if layout_payload:
            response["translation_html"] = layout_payload.get("translation_html") or ""
            response["layout"] = True
        return response

    if layout_payload:
        source = layout_payload.get("source") or ""
        translated = layout_payload.get("translation") or ""
        translation_html = layout_payload.get("translation_html") or ""
    else:
        source = extract_single_page_text(doc, page_num)
        translated = ""
        translation_html = ""

    if not source:
        if pdf_source_file_missing(doc):
            error = (
                "Source PDF not found on server. "
                "Re-upload the file in Documents, then click Reindex."
            )
        else:
            error = "No extractable text on this page."
        return {
            "page": page_num,
            "lang": target_lang,
            "source": "",
            "translation": "",
            "provider": get_translation_provider(),
            "cached": False,
            "error": error,
        }

    if not layout_payload:
        translated = translate_to_language(source, target_lang)

    provider = get_translation_provider()
    if translated:
        if row:
            row.source_text = source[:20000]
            if target_lang == "es":
                row.translated_text_es = translated
            else:
                row.translated_text = translated
            row.provider = provider
        else:
            dpt_cls = cast(Any, DocumentPageTranslation)
            db.session.add(
                dpt_cls(
                    document_id=doc.id,
                    page=page_num,
                    source_text=source[:20000],
                    translated_text=translated if target_lang == "en" else "",
                    translated_text_es=translated if target_lang == "es" else "",
                    provider=provider,
                )
            )
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()

    response = {
        "page": page_num,
        "lang": target_lang,
        "source": source,
        "translation": translated,
        "provider": provider,
        "cached": False,
    }
    if layout_payload:
        response["translation_html"] = translation_html
        response["layout"] = True
    return response


def count_pdf_pages(doc):
    offsets = load_page_offsets(doc)
    if offsets:
        return max(page for _, _, page in offsets)

    if doc.extension != "pdf" or not storage.document_exists(doc.stored_filename, DOC_FOLDER):
        return 0

    if not OCR_AVAILABLE or fitz is None:
        return 0

    try:
        with storage.open_document_local_path(doc.stored_filename, DOC_FOLDER) as file_path:
            with fitz.open(file_path) as pdf_doc:
                return len(pdf_doc)
    except Exception as exc:
        print(f"PDF page count error for doc {doc.id}: {exc}")
        return 0


def count_document_pages(doc):
    ext = doc.extension.lower()
    if ext in {"docx", "txt"}:
        return max(1, len(split_text_into_virtual_pages(get_document_full_text(doc))))
    if ext == "pdf":
        total = count_pdf_pages(doc)
        return total if total > 0 else 0
    return 0


def _split_text_for_pages(text, max_chars=3200):
    text = text or ""
    if len(text) <= max_chars:
        return [text] if text else [""]

    parts = []
    start = 0
    while start < len(text):
        end = min(len(text), start + max_chars)
        if end < len(text):
            split_at = text.rfind("\n", start, end)
            if split_at > start + 400:
                end = split_at + 1
        parts.append(text[start:end])
        start = end
    return parts


def _append_translation_text(pdf_doc, header, text):
    if fitz is None:
        raise RuntimeError("PyMuPDF is required to export translated PDFs.")
    page_width, page_height = fitz.paper_size("a4")
    margin_x = 50
    margin_top = 70
    margin_bottom = 50
    body_rect = fitz.Rect(
        margin_x, margin_top, page_width - margin_x, page_height - margin_bottom
    )
    chunks = _split_text_for_pages(text)

    for index, chunk in enumerate(chunks):
        page = pdf_doc.new_page(width=page_width, height=page_height)
        if index == 0 and header:
            page.insert_text(
                (margin_x, 42),
                header,
                fontsize=11,
                fontname="helv",
            )
        page.insert_textbox(
            body_rect,
            chunk,
            fontsize=10,
            fontname="helv",
            align=fitz.TEXT_ALIGN_LEFT,
        )


def build_translated_pdf_bytes(doc, target_lang="en"):
    if not OCR_AVAILABLE:
        raise RuntimeError("PyMuPDF is required to export translated PDFs.")
    if fitz is None:
        raise RuntimeError("PyMuPDF is required to export translated PDFs.")

    if not translation.translation_enabled():
        raise RuntimeError("Translation is disabled.")

    target_lang = normalize_translation_lang(target_lang)
    lang_label = translation.language_label(target_lang)

    total_pages = count_document_pages(doc)
    if total_pages <= 0:
        raise ValueError("Could not determine document page count.")

    export_limit = MAX_PDF_PAGES if MAX_PDF_PAGES > 0 else total_pages
    pages_to_export = min(total_pages, export_limit)
    truncated = pages_to_export < total_pages

    pdf_out = fitz.open()
    source_name = doc.original_filename or doc.stored_filename

    try:
        if truncated:
            notice = (
                f"{lang_label} translation export for {source_name}\n"
                f"Pages 1-{pages_to_export} of {total_pages}\n\n"
            )
            _append_translation_text(pdf_out, "Export notice", notice)

        for page_num in range(1, pages_to_export + 1):
            result = get_or_translate_page(doc, page_num, target_lang=target_lang)
            translation_text = (result.get("translation") or "").strip()
            if not translation_text:
                translation_text = "(No translation available for this page.)"

            page_word = "Page" if doc.extension == "pdf" else "Section"
            header = f"{source_name} - {page_word} {page_num} ({lang_label})"
            _append_translation_text(pdf_out, header, translation_text)

        if pdf_out.page_count == 0:
            raise ValueError("No translated pages were generated.")

        buffer = io.BytesIO()
        pdf_out.save(buffer, garbage=4, deflate=True)
        buffer.seek(0)
        return buffer.getvalue()
    finally:
        pdf_out.close()


def translate_to_language(text, target_lang="en"):
    return translation.translate_text(
        text,
        target_lang=normalize_translation_lang(target_lang),
        cache_get=get_cached_translation,
        cache_set=store_cached_translation,
    )


def translate_to_english(text):
    return translate_to_language(text, "en")


def build_translated_snippet(snippet_html, target_lang="en"):
    plain = strip_html(snippet_html)
    translated = translate_to_language(plain, target_lang)
    if not translated or translated == plain:
        return ""
    return highlight_terms(translated, [])


def build_english_snippet(snippet_html):
    return build_translated_snippet(snippet_html, "en")


def enrich_result_with_translation(result_dict):
    result_dict["snippet_en"] = build_translated_snippet(result_dict.get("snippet", ""), "en")
    return result_dict


def strip_html(text):
    return re.sub(r"<[^>]+>", "", text or "")

def highlight_terms(text, terms):
    if not text:
        return ""
    safe = text
    if terms:
        pattern = re.compile("|".join(re.escape(t) for t in sorted(set(terms), key=len, reverse=True)), re.IGNORECASE)
        safe = pattern.sub(lambda m: f'<span class="highlight">{m.group(0)}</span>', safe)
    return safe

def get_page_number(page_offsets, char_pos):
    if not page_offsets:
        return 1
    for start, end, page_num in page_offsets:
        if start <= char_pos < end:
            return page_num
    if char_pos >= page_offsets[-1][0]:
        return page_offsets[-1][2]
    return 1


def load_page_offsets(doc):
    raw = getattr(doc, "page_offsets_json", None) or ""
    if not raw:
        return []
    try:
        data = json.loads(raw)
        return [(int(a), int(b), int(c)) for a, b, c in data]
    except Exception:
        return []


def find_hit_position(text, expanded_tokens, query=""):
    text_lower = (text or "").lower()
    q = (query or "").lower().strip()
    if q and q in text_lower:
        return text_lower.find(q)

    exact_terms = default_exact_search_terms(query)
    for term in exact_terms:
        match = re.search(r"\b" + re.escape(term.lower()) + r"\b", text_lower)
        if match:
            return match.start()

    hit_pos = -1
    for term in expanded_tokens:
        pos = text_lower.find(term.lower())
        if pos != -1 and (hit_pos == -1 or pos < hit_pos):
            hit_pos = pos
    return hit_pos


def compact_display_text(text):
    if not text:
        return ""
    normalized = re.sub(r"\r\n?", "\n", text)
    normalized = re.sub(r"[ \t]+\n", "\n", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def resolve_result_page(doc, char_start, body_text, expanded_tokens, query, fallback_page=1):
    offsets = load_page_offsets(doc)
    if not offsets:
        return fallback_page or 1

    hit_in_body = find_hit_position(body_text, expanded_tokens, query)
    anchor = (char_start or 0) + hit_in_body if hit_in_body >= 0 else (char_start or 0)
    if anchor < 0:
        return fallback_page or 1
    return get_page_number(offsets, anchor)

def file_ext(filename):
    return filename.rsplit(".", 1)[1].lower() if "." in filename else ""

def current_time_str():
    return datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")

# =========================================================
# PDF page image helper
# =========================================================
def get_page_image(document_path, page_num=1):
    """
    Convert a single PDF page to PNG image bytes.
    Returns (bytes, mimetype) or None if conversion fails.
    """
    # First try pdf2image (requires poppler).
    if PDF2IMAGE_AVAILABLE and convert_from_path is not None:
        try:
            images = convert_from_path(
                document_path,
                first_page=page_num,
                last_page=page_num,
                dpi=150,
            )
            if images:
                img_io = io.BytesIO()
                images[0].save(img_io, format="PNG")
                img_io.seek(0)
                return img_io.getvalue(), "image/png"
        except Exception as e:
            print(f"pdf2image page render error: {document_path} page {page_num} -> {e}")

    # Fallback: render with PyMuPDF, which doesn't need poppler.
    if OCR_AVAILABLE and fitz is not None:
        try:
            with fitz.open(document_path) as pdf_doc:
                if page_num < 1 or page_num > len(pdf_doc):
                    return None
                page = pdf_doc.load_page(page_num - 1)
                pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
                return pix.tobytes("png"), "image/png"
        except Exception as e:
            print(f"fitz page render error: {document_path} page {page_num} -> {e}")

    return None


def list_docx_image_entries(document_path):
    try:
        with zipfile.ZipFile(document_path) as docx_zip:
            image_entries = [
                name
                for name in docx_zip.namelist()
                if name.startswith("word/media/")
                and file_ext(name) in {"png", "jpg", "jpeg", "gif", "bmp", "webp"}
            ]
        return sorted(image_entries)
    except Exception as exc:
        print(f"DOCX image list error: {document_path} -> {exc}")
        return []


def get_docx_image(document_path, image_index=0):
    image_entries = list_docx_image_entries(document_path)
    if image_index < 0 or image_index >= len(image_entries):
        return None

    entry_name = image_entries[image_index]
    try:
        with zipfile.ZipFile(document_path) as docx_zip:
            image_bytes = docx_zip.read(entry_name)
        mime_type = mimetypes.guess_type(entry_name)[0] or "application/octet-stream"
        return image_bytes, mime_type
    except Exception as exc:
        print(f"DOCX image read error: {document_path} {entry_name} -> {exc}")
        return None


def get_best_docx_image_index(document_path):
    image_entries = list_docx_image_entries(document_path)
    if not image_entries:
        return None

    best_idx = 0
    best_area = 0
    for idx, entry_name in enumerate(image_entries):
        try:
            with zipfile.ZipFile(document_path) as docx_zip:
                raw = docx_zip.read(entry_name)
            with Image.open(io.BytesIO(raw)) as img:
                width, height = img.size
                area = max(0, int(width)) * max(0, int(height))
            if area > best_area:
                best_area = area
                best_idx = idx
        except Exception:
            continue
    return best_idx


def get_top_docx_image_indexes(document_path, max_images=4, min_area=50000):
    image_entries = list_docx_image_entries(document_path)
    if not image_entries:
        return []

    scored = []
    for idx, entry_name in enumerate(image_entries):
        try:
            with zipfile.ZipFile(document_path) as docx_zip:
                raw = docx_zip.read(entry_name)
            with Image.open(io.BytesIO(raw)) as img:
                width, height = img.size
                area = max(0, int(width)) * max(0, int(height))
            if area >= min_area:
                scored.append((area, idx))
        except Exception:
            continue

    scored.sort(reverse=True)
    return [idx for _area, idx in scored[:max_images]]


def _read_docx_relationships(docx_zip, rels_path):
    try:
        root = ET.fromstring(docx_zip.read(rels_path))
    except Exception:
        return {}
    rels = {}
    for rel in root:
        rel_id = rel.attrib.get("Id")
        target = rel.attrib.get("Target", "")
        if rel_id and target:
            rels[rel_id] = target
    return rels


def _docx_target_to_media_path(target):
    target = (target or "").lstrip("/")
    if target.startswith("word/"):
        return target
    return f"word/{target}"


def _docx_paragraph_text_and_images(paragraph_elem):
    texts = []
    image_rids = []
    has_page_break = False
    for elem in paragraph_elem.iter():
        tag = elem.tag
        if tag == f"{DOCX_W_NS}t":
            if elem.text:
                texts.append(elem.text)
            if elem.tail:
                texts.append(elem.tail)
        elif tag == f"{DOCX_W_NS}br":
            if elem.attrib.get(f"{DOCX_W_NS}type") == "page":
                has_page_break = True
        elif tag == f"{DOCX_A_NS}blip":
            rid = elem.attrib.get(f"{DOCX_R_NS}embed") or elem.attrib.get(f"{DOCX_R_NS}link")
            if rid:
                image_rids.append(rid)
    return "".join(texts).strip(), image_rids, has_page_break


def _docx_table_text(table_elem):
    parts = []
    for text_elem in table_elem.iter(f"{DOCX_W_NS}t"):
        if text_elem.text:
            parts.append(text_elem.text)
        if text_elem.tail:
            parts.append(text_elem.tail)
    return "".join(parts).strip()


def _is_drawing_heading(text):
    normalized = re.sub(r"\s+", " ", (text or "").strip().lower())
    return normalized in DRAWING_MARKERS or normalized.startswith("drawing ") or normalized.startswith("tekening ")


def _docx_image_area(document_path, image_index, min_area=50000):
    image_entries = list_docx_image_entries(document_path)
    if image_index < 0 or image_index >= len(image_entries):
        return 0
    try:
        with zipfile.ZipFile(document_path) as docx_zip:
            raw = docx_zip.read(image_entries[image_index])
        with Image.open(io.BytesIO(raw)) as img:
            width, height = img.size
            area = max(0, int(width)) * max(0, int(height))
        return area if area >= min_area else 0
    except Exception:
        return 0


def collect_docx_drawing_image_candidates(document_path, min_area=50000):
    candidates = []
    try:
        with zipfile.ZipFile(document_path) as docx_zip:
            rels = _read_docx_relationships(docx_zip, "word/_rels/document.xml.rels")
            media_entries = list_docx_image_entries(document_path)
            media_index = {name: idx for idx, name in enumerate(media_entries)}

            root = ET.fromstring(docx_zip.read("word/document.xml"))
            body = root.find(f".//{DOCX_W_NS}body")
            if body is None:
                return []

            body_children = [child for child in body if child.tag in {f"{DOCX_W_NS}p", f"{DOCX_W_NS}tbl"}]
            total_blocks = len(body_children)

            recent_text = []
            drawing_caption = ""
            blocks_since_drawing_marker = 999

            for block_index, child in enumerate(body_children):
                text = ""
                image_indexes = []
                has_page_break = False

                if child.tag == f"{DOCX_W_NS}p":
                    text, image_rids, has_page_break = _docx_paragraph_text_and_images(child)
                    for rid in image_rids:
                        target = rels.get(rid, "")
                        media_path = _docx_target_to_media_path(target)
                        if media_path in media_index:
                            image_indexes.append(media_index[media_path])
                elif child.tag == f"{DOCX_W_NS}tbl":
                    text = _docx_table_text(child)

                if text:
                    lower = text.lower()
                    if _is_drawing_heading(text) or any(marker in lower for marker in DRAWING_MARKERS):
                        blocks_since_drawing_marker = 0
                        drawing_caption = ""
                    elif blocks_since_drawing_marker < 8:
                        blocks_since_drawing_marker += 1
                        if blocks_since_drawing_marker <= 3 and not _is_drawing_heading(text):
                            drawing_caption = text
                    else:
                        blocks_since_drawing_marker = 999
                        drawing_caption = ""

                    recent_text.append(text)
                    if len(recent_text) > 10:
                        recent_text = recent_text[-10:]

                if has_page_break and blocks_since_drawing_marker < 8:
                    blocks_since_drawing_marker += 1

                if not image_indexes:
                    continue

                context = " ".join(recent_text)
                for image_index in image_indexes:
                    if _docx_image_area(document_path, image_index, min_area=min_area) <= 0:
                        continue
                    candidates.append(
                        {
                            "image_index": image_index,
                            "context": context,
                            "caption": drawing_caption,
                            "in_drawing_zone": blocks_since_drawing_marker <= 8,
                            "block_index": block_index,
                            "total_blocks": total_blocks,
                        }
                    )
    except Exception as exc:
        print(f"DOCX drawing candidate error: {document_path} -> {exc}")
    return candidates


def score_docx_drawing_candidate(candidate, query, active_terms):
    context = f"{candidate.get('context', '')} {candidate.get('caption', '')}".lower()
    if any(marker in context for marker in CONTENTS_MARKERS):
        return -1.0
    if not candidate.get("in_drawing_zone") and not any(marker in context for marker in DRAWING_MARKERS):
        return -1.0

    score = 4.0 if candidate.get("in_drawing_zone") else 0.0
    if any(marker in context for marker in DRAWING_MARKERS):
        score += 4.0

    q = (query or "").lower().strip()
    if q and q in context:
        score += 6.0

    for term in active_terms or []:
        term_lower = (term or "").lower().strip()
        if len(term_lower) > 1 and term_lower in context:
            score += 3.0

    query_tokens = [t for t in re.findall(r"[a-z0-9\-]+", q) if len(t) > 2]
    for token in query_tokens:
        if token in context:
            score += 1.5

    return score


def select_docx_drawing_images(document_path, query, active_terms, max_images=4, min_area=50000):
    candidates = collect_docx_drawing_image_candidates(document_path, min_area=min_area)
    if not candidates:
        return []

    scored = []
    seen = set()
    for candidate in candidates:
        image_index = candidate["image_index"]
        if image_index in seen:
            continue
        score = score_docx_drawing_candidate(candidate, query, active_terms)
        if score <= 0:
            continue
        seen.add(image_index)
        scored.append((score, image_index))

    scored.sort(reverse=True)
    if scored:
        return [idx for _score, idx in scored[:max_images]]
    return []
# =========================================================
# Text extraction
# =========================================================
def ensure_indexable_file(path):
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"Stored file not found: {file_path}")

    if MAX_INDEX_FILE_MB > 0:
        size_mb = file_path.stat().st_size / (1024 * 1024)
        if size_mb > MAX_INDEX_FILE_MB:
            raise ValueError(
                f"File is {size_mb:.1f} MB; this server limits indexing to "
                f"{MAX_INDEX_FILE_MB} MB. Split the PDF or upgrade the hosting plan."
            )


def extract_text_from_pdf_fitz(path, max_pages=0):
    if not OCR_AVAILABLE or fitz is None:
        return None

    full_text = ""
    page_offsets = []
    current = 0

    try:
        with fitz.open(path) as doc:
            total_pages = len(doc)
            page_limit = total_pages if max_pages <= 0 else min(total_pages, max_pages)
            for page_num in range(page_limit):
                page = doc.load_page(page_num)
                text = str(page.get_text("text") or "")
                start = current
                full_text += text + "\n"
                current += len(text) + 1
                page_offsets.append((start, current, page_num + 1))
        return full_text, page_offsets
    except Exception as e:
        print(f"PyMuPDF extract error: {path} -> {e}")
        return None


def extract_text_from_pdf(path):
    fitz_result = extract_text_from_pdf_fitz(path, MAX_PDF_PAGES)
    if fitz_result is not None:
        return fitz_result

    full_text = ""
    page_offsets = []
    current = 0

    try:
        with pdfplumber.open(path) as pdf:
            pages = pdf.pages
            if MAX_PDF_PAGES > 0:
                pages = pages[:MAX_PDF_PAGES]
            for page_num, page in enumerate(pages, start=1):
                text = page.extract_text() or ""
                start = current
                full_text += text + "\n"
                current += len(text) + 1
                end = current
                page_offsets.append((start, end, page_num))
    except Exception as e:
        print(f"PDF extract error: {path} -> {e}")
        return "", [(0, 0, 1)]

    return full_text, page_offsets

def extract_text_from_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return text, [(0, len(text), 1)]

def extract_text_from_docx(path):
    if not DOCX_AVAILABLE or DocxDocument is None:
        return "", [(0, 0, 1)]
    doc = DocxDocument(path)
    parts = []
    for paragraph in doc.paragraphs:
        value = (paragraph.text or "").strip()
        if value:
            parts.append(value)
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)
    text = "\n".join(parts)
    return text, [(0, len(text), 1)]

def extract_text_from_csv(path):
    try:
        df = pd.read_csv(path, dtype=str, keep_default_na=False, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(path, dtype=str, keep_default_na=False, encoding="latin1")
    df = clean_dataframe_for_display(df, max_rows=200)
    text = df.to_string(index=False)
    return text, [(0, len(text), 1)], [("CSV", df)]
    
def clean_dataframe_for_display(df, max_rows=50):
    if df is None:
        return pd.DataFrame()

    cleaned = df.copy()
    cleaned = cleaned.fillna("")
    cleaned = cleaned.astype(str)
    for bad in ("nan", "NaN", "None", "<NA>"):
        cleaned = cleaned.replace(bad, "")

    cleaned.columns = [
        (str(col).strip() if not str(col).startswith("Unnamed:") else f"Column {idx + 1}")
        for idx, col in enumerate(cleaned.columns)
    ]

    if not cleaned.empty:
        row_keep = cleaned.apply(
            lambda row: any(normalize_cell_text(value) for value in row), axis=1
        )
        cleaned = cleaned.loc[row_keep]
        col_keep = cleaned.apply(
            lambda col: any(normalize_cell_text(value) for value in col), axis=0
        )
        cleaned = cleaned.loc[:, col_keep]

    return cleaned.head(max_rows).reset_index(drop=True)


def extract_text_from_xlsx(path):
    workbook = pd.ExcelFile(path)
    all_parts = []
    tables = []
    for sheet_name in workbook.sheet_names:
        raw_df = pd.read_excel(workbook, sheet_name=sheet_name, header=None, dtype=object)
        df = normalize_excel_sheet(raw_df)
        if df.empty:
            continue
        tables.append((sheet_name, df))
        all_parts.append(f"[Sheet: {sheet_name}]")
        all_parts.append(df.to_string(index=False))
    text = "\n".join(all_parts)
    return text, [(0, len(text), 1)], tables


def normalize_excel_sheet(raw_df, max_rows=200):
    if raw_df is None or raw_df.empty:
        return pd.DataFrame()

    sheet = raw_df.fillna("").astype(str)
    for bad in ("nan", "NaN", "None", "<NA>"):
        sheet = sheet.replace(bad, "")

    if sheet.empty:
        return pd.DataFrame()

    # Trim surrounding blank canvas introduced by styled spreadsheets.
    non_empty_mask = sheet.apply(lambda col: col.map(lambda value: bool(normalize_cell_text(value))))
    if non_empty_mask.values.any():
        non_empty_rows = non_empty_mask.any(axis=1)
        non_empty_cols = non_empty_mask.any(axis=0)
        row_positions = [idx for idx, keep in enumerate(non_empty_rows.tolist()) if keep]
        col_positions = [idx for idx, keep in enumerate(non_empty_cols.tolist()) if keep]
        if row_positions and col_positions:
            row_start, row_end = row_positions[0], row_positions[-1]
            col_start, col_end = col_positions[0], col_positions[-1]
            sheet = sheet.iloc[row_start : row_end + 1, col_start : col_end + 1].reset_index(
                drop=True
            )

    scan_rows = min(len(sheet), 40)
    header_row_idx = max(
        range(scan_rows),
        key=lambda idx: sum(1 for value in sheet.iloc[idx].tolist() if normalize_cell_text(value)),
        default=0,
    )
    header_cells = [normalize_cell_text(value) for value in sheet.iloc[header_row_idx].tolist()]
    if sum(1 for value in header_cells if value) < 2:
        return finalize_display_dataframe(sheet, max_rows=max_rows)

    header_keys = {
        re.sub(r"[^a-z0-9]+", "", value.strip().lower())
        for value in header_cells
        if value.strip()
    }
    is_master_sheet = "speccode" in header_keys and "sourcedocument" in header_keys
    effective_max_rows = len(sheet) if is_master_sheet else max_rows

    data_end_row = min(len(sheet), header_row_idx + effective_max_rows + 5)
    block = sheet.iloc[header_row_idx:data_end_row]

    keep_col_indexes = []
    for col_idx in range(block.shape[1]):
        column_values = [normalize_cell_text(block.iloc[row_idx, col_idx]) for row_idx in range(len(block))]
        if any(column_values):
            keep_col_indexes.append(col_idx)

    if not keep_col_indexes:
        return pd.DataFrame()

    block = block.iloc[:, keep_col_indexes].copy()
    headers = []
    for col_idx, value in enumerate(block.iloc[0].tolist()):
        label = normalize_cell_text(value) or f"Column {col_idx + 1}"
        headers.append(label)

    body = block.iloc[1:].copy()
    body.columns = dedupe_column_names(headers)
    return finalize_display_dataframe(body, max_rows=effective_max_rows)

def ocr_pdf(path):
    if not TESSERACT_AVAILABLE or pytesseract is None or convert_from_path is None:
        return "", [(0, 0, 1)]

    full_text = ""
    page_offsets = []
    current = 0

    try:
        images = convert_from_path(path)
        for i, image in enumerate(images, start=1):
            text = pytesseract.image_to_string(image) or ""
            text = str(text)
            start = current
            full_text += text + "\n"
            current += len(text) + 1
            end = current
            page_offsets.append((start, end, i))
    except Exception as e:
        print(f"OCR error: {path} -> {e}")
        return "", [(0, 0, 1)]

    return full_text, page_offsets

def extract_text_by_extension(path):
    ext = file_ext(path)

    if ext == "pdf":
        text, page_offsets = extract_text_from_pdf(path)
        used_ocr = False
        merged_text = text

        if not OCR_ENABLED:
            return merged_text, page_offsets, [], used_ocr

        ocr_text, ocr_offsets = ocr_pdf(path)
        if ocr_text.strip():
            used_ocr = True
            merged_text = text + "\n\n[OCR_LAYER]\n" + ocr_text if text.strip() else ocr_text

            # if using merged text, page mapping can just use OCR offsets
            # or keep PDF offsets if native text is dominant
            if len(ocr_text) > len(text):
                page_offsets = ocr_offsets

        return merged_text, page_offsets, [], used_ocr


    if ext == "txt":
        text, page_offsets = extract_text_from_txt(path)
        return text, page_offsets, [], False

    if ext == "docx":
        text, page_offsets = extract_text_from_docx(path)
        return text, page_offsets, [], False

    if ext == "csv":
        text, page_offsets, tables = extract_text_from_csv(path)
        return text, page_offsets, tables, False

    if ext == "xlsx":
        text, page_offsets, tables = extract_text_from_xlsx(path)
        return text, page_offsets, tables, False

    if ext in {"png", "jpg", "jpeg"}:
        text, page_offsets, tables = extract_text_from_image(path)
        return text, page_offsets, tables, True

    return "", [(0, 0, 1)], [], False


# =========================================================
# Requirement parsing
# =========================================================
def find_definition_nearby(lines, start_index):
    for i in range(start_index, min(start_index + 8, len(lines))):
        line = lines[i].strip()
        low = line.lower()
        if low.startswith("definition") or low.startswith("definitie"):
            if i + 1 < len(lines):
                return lines[i + 1].strip()
        if "definition:" in low or "definitie:" in low:
            parts = re.split(r"definition:|definitie:", line, flags=re.IGNORECASE)
            if len(parts) > 1:
                return parts[1].strip()
    return ""

def normalize_requirement_text(content):
    """Repair common PDF line-break splits inside requirement IDs."""
    text = content or ""
    text = re.sub(
        r"([A-Z]{1,10}-Req-)\s+(\d+(?:\.\d+)?)",
        r"\1\2",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"([A-Z]{1,10}-Req-\d+(?:\.\d+)?)\s+-\s+",
        r"\1 - ",
        text,
        flags=re.IGNORECASE,
    )
    return text


SECTION_HEADER_REGEX = re.compile(
    r"(?m)^(\d+(?:\.\d+)*)\s+(.+)$"
)


def clean_section_title(title):
    value = (title or "").strip()
    value = re.sub(r"\.{2,}.*$", "", value).strip()
    value = re.sub(r"\s+\d+$", "", value).strip()
    return value


def parse_document_sections(content):
    sections = []
    for match in SECTION_HEADER_REGEX.finditer(content or ""):
        num = match.group(1).strip()
        title = clean_section_title(match.group(2))
        if not title or len(title) < 3 or len(title) > 200:
            continue
        if title.lower().startswith(("page ", "pagina ")):
            continue
        sections.append(
            {
                "char_start": match.start(),
                "num": num,
                "title": title,
                "label": f"{num} {title}",
            }
        )
    return sections


def get_major_section_for_position(sections, char_pos):
    major = ""
    for sec in sections:
        if sec["char_start"] > char_pos:
            break
        if re.fullmatch(r"\d+", sec["num"]):
            major = sec["label"]
    return major


def get_nearest_section_for_position(sections, char_pos):
    nearest = ""
    for sec in sections:
        if sec["char_start"] > char_pos:
            break
        nearest = sec["label"]
    return nearest


def parse_section_blocks(content, page_offsets):
    matches = list(SECTION_HEADER_REGEX.finditer(content))
    if not matches:
        return []

    doc_sections = parse_document_sections(content)
    blocks = []
    for i, match in enumerate(matches):
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
        block_text = content[start:end].strip()
        if len(block_text) < 40:
            continue

        section_num = match.group(1)
        section_title = clean_section_title(match.group(2))
        page = get_page_number(page_offsets, start)
        lines = [line.strip() for line in block_text.splitlines() if line.strip()]
        summary = ""
        for line in lines[1:6]:
            if len(line) > 25:
                summary = line
                break

        blocks.append({
            "requirement_id": f"Section-{section_num}",
            "title": section_title,
            "section": f"{section_num} {section_title}",
            "major_section": get_major_section_for_position(doc_sections, start),
            "page": page,
            "char_start": start,
            "category": detect_category(block_text, section_title),
            "definition": "",
            "summary": summary,
            "full_text": block_text,
            "token_blob": " ".join(preprocess(block_text)),
        })

    return blocks


def requirement_id_from_filename(filename):
    match = REQ_ID_REGEX.search(filename or "")
    return match.group(1) if match else ""


def normalize_requirement_id_key(req_id):
    return (req_id or "").strip().lower()


def is_trackable_am_req_id(req_id):
    text = (req_id or "").strip()
    if not text:
        return False
    match = REQ_ID_REGEX.fullmatch(text)
    if match:
        return True
    match = REQ_ID_REGEX.search(text)
    return bool(match and match.group(1).lower() == text.lower())


def requirement_block_keep_score(block):
    score = len(block.full_text or "")
    doc = block.document
    if doc:
        if is_tennet_requirements_document(doc):
            score += 10_000_000
        filename_req = requirement_id_from_filename(doc.original_filename)
        block_req = normalize_requirement_id_key(block.requirement_id)
        if filename_req and normalize_requirement_id_key(filename_req) == block_req:
            score += 5_000_000
        if not is_spreadsheet_document(doc):
            score += 50_000
    score += (block.id or 0) * 0.001
    return score


def dedupe_requirement_blocks_by_id():
    rows = RequirementBlock.query.filter(
        RequirementBlock.requirement_id.isnot(None),
        RequirementBlock.requirement_id != "",
    ).all()
    groups = {}
    for row in rows:
        if not is_trackable_am_req_id(row.requirement_id):
            continue
        key = normalize_requirement_id_key(row.requirement_id)
        groups.setdefault(key, []).append(row)

    removed = 0
    for group in groups.values():
        if len(group) <= 1:
            continue
        keeper = max(group, key=requirement_block_keep_score)
        for row in group:
            if row.id != keeper.id:
                db.session.delete(row)
                removed += 1
    if removed:
        db.session.commit()
    return removed


def build_spreadsheet_requirement_blocks(doc_record, text, tables):
    if is_tennet_requirements_document(doc_record):
        blocks = []
        seen_req_ids = set()
        for sheet_name, df in (tables or []):
            if df is None or df.empty or not is_requirement_master_table_dataframe(df):
                continue

            columns = [str(col) for col in df.columns]
            normalized_map = {
                re.sub(r"[^a-z0-9]+", "", col.strip().lower()): col
                for col in columns
                if col.strip()
            }

            spec_col = normalized_map.get("speccode")
            source_col = normalized_map.get("sourcedocument")
            title_nl_col = normalized_map.get("dutchtitle")
            title_en_col = normalized_map.get("englishtitle")
            desc_nl_col = normalized_map.get("dutchdescription")
            desc_en_col = normalized_map.get("englishdescription")
            standard_desc_col = normalized_map.get("referredstandarddescription")
            statement_col = normalized_map.get("amstatement")
            domain_col = normalized_map.get("vakgebied")
            phase_col = normalized_map.get("fase")

            if not spec_col:
                continue

            for row_idx, row in df.iterrows():
                req_id = cell_to_dictionary_text(row.get(spec_col))
                req_id = normalize_requirement_lookup_id(req_id)
                if not is_trackable_am_req_id(req_id):
                    continue

                req_key = normalize_requirement_id_key(req_id)
                if req_key in seen_req_ids:
                    continue
                seen_req_ids.add(req_key)

                source_document = cell_to_dictionary_text(row.get(source_col)) if source_col else ""
                title_nl = cell_to_dictionary_text(row.get(title_nl_col)) if title_nl_col else ""
                title_en = cell_to_dictionary_text(row.get(title_en_col)) if title_en_col else ""
                desc_nl = cell_to_dictionary_text(row.get(desc_nl_col)) if desc_nl_col else ""
                desc_en = cell_to_dictionary_text(row.get(desc_en_col)) if desc_en_col else ""
                standard_desc = cell_to_dictionary_text(row.get(standard_desc_col)) if standard_desc_col else ""
                am_statement = cell_to_dictionary_text(row.get(statement_col)) if statement_col else ""
                domain = cell_to_dictionary_text(row.get(domain_col)) if domain_col else ""
                phase = cell_to_dictionary_text(row.get(phase_col)) if phase_col else ""

                full_text_parts = []
                for col_name in columns:
                    value = cell_to_dictionary_text(row.get(col_name))
                    if value:
                        full_text_parts.append(f"{col_name}: {value}")
                full_text = "\n".join(full_text_parts)

                summary_parts = [part for part in [desc_nl, desc_en, am_statement] if part]
                section_parts = [part for part in [source_document, domain, phase] if part]
                title = title_nl or title_en or req_id

                row_number = int(cast(Any, row_idx)) + 1
                blocks.append(
                    {
                        "requirement_id": req_id,
                        "title": title,
                        "section": " · ".join(section_parts),
                        "major_section": str(sheet_name or ""),
                        "page": 1,
                        "char_start": row_number,
                        "category": "Requirement master",
                        "definition": standard_desc,
                        "summary": " | ".join(summary_parts),
                        "full_text": full_text,
                        "token_blob": " ".join(preprocess(full_text)),
                    }
                )
        if blocks:
            return blocks

    req_id = requirement_id_from_filename(doc_record.original_filename)
    if not req_id:
        return []

    sheet_names = [str(name) for name, _df in (tables or []) if name]
    filename = doc_record.original_filename or ""
    if sheet_names:
        shown = ", ".join(sheet_names[:6])
        if len(sheet_names) > 6:
            shown += f" (+{len(sheet_names) - 6} more)"
        summary = f"Spreadsheet table · Sheets: {shown}"
    else:
        summary = "Spreadsheet table"

    full_text = sanitize_text_for_db((text or "")[:20000])
    if not full_text.strip():
        full_text = sanitize_text_for_db(summary)

    return [
        {
            "requirement_id": req_id,
            "title": filename,
            "section": sheet_names[0] if sheet_names else "",
            "major_section": "",
            "page": 1,
            "char_start": 0,
            "category": "Table",
            "definition": "",
            "summary": summary,
            "full_text": full_text,
            "token_blob": " ".join(preprocess(full_text)),
        }
    ]


def parse_requirement_blocks(content, page_offsets):
    content = normalize_requirement_text(content)
    doc_sections = parse_document_sections(content)
    matches = list(REQ_ID_REGEX.finditer(content))
    blocks = []

    if matches:
        for i, m in enumerate(matches):
            start = m.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
            block_text = content[start:end].strip()
            req_id = m.group(1)
            page = get_page_number(page_offsets, start)

            lines = [line.strip() for line in block_text.splitlines() if line.strip()]
            title = ""
            summary = ""
            section = get_nearest_section_for_position(doc_sections, start)
            major_section = get_major_section_for_position(doc_sections, start)
            definition = ""

            if len(lines) >= 2:
                if req_id.lower() in lines[0].lower():
                    title = lines[1]
                else:
                    title = lines[0]

            for idx, line in enumerate(lines):
                low = line.lower()
                if not section and re.match(r"^\d+(?:\.\d+)*\s+", line):
                    section = clean_section_title(line)
                    break
                if low.startswith(("de ", "het ", "een ", "opdrachtnemer ", "the ")):
                    summary = line
                    definition = find_definition_nearby(lines, idx)
                    break

            if not summary:
                for line in lines[1:6]:
                    if len(line) > 20:
                        summary = line
                        break

            if not definition:
                definition = find_definition_nearby(lines, 0)

            category = detect_category(block_text, title)
            blocks.append({
                "requirement_id": req_id,
                "title": title,
                "section": section,
                "major_section": major_section,
                "page": page,
                "char_start": start,
                "category": category,
                "definition": definition,
                "summary": summary,
                "full_text": block_text,
                "token_blob": " ".join(preprocess(block_text)),
            })
        return blocks

    section_blocks = parse_section_blocks(content, page_offsets)
    if section_blocks:
        return section_blocks

    # Fallback: paragraph blocks
    doc_sections = parse_document_sections(content)
    parts = re.split(r"\n\s*\n+", content)
    cursor = 0
    for part in parts:
        text = part.strip()
        if not text:
            continue

        pos = content.find(part, cursor)
        if pos == -1:
            pos = cursor
        page = get_page_number(page_offsets, pos)
        lines = [x.strip() for x in text.splitlines() if x.strip()]
        title = lines[0][:160] if lines else ""
        category = detect_category(text, title)

        blocks.append({
            "requirement_id": "",
            "title": title,
            "section": get_nearest_section_for_position(doc_sections, pos),
            "major_section": get_major_section_for_position(doc_sections, pos),
            "page": page,
            "char_start": pos,
            "category": category,
            "definition": "",
            "summary": lines[1][:500] if len(lines) > 1 else title,
            "full_text": text,
            "token_blob": " ".join(preprocess(text)),
        })
        cursor = pos + len(part)

    return blocks


# =========================================================
# Database indexing
# =========================================================
def create_default_admin():
    admin_username = os.environ.get("ADMIN_USERNAME", "admin").strip() or "admin"
    admin_password = os.environ.get("ADMIN_PASSWORD", "").strip()
    admin_auto_create = env_flag("ADMIN_AUTO_CREATE", default=True)

    if not admin_auto_create:
        print("Admin bootstrap skipped because ADMIN_AUTO_CREATE is disabled.")
        return

    if not admin_password:
        print(
            "Admin bootstrap skipped. Set ADMIN_USERNAME and ADMIN_PASSWORD "
            "before first start to create the initial admin account."
        )
        return

    user = User.query.filter_by(username=admin_username).first()
    if not user:
        user_cls = cast(Any, User)
        user = user_cls(username=admin_username, role=ADMIN_ROLE)
        user.set_password(admin_password)
        db.session.add(user)
        db.session.commit()
        print(f"Created initial admin user '{admin_username}'.")

def normalize_cell_text(value):
    text = str(value or "").strip()
    text = text.replace("\x00", "")
    text = re.sub(r"[\u200B-\u200D\uFEFF]", "", text)
    text = text.replace("\xa0", " ")
    text = text.replace("\\n", " ").replace("\n", " ")
    return re.sub(r"\s+", " ", text).strip()


def sanitize_text_for_db(value):
    return str(value or "").replace("\x00", "").strip()


def _is_spacer_like_row(values):
    non_empty = [normalize_cell_text(v) for v in values if normalize_cell_text(v)]
    if not non_empty:
        return True
    if len(non_empty) == 1:
        token = non_empty[0]
        # Drop visual spacer artifacts such as lone symbols or short numeric markers.
        if re.fullmatch(r"[-_=|./\\]+", token):
            return True
        if not any(ch.isalpha() for ch in token) and len(token) <= 2:
            return True
    return False


def dedupe_column_names(columns):
    seen = {}
    result = []
    for col in columns:
        base = normalize_cell_text(col) or "Column"
        count = seen.get(base, 0) + 1
        seen[base] = count
        result.append(base if count == 1 else f"{base} ({count})")
    return result


def trim_sparse_display_columns(df, min_fill_ratio=0.05):
    if df is None or df.empty:
        return df
    keep_cols = []
    for idx, col in enumerate(df.columns):
        col_name = normalize_cell_text(col)
        series = df.iloc[:, idx]
        values = [normalize_cell_text(v) for v in series.tolist()]
        filled = sum(1 for value in values if value)
        if filled == 0:
            continue
        non_empty_values = [value for value in values if value]
        has_alpha_content = any(any(ch.isalpha() for ch in value) for value in non_empty_values)
        max_len = max((len(value) for value in non_empty_values), default=0)
        # Drop noise columns that only carry a single numeric marker (e.g. page index).
        if filled <= 1 and not has_alpha_content and max_len <= 8:
            continue
        if (col_name.startswith("Column ") or col_name.startswith("Unnamed:")) and (
            filled / max(len(values), 1) < min_fill_ratio
        ):
            continue
        keep_cols.append(col)
    if not keep_cols:
        return df
    return df.loc[:, keep_cols]


def finalize_display_dataframe(df, max_rows=40):
    if df is None or df.empty:
        return pd.DataFrame()

    cleaned = df.copy()
    cleaned.columns = dedupe_column_names(cleaned.columns)
    for idx in range(cleaned.shape[1]):
        cleaned.iloc[:, idx] = cleaned.iloc[:, idx].map(normalize_cell_text)

    cleaned = clean_dataframe_for_display(cleaned, max_rows=max_rows)
    cleaned = trim_sparse_display_columns(cleaned)
    if cleaned.empty:
        return cleaned

    keep_indexes = [
        idx
        for idx in range(cleaned.shape[1])
        if any(str(value).strip() for value in cleaned.iloc[:, idx].tolist())
    ]
    if keep_indexes:
        cleaned = cleaned.iloc[:, keep_indexes]

    # After sparse/noise columns are removed, some rows may become fully empty.
    if not cleaned.empty:
        row_keep = cleaned.apply(
            lambda row: any(normalize_cell_text(value) for value in row), axis=1
        )
        cleaned = cleaned.loc[row_keep]

    if not cleaned.empty:
        cleaned = cleaned.loc[
            cleaned.apply(lambda row: not _is_spacer_like_row(row.tolist()), axis=1)
        ]

    if cleaned.empty:
        return cleaned

    cleaned.columns = dedupe_column_names(cleaned.columns)
    return cleaned.head(max_rows).reset_index(drop=True)


def split_table_footnote_rows(df):
    if df is None or df.empty:
        return df, []

    footnotes = []
    keep_rows = []
    total_cols = max(len(df.columns), 1)
    for _, row in df.iterrows():
        values = [normalize_cell_text(v) for v in row.tolist()]
        non_empty = [v for v in values if v]
        if not non_empty:
            continue
        joined = " ".join(non_empty)
        sparse_row = len(non_empty) <= max(2, int(total_cols * 0.15))
        is_footnote = (
            (len(non_empty) == 1 and len(non_empty[0]) >= 70)
            or (len(non_empty) <= 2 and max(len(v) for v in non_empty) >= 100)
            or (sparse_row and len(joined) >= 80)
            or bool(re.match(r"^\d+\)\s", non_empty[0]))
        )
        if is_footnote:
            footnotes.append(joined)
        else:
            keep_rows.append(row)

    if keep_rows:
        trimmed = pd.DataFrame(keep_rows).reset_index(drop=True)
        trimmed.columns = dedupe_column_names(df.columns[: trimmed.shape[1]])
    else:
        trimmed = df.iloc[0:0].copy()
    return trimmed, footnotes


def prepare_table_display(full_df, footnotes_override=None):
    if full_df is None or full_df.empty:
        return "", pd.DataFrame(), []

    preview_df, footnotes = split_table_footnote_rows(full_df)
    if footnotes_override is not None:
        footnotes = footnotes_override
    html_table = preview_df.to_html(
        index=False, classes="data-table compact-table", border=0, na_rep=""
    )
    if footnotes:
        html_table += "".join(
            f'<div class="table-footnote">{html.escape(note)}</div>' for note in footnotes
        )
    return html_table, preview_df, footnotes


def build_table_html(df):
    full_max_rows = max(len(df.index) + 5, 40) if df is not None else 40
    full_df = finalize_display_dataframe(df, max_rows=full_max_rows)
    if full_df.empty:
        return "", pd.DataFrame()

    preview_df = finalize_display_dataframe(full_df, max_rows=40)
    html_table, _preview_df, _footnotes = prepare_table_display(preview_df)
    return html_table, full_df


def dataframe_from_table_csv(csv_text):
    if not csv_text or not str(csv_text).strip():
        return None
    try:
        return pd.read_csv(io.StringIO(csv_text), dtype=str, keep_default_na=False)
    except Exception:
        return None


def render_table_preview_html(table_row, lang="nl"):
    csv_text = (table_row.csv_text or "").strip()
    csv_text_en = (table_row.csv_text_en or "").strip()
    if lang == "en" and csv_text_en:
        df_en = dataframe_from_table_csv(csv_text_en)
        if df_en is not None and not df_en.empty:
            full_df_en = finalize_display_dataframe(df_en, max_rows=40)
            footnotes_override = None
            if csv_text:
                df_nl = dataframe_from_table_csv(csv_text)
                if df_nl is not None and not df_nl.empty:
                    _, original_footnotes = split_table_footnote_rows(
                        finalize_display_dataframe(df_nl, max_rows=40)
                    )
                    if original_footnotes and translation.translation_enabled():
                        footnotes_override = [
                            translate_to_language(note, "en") or note
                            for note in original_footnotes
                        ]
            html_table, _, _ = prepare_table_display(
                full_df_en, footnotes_override=footnotes_override
            )
            if html_table:
                return html_table
        return table_row.html_table_en or ""

    df = dataframe_from_table_csv(csv_text)
    if df is not None and not df.empty:
        html_table, _ = build_table_html(df)
        if html_table:
            return html_table

    if lang == "en":
        return table_row.html_table_en or ""
    return table_row.html_table or ""


def translate_preview_df(full_df, target_lang="en"):
    if full_df is None or full_df.empty or not translation.translation_enabled():
        return "", ""
    try:
        display_df = finalize_display_dataframe(full_df, max_rows=40)
        translated_df = translation.translate_dataframe_values(
            display_df,
            target_lang=target_lang,
            cache_get=get_cached_translation,
            cache_set=store_cached_translation,
            max_cells=TRANSLATE_MAX_CELLS,
        )
        translated_df = finalize_display_dataframe(translated_df, max_rows=40)
        _, original_footnotes = split_table_footnote_rows(display_df)
        translated_footnotes = None
        if original_footnotes:
            translated_footnotes = [
                translate_to_language(note, target_lang) or note for note in original_footnotes
            ]
        html_table, _preview_df, _footnotes = prepare_table_display(
            translated_df, footnotes_override=translated_footnotes
        )

        csv_buf = io.StringIO()
        translated_df.to_csv(csv_buf, index=False)
        return html_table, csv_buf.getvalue()
    except Exception as exc:
        print(f"Table translation skipped ({target_lang}): {exc}")
        return "", ""


def translate_preview_df_to_html(preview_df, target_lang="en"):
    html_table, _ = translate_preview_df(preview_df, target_lang)
    return html_table


def save_table_previews(doc_id, tables):
    doc = db.session.get(DocumentRecord, doc_id)
    skip_translation = should_skip_table_translation(doc)
    for sheet_name, df in tables:
        if df is None:
            continue
        html_table, full_df = build_table_html(df)
        if full_df.empty:
            continue
        html_table_en = ""
        csv_text_en = ""
        if translation.translation_enabled() and not skip_translation:
            html_table_en, csv_text_en = translate_preview_df(full_df, "en")
        csv_buf = io.StringIO()
        full_df.to_csv(csv_buf, index=False)
        table_preview_cls = cast(Any, TablePreview)
        row = table_preview_cls(
            document_id=doc_id,
            sheet_name=str(sheet_name),
            page=1,
            table_format="xlsx" if sheet_name != "CSV" else "csv",
            html_table=html_table,
            html_table_en=html_table_en,
            html_table_es="",
            csv_text=csv_buf.getvalue(),
            csv_text_en=csv_text_en,
            csv_text_es="",
            preview_title=f"{sheet_name} preview"
        )
        db.session.add(row)


def table_preview_needs_regenerate(tables):
    if not tables:
        return True
    for row in tables:
        html = row.html_table or ""
        if "NaN" in html or "Unnamed:" in html:
            return True
    return False


def table_preview_needs_translation(tables):
    if not translation.translation_enabled():
        return False
    for row in tables:
        if not (row.html_table_en or "").strip():
            return True
    return False


def refresh_table_translations_from_cache(tables, skip_translation=False):
    if skip_translation:
        return False
    updated = False
    for row in tables:
        needs_en = not (row.html_table_en or "").strip()
        if (
            not needs_en
            and "NaN" not in (row.html_table or "")
        ):
            continue
        if not (row.csv_text or "").strip():
            continue
        try:
            df = pd.read_csv(
                io.StringIO(row.csv_text), dtype=str, keep_default_na=False
            )
            html_table, full_df = build_table_html(df)
            if full_df.empty:
                continue
            row.html_table = html_table
            if needs_en:
                row.html_table_en, row.csv_text_en = translate_preview_df(full_df, "en")
            updated = True
        except Exception as exc:
            print(f"Cached table translation failed for preview {row.id}: {exc}")
    if updated:
        db.session.commit()
    return updated


def regenerate_table_previews(doc_record):
    ext = doc_record.extension.lower()
    if ext not in {"csv", "xlsx"}:
        return []

    if not storage.document_exists(doc_record.stored_filename, DOC_FOLDER):
        raise FileNotFoundError(
            f"Stored file not found: {doc_record.stored_filename}"
        )

    with storage.open_document_local_path(
        doc_record.stored_filename, DOC_FOLDER
    ) as full_path:
        ensure_indexable_file(full_path)
        if ext == "xlsx":
            _, _, tables = extract_text_from_xlsx(full_path)
        else:
            _, _, tables = extract_text_from_csv(full_path)

    if not tables:
        return []

    TablePreview.query.filter_by(document_id=doc_record.id).delete()
    save_table_previews(doc_record.id, tables)
    db.session.commit()
    return TablePreview.query.filter_by(document_id=doc_record.id).all()

def remove_existing_blocks(doc_id):
    RequirementBlock.query.filter_by(document_id=doc_id).delete()
    TablePreview.query.filter_by(document_id=doc_id).delete()
    db.session.commit()

def index_document_record(doc_record):
    try:
        _index_document_record_impl(doc_record)
    except Exception as exc:
        db.session.rollback()
        raise


def _index_document_record_impl(doc_record):
    if not storage.document_exists(doc_record.stored_filename, DOC_FOLDER):
        raise FileNotFoundError(
            f"Stored file not found: {doc_record.stored_filename}. "
            "The upload may have been lost after a service restart."
        )

    with storage.open_document_local_path(
        doc_record.stored_filename, DOC_FOLDER
    ) as full_path:
        ensure_indexable_file(full_path)
        text, page_offsets, tables, used_ocr = extract_text_by_extension(full_path)

    doc_record.is_ocr = used_ocr
    doc_record.text_preview = text[:TEXT_PREVIEW_LIMIT]
    doc_record.page_offsets_json = json.dumps(page_offsets)
    doc_record.status = "indexed"
    db.session.commit()

    remove_existing_blocks(doc_record.id)

    if is_spreadsheet_document(doc_record):
        blocks = build_spreadsheet_requirement_blocks(doc_record, text, tables)
    else:
        blocks = parse_requirement_blocks(text, page_offsets)
    seen_hashes = set()
    seen_req_ids = set()
    pending = 0

    for b in blocks:
        requirement_id = sanitize_text_for_db(b.get("requirement_id", ""))
        if requirement_id and is_trackable_am_req_id(requirement_id):
            req_key = normalize_requirement_id_key(requirement_id)
            if req_key in seen_req_ids:
                continue
            seen_req_ids.add(req_key)
        title = sanitize_text_for_db(b.get("title", ""))
        section = sanitize_text_for_db(b.get("section", ""))
        major_section = sanitize_text_for_db(b.get("major_section", ""))
        category = sanitize_text_for_db(b.get("category", ""))
        definition = sanitize_text_for_db(b.get("definition", ""))
        summary = sanitize_text_for_db(b.get("summary", ""))
        full_text = sanitize_text_for_db(b.get("full_text", ""))
        token_blob = sanitize_text_for_db(b.get("token_blob", ""))

        th = hashlib.sha256(
            full_text.encode("utf-8", errors="ignore")
        ).hexdigest()
        if th in seen_hashes:
            continue
        seen_hashes.add(th)

        req_block_cls = cast(Any, RequirementBlock)
        row = req_block_cls(
            document_id=doc_record.id,
            requirement_id=requirement_id,
            title=title,
            section=section,
            major_section=major_section,
            page=b["page"],
            char_start=b.get("char_start", 0),
            category=category,
            definition=definition,
            summary=summary,
            full_text=full_text,
            token_blob=token_blob,
            text_hash=th,
            semantic_text=" ".join([title, summary, definition, full_text[:2000]]),
        )
        db.session.add(row)
        pending += 1
        if pending >= INDEX_BATCH_SIZE:
            db.session.commit()
            pending = 0

    save_table_previews(doc_record.id, tables)
    db.session.commit()
    dedupe_requirement_blocks_by_id()

def dedupe_lookup(file_hash):
    return DocumentRecord.query.filter_by(file_hash=file_hash).first()


def restore_document_file_from_temp(doc, temp_path):
    size_bytes = storage.save_document(doc.stored_filename, temp_path, DOC_FOLDER)
    storage.verify_document_stored(doc.stored_filename, DOC_FOLDER)
    doc.size_bytes = size_bytes
    db.session.commit()
    return size_bytes


def format_storage_size(bytes_value):
    size = max(int(bytes_value or 0), 0)
    if size >= 1024 ** 3:
        return f"{size / (1024 ** 3):.2f} GB"
    if size >= 1024 ** 2:
        return f"{size / (1024 ** 2):.1f} MB"
    if size >= 1024:
        return f"{size / 1024:.0f} KB"
    return f"{size} B"


def build_storage_usage_summary():
    from sqlalchemy import func

    quota_bytes = max(STORAGE_QUOTA_MB, 1) * 1024 * 1024
    doc_bytes = int(
        db.session.query(func.coalesce(func.sum(DocumentRecord.size_bytes), 0)).scalar() or 0
    )
    dict_bytes = int(
        db.session.query(func.coalesce(func.sum(DictionarySource.size_bytes), 0)).scalar()
        or 0
    )
    used_bytes = max(doc_bytes + dict_bytes, 0)
    percent = (used_bytes / quota_bytes) * 100 if quota_bytes else 0.0
    remaining_bytes = max(quota_bytes - used_bytes, 0)
    return {
        "used_bytes": used_bytes,
        "quota_bytes": quota_bytes,
        "used_label": format_storage_size(used_bytes),
        "quota_label": format_storage_size(quota_bytes),
        "remaining_label": format_storage_size(remaining_bytes),
        "percent": round(min(percent, 999.9), 1),
        "doc_count": DocumentRecord.query.count(),
        "dict_count": DictionarySource.query.count(),
        "near_limit": percent >= 85,
        "over_limit": used_bytes > quota_bytes,
    }


# =========================================================
# Dictionary / reference lookup
# =========================================================
DICTIONARY_TERM_HEADERS = {
    "abreviation",
    "abbreviation",
    "abbrev",
    "abbr",
    "code",
    "document",
    "term",
    "key",
    "query",
    "title",
    "name",
}
DICTIONARY_CONTENT_HEADERS = {
    "english translation",
    "english",
    "translation",
    "content",
    "result",
    "definition",
    "meaning",
    "description",
}
DICTIONARY_NL_HEADERS = {
    "full name",
    "naam",
    "dutch",
    "nl",
    "omschrijving",
    "function",
    "beschrijving",
    "description nl",
}
DICTIONARY_ABBREV_TERM_HEADERS = {
    "abreviation",
    "abbreviation",
    "abbrev",
    "abbr",
    "code",
}
CABINET_CODE_RE = re.compile(r"=\__\+A", re.IGNORECASE)
CABINET_LOOSE_CODE_RE = re.compile(r"^A\d[\w]*$", re.IGNORECASE)
CABINET_SUFFIX_CODE_RE = re.compile(r"^\d+[A-Z]?$", re.IGNORECASE)


def normalize_dictionary_header(value):
    return re.sub(r"\s+", " ", str(value or "").strip().lower())


def cell_to_dictionary_text(value):
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    text = str(value).replace("\r\n", "\n").replace("\r", "\n").strip()
    return text


def find_dictionary_header_row(df_raw, sheet_name="", original_filename="", max_rows=8):
    if sheet_is_cabinet_data(df_raw, sheet_name, original_filename):
        return None
    for row_idx in range(min(max_rows, len(df_raw))):
        row_values = [
            normalize_dictionary_header(value) for value in df_raw.iloc[row_idx].tolist()
        ]
        has_term = any(value in DICTIONARY_TERM_HEADERS for value in row_values if value)
        has_content = any(
            value in DICTIONARY_CONTENT_HEADERS | DICTIONARY_NL_HEADERS
            for value in row_values
            if value
        )
        if has_term and has_content:
            return row_idx
    return 0


def is_cabinet_code(text, cabinet_context=False):
    value = (text or "").strip()
    if not value:
        return False
    if CABINET_CODE_RE.search(value):
        return True
    if re.search(r"__\+A", value, re.IGNORECASE):
        return True
    if cabinet_context and CABINET_LOOSE_CODE_RE.fullmatch(value):
        return True
    if cabinet_context and CABINET_SUFFIX_CODE_RE.fullmatch(value):
        return True
    return False


def normalize_cabinet_term(text):
    value = (text or "").strip()
    if not value:
        return value
    if CABINET_CODE_RE.search(value) or re.search(r"__\+A", value, re.IGNORECASE):
        if not value.startswith("=") and value.startswith("__+"):
            return f"={value}"
        return value
    if CABINET_SUFFIX_CODE_RE.fullmatch(value):
        return f"=__+A{value.upper()}"
    if CABINET_LOOSE_CODE_RE.fullmatch(value):
        return f"=__+{value.upper()}"
    return value


def sheet_is_glossary_data(df_raw, sheet_name=""):
    name = (sheet_name or "").lower()
    if any(token in name for token in ("acronim", "glossary", "abbrev", "acronym")):
        return True
    for row_idx in range(min(8, len(df_raw))):
        row_values = [
            normalize_dictionary_header(value) for value in df_raw.iloc[row_idx].tolist()
        ]
        if "abreviation" in row_values or "abbreviation" in row_values:
            return True
    return False


def sheet_is_cabinet_data(df_raw, sheet_name="", original_filename=""):
    if "cabinet" in (sheet_name or "").lower():
        return True
    if "cabinet" in Path(original_filename or "").stem.lower():
        return True
    for row_idx in range(min(8, len(df_raw))):
        for value in df_raw.iloc[row_idx].tolist():
            if is_cabinet_code(cell_to_dictionary_text(value), cabinet_context=True):
                return True
    return False


def build_cabinet_search_keys(term):
    text = normalize_cabinet_term((term or "").strip())
    keys = set()
    if not text:
        return []
    keys.update({text, text.lower(), text.upper()})
    for pattern in (r"\+A(.+)$", r"__\+A(.+)$"):
        match = re.search(pattern, text, re.IGNORECASE)
        if not match:
            continue
        suffix = match.group(1).strip()
        if not suffix:
            continue
        keys.update(
            {
                suffix,
                suffix.lower(),
                suffix.upper(),
                f"A{suffix}",
                f"a{suffix}",
                f"A{suffix}".upper(),
            }
        )
    return [key for key in keys if key]


def extract_cabinet_row_fields(row_values, cabinet_context=False):
    cells = [cell_to_dictionary_text(value) for value in (row_values or [])]
    if cells and normalize_dictionary_header(cells[0]) in DICTIONARY_TERM_HEADERS:
        return None

    term = ""
    content_nl = ""
    content_en = ""

    for cell in cells:
        if not cell:
            continue
        if normalize_dictionary_header(cell) == "cabinets":
            continue
        if not term and is_cabinet_code(cell, cabinet_context=cabinet_context):
            term = normalize_cabinet_term(cell)
            continue
        if not content_nl:
            content_nl = cell
        elif not content_en and cell != content_nl:
            content_en = cell
            break

    if term and content_nl:
        return term, content_nl, content_en
    return None


def pack_dictionary_search_keys(keys):
    normalized = sorted({(key or "").strip().lower() for key in keys if (key or "").strip()})
    if not normalized:
        return ""
    return "|" + "|".join(normalized) + "|"


def ensure_dictionary_english_text(nl_text, en_text="", translate_if_missing=True):
    english = (en_text or "").strip()
    dutch = (nl_text or "").strip()
    if english:
        return english
    if not dutch or not translate_if_missing:
        return ""
    if not translation.translation_enabled():
        return ""
    try:
        translated = translate_to_english(dutch)
        return (translated or "").strip()
    except Exception as exc:
        print(f"Dictionary translation error: {exc}")
        return ""


def detect_dictionary_columns(df):
    columns = list(df.columns)
    normalized = {normalize_dictionary_header(col): col for col in columns}

    term_col = None
    content_nl_col = None
    content_en_col = None
    for header in DICTIONARY_TERM_HEADERS:
        if header in normalized:
            term_col = normalized[header]
            break
    for header in DICTIONARY_NL_HEADERS:
        if header in normalized:
            content_nl_col = normalized[header]
            break
    for header in DICTIONARY_CONTENT_HEADERS:
        if header in normalized:
            content_en_col = normalized[header]
            break

    usable_columns = [
        col
        for col in columns
        if df[col].apply(lambda value: bool(cell_to_dictionary_text(value))).any()
    ]
    if not term_col and usable_columns:
        term_col = usable_columns[0]
    if not content_nl_col and len(usable_columns) >= 2:
        content_nl_col = usable_columns[1]
    if not content_en_col and len(usable_columns) >= 3:
        content_en_col = usable_columns[2]

    entry_kind = "reference"
    if any(header in DICTIONARY_ABBREV_TERM_HEADERS for header in normalized):
        entry_kind = "abbreviation"
    elif content_nl_col and content_en_col and term_col:
        entry_kind = "abbreviation"
    elif term_col and content_en_col and not content_nl_col:
        entry_kind = "abbreviation"
    elif term_col and content_nl_col and not content_en_col:
        sample_term = ""
        for _, row in df.head(5).iterrows():
            sample_term = cell_to_dictionary_text(row.get(term_col))
            if sample_term:
                break
        if is_cabinet_code(sample_term, cabinet_context=True):
            entry_kind = "cabinet"

    return term_col, content_nl_col, content_en_col, entry_kind


def build_dictionary_dataframe_from_raw(raw_df, header_row):
    header_values = raw_df.iloc[header_row].tolist()
    headers = []
    seen = {}
    for idx, value in enumerate(header_values):
        label = cell_to_dictionary_text(value) or f"Column {idx + 1}"
        base = label
        count = seen.get(base, 0)
        if count:
            label = f"{base}_{count + 1}"
        seen[base] = count + 1
        headers.append(label)

    data = raw_df.iloc[header_row + 1 :].copy()
    data = data.dropna(how="all")
    if data.empty:
        return pd.DataFrame()

    width = data.shape[1]
    data.columns = headers[:width]
    return data


def load_dictionary_sheet_matrix(file_path, sheet_name):
    try:
        from openpyxl import load_workbook
    except ImportError:
        return []

    wb = load_workbook(file_path, data_only=False, read_only=False)
    try:
        if sheet_name not in wb.sheetnames:
            return []
        ws = wb[sheet_name]
        matrix = []
        iter_rows_fn = getattr(ws, "iter_rows", None)
        if not callable(iter_rows_fn):
            return []
        for row in cast(Any, iter_rows_fn()):
            row_values = []
            for cell in row:
                value = cell.value
                if value is None and getattr(cell, "data_type", None) == "f":
                    value = cell.value
                row_values.append(cell_to_dictionary_text(value))
            if any(row_values):
                matrix.append(row_values)
        return matrix
    finally:
        wb.close()


def parse_cabinet_matrix(matrix, sheet_name="", original_filename=""):
    entries = []
    cabinet_context = sheet_is_cabinet_data(
        pd.DataFrame(matrix or []), sheet_name, original_filename
    )
    for row_idx, row_values in enumerate(matrix or []):
        parsed = extract_cabinet_row_fields(row_values, cabinet_context=cabinet_context)
        if not parsed:
            continue
        term, content_nl, content_en = parsed
        content_en = ensure_dictionary_english_text(
            content_nl, content_en, translate_if_missing=False
        )
        entries.append(
            {
                "term": term,
                "content": content_nl,
                "content_nl": content_nl,
                "content_en": content_en,
                "search_keys": pack_dictionary_search_keys(build_cabinet_search_keys(term)),
                "entry_kind": "cabinet",
                "row_number": row_idx + 1,
            }
        )
    return entries, "cabinet"


def parse_cabinet_sheet(df_raw, sheet_name="", original_filename=""):
    matrix = []
    for _, row in df_raw.iterrows():
        matrix.append([cell_to_dictionary_text(value) for value in row.tolist()])
    return parse_cabinet_matrix(matrix, sheet_name=sheet_name, original_filename=original_filename)


def parse_dictionary_sheet(df, entry_kind=None):
    df = df.copy()
    df = df.dropna(how="all")
    df = df.loc[:, ~df.columns.astype(str).str.match(r"^Unnamed", na=False)]
    df = df.dropna(axis=1, how="all")
    if df.empty:
        return [], entry_kind or "reference"

    term_col, content_nl_col, content_en_col, detected_kind = detect_dictionary_columns(df)
    if not term_col or not (content_nl_col or content_en_col):
        return [], entry_kind or detected_kind

    kind = entry_kind or detected_kind
    entries = []
    for row_idx, row in df.iterrows():
        term = cell_to_dictionary_text(row.get(term_col))
        content_nl = cell_to_dictionary_text(row.get(content_nl_col)) if content_nl_col else ""
        content_en = cell_to_dictionary_text(row.get(content_en_col)) if content_en_col else ""
        if not term and not content_nl and not content_en:
            continue
        if normalize_dictionary_header(term) in DICTIONARY_TERM_HEADERS:
            continue
        if not term:
            continue

        if kind == "reference":
            if content_en and not content_nl:
                content_nl = content_en
                content_en = ""
        elif kind == "cabinet":
            content_en = ensure_dictionary_english_text(
                content_nl, content_en, translate_if_missing=False
            )

        legacy_content = content_nl or content_en
        search_keys = (
            pack_dictionary_search_keys(build_cabinet_search_keys(term))
            if kind == "cabinet"
            else pack_dictionary_search_keys([term])
        )

        entries.append(
            {
                "term": term,
                "content": legacy_content,
                "content_nl": content_nl,
                "content_en": content_en,
                "search_keys": search_keys,
                "entry_kind": kind,
                "row_number": int(row_idx) + 1,
            }
        )
    return entries, kind


def import_dictionary_excel(file_path, original_filename):
    workbook = pd.ExcelFile(file_path)
    imported_sources = []
    skipped_sheets = []

    for sheet_name in workbook.sheet_names:
        raw_df = pd.read_excel(file_path, sheet_name=sheet_name, header=None, dtype=object)
        if raw_df.empty:
            skipped_sheets.append(sheet_name)
            continue

        matrix = load_dictionary_sheet_matrix(file_path, sheet_name)
        header_row = find_dictionary_header_row(
            raw_df, sheet_name=sheet_name, original_filename=original_filename
        )
        cabinet_sheet = sheet_is_cabinet_data(
            raw_df, sheet_name, original_filename
        )
        if cabinet_sheet:
            entries, entry_kind = parse_cabinet_matrix(
                matrix, sheet_name=sheet_name, original_filename=original_filename
            )
            if not entries and matrix:
                entries, entry_kind = parse_cabinet_matrix(
                    raw_df.values.tolist(),
                    sheet_name=sheet_name,
                    original_filename=original_filename,
                )
        elif header_row is None:
            entries, entry_kind = parse_cabinet_matrix(
                matrix or raw_df.values.tolist(),
                sheet_name=sheet_name,
                original_filename=original_filename,
            )
        else:
            df = build_dictionary_dataframe_from_raw(raw_df, header_row)
            forced_kind = "abbreviation" if sheet_is_glossary_data(raw_df, sheet_name) else None
            entries, entry_kind = parse_dictionary_sheet(df, entry_kind=forced_kind)
        if not entries:
            skipped_sheets.append(sheet_name)
            continue

        display_name = f"{Path(original_filename).stem}"
        if len(workbook.sheet_names) > 1:
            display_name = f"{display_name} ({sheet_name})"
        file_size = os.path.getsize(file_path)

        existing = DictionarySource.query.filter_by(
            original_filename=original_filename,
            sheet_name=sheet_name,
        ).first()
        if existing:
            DictionaryEntry.query.filter_by(source_id=existing.id).delete()
            source = existing
            source.name = display_name
            source.entry_kind = entry_kind
            source.size_bytes = file_size
            source.uploaded_at = datetime.utcnow()
            storage.save_document(existing.stored_filename, file_path, DICT_FOLDER)
        else:
            stored_filename = f"{uuid.uuid4().hex}.xlsx"
            source_cls = cast(Any, DictionarySource)
            source = source_cls(
                name=display_name,
                original_filename=original_filename,
                stored_filename=stored_filename,
                sheet_name=sheet_name,
                entry_kind=entry_kind,
                size_bytes=file_size,
            )
            db.session.add(source)
            db.session.flush()
            storage.save_document(stored_filename, file_path, DICT_FOLDER)

        for entry in entries:
            entry_cls = cast(Any, DictionaryEntry)
            db.session.add(
                entry_cls(
                    source_id=source.id,
                    term=entry["term"],
                    content=entry["content"],
                    content_nl=entry.get("content_nl", ""),
                    content_en=entry.get("content_en", ""),
                    search_keys=entry.get("search_keys", ""),
                    entry_kind=entry["entry_kind"],
                    row_number=entry["row_number"],
                )
            )

        source.entry_count = len(entries)
        imported_sources.append(
            {
                "source": source,
                "entry_count": len(entries),
                "entry_kind": entry_kind,
            }
        )

    db.session.commit()
    return imported_sources, skipped_sheets


def dictionary_entry_kind_label(entry_kind):
    return {
        "abbreviation": "Glossary",
        "cabinet": "Cabinets",
        "reference": "Project Documents",
    }.get(entry_kind, "Definition")


DICTIONARY_TYPE_OPTIONS = [
    {"id": "glossary", "label": "Glossary", "kinds": ["abbreviation"]},
    {"id": "project_documents", "label": "Project Documents", "kinds": ["reference"]},
    {"id": "cabinets", "label": "Cabinets", "kinds": ["cabinet"]},
]
DICTIONARY_TYPE_IDS = {option["id"] for option in DICTIONARY_TYPE_OPTIONS}


def resolve_dictionary_entry_kinds(selected_types):
    if not selected_types:
        return []
    kinds = []
    for option in DICTIONARY_TYPE_OPTIONS:
        if option["id"] in selected_types:
            kinds.extend(option["kinds"])
    return kinds


def search_dictionary_entries(query, limit=50, entry_kinds=None):
    q = (query or "").strip()
    if not q:
        return []
    if entry_kinds is not None and not entry_kinds:
        return []

    q_lower = q.lower()
    seen_ids = set()
    ranked = []

    def scoped_query():
        query_obj = DictionaryEntry.query
        if entry_kinds:
            query_obj = query_obj.filter(DictionaryEntry.entry_kind.in_(entry_kinds))
        return query_obj

    def add_matches(matches, score):
        for entry in matches:
            if entry.id in seen_ids:
                continue
            seen_ids.add(entry.id)
            ranked.append((score, entry))

    exact_matches = (
        scoped_query()
        .filter(db.func.lower(DictionaryEntry.term) == q_lower)
        .limit(limit)
        .all()
    )
    add_matches(exact_matches, 100)

    key_matches = (
        scoped_query()
        .filter(DictionaryEntry.search_keys.ilike(f"%|{q_lower}|%"))
        .limit(limit)
        .all()
    )
    add_matches(key_matches, 98)

    if len(q) <= 12:
        prefix_matches = (
            scoped_query()
            .filter(DictionaryEntry.term.ilike(f"{q}%"))
            .limit(limit)
            .all()
        )
        add_matches(prefix_matches, 80)

    term_contains = (
        scoped_query()
        .filter(DictionaryEntry.term.ilike(f"%{q}%"))
        .limit(limit)
        .all()
    )
    add_matches(term_contains, 60)

    for column in (
        DictionaryEntry.content_nl,
        DictionaryEntry.content_en,
        DictionaryEntry.content,
    ):
        content_contains = (
            scoped_query().filter(column.ilike(f"%{q}%")).limit(limit).all()
        )
        add_matches(content_contains, 40)

    ranked.sort(key=lambda item: (-item[0], item[1].term.lower(), item[1].id))
    return [entry for _, entry in ranked[:limit]]


# =========================================================
# Search
# =========================================================
def lexical_score(block, query, active_terms, exact_terms=None):
    score = 0.0
    q = query.lower().strip()
    exact_terms = exact_terms or default_exact_search_terms(query)
    synonym_terms = [t for t in active_terms if t.lower() not in {e.lower() for e in exact_terms}]

    doc = block.document
    filename = (doc.original_filename or "").lower()
    stored_filename = (doc.stored_filename or "").lower()
    filename_no_ext = filename.rsplit(".", 1)[0] if "." in filename else filename

    title = (block.title or "").lower()
    summary = (block.summary or "").lower()
    definition = (block.definition or "").lower()
    full_text = (block.full_text or "").lower()
    req_id = (block.requirement_id or "").lower()
    section = (block.section or "").lower()
    major_section = (getattr(block, "major_section", "") or "").lower()
    category = (block.category or "").lower()
    combined = " ".join([title, summary, definition, full_text, req_id, section, major_section, category])

    matched_active = any(term_matches_text(term, combined) for term in active_terms)
    if not matched_active and not any(
        term_matches_text(term, " ".join([filename, stored_filename, filename_no_ext]))
        for term in active_terms
    ):
        return 0.0

    # Strong exact filename matches
    if q and q == filename:
        score += 100
    if q and q == stored_filename:
        score += 100
    if q and q == filename_no_ext:
        score += 95

    # Partial filename matches
    if q and q in filename:
        score += 40
    if q and q in stored_filename:
        score += 35
    if q and q in filename_no_ext:
        score += 35

    # Exact query phrase / token matches (highest priority in content)
    if q and q in title:
        score += 45
    if q and q in summary:
        score += 35
    if q and q in definition:
        score += 35
    if q and q in full_text:
        score += 40
    if q and len(q.split()) > 1 and q in full_text:
        score += 15

    for term in exact_terms:
        if term_matches_text(term, title):
            score += 28
        if term_matches_text(term, summary):
            score += 22
        if term_matches_text(term, definition):
            score += 22
        if term_matches_text(term, full_text):
            score += 30
        if term_matches_text(term, section):
            score += 18
        if term_matches_text(term, major_section):
            score += 24
        if term_matches_text(term, req_id):
            score += 25
        if term_matches_text(term, category):
            score += 10

    # Existing ID / metadata matches for exact query
    if q and q == req_id:
        score += 30
    if q and q in section:
        score += 10
    if q and q in major_section:
        score += 14
    if q and q in category:
        score += 6

    query_tokens = preprocess(query)
    if len(query_tokens) > 1 and all(t in full_text for t in query_tokens):
        score += 10
    if len(query_tokens) > 1 and all(t in title for t in query_tokens):
        score += 8

    # Selected synonym / expanded terms (lower weight)
    for term in set(synonym_terms):
        if term_matches_text(term, filename) or term_matches_text(term, filename_no_ext):
            score += 8
        if term_matches_text(term, title):
            score += 5
        if term_matches_text(term, summary):
            score += 4
        if term_matches_text(term, definition):
            score += 4
        if term_matches_text(term, full_text):
            score += 4
        if term_matches_text(term, category):
            score += 2

    if any(t.startswith("aard") for t in active_terms) and "ground" in category:
        score += 2

    return score

def semantic_score(query, block):
    model = get_semantic_model()
    if not model:
        return 0.0
    try:
        q_emb = model.encode([query])[0]
        b_emb = model.encode([block.semantic_text[:3000]])[0]

        dot = float(sum(a * b for a, b in zip(q_emb, b_emb)))
        q_norm = math.sqrt(sum(a * a for a in q_emb))
        b_norm = math.sqrt(sum(a * a for a in b_emb))
        if q_norm == 0 or b_norm == 0:
            return 0.0
        return (dot / (q_norm * b_norm)) * 20.0
    except Exception:
        return 0.0

def get_result_snippet(block, active_terms, query=""):
    text = compact_display_text(block.full_text or "")
    hit_pos = find_hit_position(text, active_terms, query)

    window = 900
    if hit_pos == -1:
        snippet = text[:window]
        if len(text) > window:
            snippet += "..."
    else:
        start = max(0, hit_pos - 250)
        end = min(len(text), hit_pos + 650)
        snippet = text[start:end]
        if start > 0:
            snippet = "..." + snippet
        if end < len(text):
            snippet = snippet + "..."

    highlight = list(dict.fromkeys([query] + active_terms))
    return highlight_terms(snippet, highlight)


def get_text_snippet(text, expanded_tokens, query=""):
    text = compact_display_text(text or "")
    text_lower = text.lower()
    hit_pos = find_hit_position(text, expanded_tokens, query)

    window = 900
    if hit_pos == -1:
        snippet = text[:window]
        if len(text) > window:
            snippet += "..."
    else:
        start = max(0, hit_pos - 250)
        end = min(len(text), hit_pos + 650)
        snippet = text[start:end]
        if start > 0:
            snippet = "..." + snippet
        if end < len(text):
            snippet += "..."

    return highlight_terms(snippet, list(dict.fromkeys([query] + expanded_tokens)))


def score_document_text(doc, query, active_terms, exact_terms=None):
    q = query.lower().strip()
    exact_terms = exact_terms or default_exact_search_terms(query)
    synonym_terms = [t for t in active_terms if t.lower() not in {e.lower() for e in exact_terms}]
    text = doc.text_preview or ""
    if not text.strip():
        return 0

    text_lower = text.lower()
    filename = (doc.original_filename or "").lower()
    base_name = filename.rsplit(".", 1)[0] if "." in filename else filename
    score = 0.0

    if not any(
        term_matches_text(term, text)
        or term_matches_text(term, filename)
        or term_matches_text(term, base_name)
        for term in active_terms
    ):
        return 0.0

    if q and q in text_lower:
        score += 35
    if q and q in filename:
        score += 25
    if q and q in base_name:
        score += 20

    for term in exact_terms:
        if term_matches_text(term, text_lower):
            score += 20
        if term_matches_text(term, filename) or term_matches_text(term, base_name):
            score += 12

    matched_terms = 0
    for term in set(synonym_terms):
        if term_matches_text(term, text_lower):
            matched_terms += 1
            score += 4
        if term_matches_text(term, filename) or term_matches_text(term, base_name):
            score += 3

    query_tokens = preprocess(query)
    if len(query_tokens) > 1 and all(t in text_lower for t in query_tokens):
        score += 14

    return score


def count_exact_term_hits_simple(result_item, query, exact_terms):
    combined = " ".join(
        [
            result_item.get("title") or "",
            result_item.get("summary") or "",
            result_item.get("full_text") or "",
            result_item.get("filename") or "",
        ]
    )
    hits = 0
    q = (query or "").strip().lower()
    if q and q in combined.lower():
        hits += 3
    for term in exact_terms:
        if term_matches_text(term, combined):
            hits += 1
    return hits


def is_spreadsheet_document(doc):
    return (doc.extension or "").lower() in {"csv", "xlsx"}


def build_document_fallback_results(query, active_terms, seen_keys, top_k=10):
    exact_terms = default_exact_search_terms(query)
    results = []
    for doc in DocumentRecord.query.all():
        if is_spreadsheet_document(doc):
            continue
        score = score_document_text(doc, query, active_terms, exact_terms)
        if score <= 0:
            continue

        page = resolve_result_page(
            doc,
            0,
            doc.text_preview,
            active_terms,
            query,
            1,
        )
        key = (doc.id, page, doc.original_filename)
        if key in seen_keys:
            continue
        seen_keys.add(key)

        results.append(enrich_result_with_translation({
            "block_id": None,
            "document_id": doc.id,
            "filename": doc.original_filename,
            "stored_filename": doc.stored_filename,
            "page": page,
            "requirement_id": "",
            "title": doc.original_filename,
            "section": "Document match",
            "category": "Document",
            "definition": "",
            "summary": (doc.text_preview or "")[:240],
            "full_text": doc.text_preview or "",
            "snippet": get_text_snippet(doc.text_preview, active_terms, query),
            "relevance": round(score, 2),
            "is_pdf": doc.extension.lower() == "pdf",
            "is_docx": doc.extension.lower() == "docx",
            "is_txt": doc.extension.lower() == "txt",
            "has_table": doc.extension.lower() in {"csv", "xlsx"},
            "ocr_used": doc.is_ocr,
            "is_image": doc.extension.lower() in {"png", "jpg", "jpeg"},
            "is_document_fallback": True,
        }))

    results.sort(
        key=lambda item: (
            count_exact_term_hits_simple(item, query, exact_terms),
            item["relevance"],
        ),
        reverse=True,
    )
    return results[:top_k]


def grouped_search_summary(results):
    category_counts = Counter(r["category"] for r in results if r.get("category"))
    top_categories = category_counts.most_common(5)
    req_ids = [r["requirement_id"] for r in results if r.get("requirement_id")]
    return {
        "top_categories": top_categories,
        "top_requirement_ids": req_ids[:8]
    }

def search_requirements(query, top_k=30, active_terms=None):
    if not query.strip():
        return [], []

    all_expanded_terms = get_all_expanded_terms(query)
    exact_terms = default_exact_search_terms(query)
    if active_terms is None:
        active_terms = exact_terms
    else:
        active_terms = resolve_active_search_terms(query, active_terms)

    q = query.lower().strip()

    direct_docs = DocumentRecord.query.all()
    matched_doc_ids = []

    for d in direct_docs:
        original_name = (d.original_filename or "").lower()
        stored_name = (d.stored_filename or "").lower()
        base_name = original_name.rsplit(".", 1)[0] if "." in original_name else original_name

        if q in {original_name, stored_name, base_name}:
            matched_doc_ids.append(d.id)

    rows = RequirementBlock.query.all()

    scored = []
    for row in rows:
        score = lexical_score(row, query, active_terms, exact_terms)

        if row.document_id in matched_doc_ids:
            score += 120

        score += semantic_score(query, row)

        if score > 0:
            exact_hits = count_exact_term_hits(row, query, exact_terms)
            scored.append((row, score, exact_hits))

    scored.sort(key=lambda x: (x[2], x[1]), reverse=True)
    results = []
    seen_keys = set()

    for row, score, _exact_hits in scored[:top_k]:
        doc = row.document
        if is_spreadsheet_document(doc):
            continue
        seen_keys.add((doc.id, row.page or 1, row.requirement_id or row.title))
        display_page = resolve_result_page(
            doc,
            row.char_start or 0,
            row.full_text,
            active_terms,
            query,
            row.page,
        )
        results.append(enrich_result_with_translation({
            "block_id": row.id,
            "document_id": doc.id,
            "filename": doc.original_filename,
            "stored_filename": doc.stored_filename,
            "page": display_page,
            "requirement_id": row.requirement_id,
            "title": row.title,
            "section": row.section,
            "major_section": getattr(row, "major_section", "") or "",
            "category": row.category,
            "definition": row.definition,
            "summary": row.summary,
            "full_text": row.full_text,
            "snippet": get_result_snippet(row, active_terms, query),
            "relevance": round(score, 2),
            "is_pdf": doc.extension.lower() == "pdf",
            "is_docx": doc.extension.lower() == "docx",
            "is_txt": doc.extension.lower() == "txt",
            "has_table": doc.extension.lower() in {"csv", "xlsx"},
            "ocr_used": doc.is_ocr,
            "is_image": doc.extension.lower() in {"png", "jpg", "jpeg"},
            "is_document_fallback": False,
        }))

    if len(results) < top_k:
        fallback = build_document_fallback_results(
            query, active_terms, seen_keys, top_k=top_k - len(results)
        )
        results.extend(fallback)

    results.sort(
        key=lambda r: (
            count_exact_term_hits_simple(r, query, exact_terms),
            r["relevance"],
        ),
        reverse=True,
    )
    return results[:top_k], all_expanded_terms


def search_table_results(query, top_k=20, active_terms=None):
    if not query.strip():
        return []

    exact_terms = default_exact_search_terms(query)
    if active_terms is None:
        active_terms = exact_terms
    else:
        active_terms = resolve_active_search_terms(query, active_terms)

    q = query.lower().strip()
    results = []
    for row in TablePreview.query.all():
        doc = row.document
        if not doc:
            continue

        searchable_parts = [
            row.sheet_name or "",
            row.preview_title or "",
            row.csv_text or "",
            row.csv_text_en or "",
            strip_html_legacy(row.html_table or ""),
            strip_html_legacy(row.html_table_en or ""),
            doc.original_filename or "",
        ]
        combined = "\n".join(part for part in searchable_parts if part).lower()
        if not combined:
            continue

        if not any(
            term_matches_text(term, combined)
            for term in active_terms
        ) and not (q and q in combined):
            continue

        score = 0.0
        if q and q in combined:
            score += 35.0
        for term in exact_terms:
            if term_matches_text(term, combined):
                score += 20.0
        for term in set(active_terms):
            if term_matches_text(term, combined):
                score += 6.0
            if term_matches_text(term, row.sheet_name or ""):
                score += 8.0
            if term_matches_text(term, doc.original_filename or ""):
                score += 5.0

        results.append(
            {
                "table_id": row.id,
                "document_id": doc.id,
                "filename": doc.original_filename,
                "sheet_name": row.sheet_name or "-",
                "preview_title": row.preview_title or doc.original_filename,
                "table_format": row.table_format or "xlsx",
                "html_table": render_table_preview_html(row, "nl"),
                "html_table_en": "" if should_skip_table_translation(doc) else render_table_preview_html(row, "en"),
                "relevance": round(score, 2),
            }
        )

    results.sort(key=lambda item: item["relevance"], reverse=True)
    return results[:top_k]


def search_documents(query, top_k=10, active_terms=None):
    q = query.lower().strip()
    exact_terms = default_exact_search_terms(query)
    if active_terms is None:
        active_terms = exact_terms
    else:
        active_terms = resolve_active_search_terms(query, active_terms)
    scored = []

    for doc in DocumentRecord.query.all():
        score = score_document_text(doc, query, active_terms, exact_terms)
        filename = (doc.original_filename or "").lower()
        stored = (doc.stored_filename or "").lower()
        base = filename.rsplit(".", 1)[0] if "." in filename else filename

        if q == filename:
            score += 100
        if q == stored:
            score += 100
        if q == base:
            score += 95
        if q in filename:
            score += 40
        if q in stored:
            score += 35
        if q in base:
            score += 35

        if score > 0:
            scored.append((doc, score))

    scored.sort(key=lambda x: x[1], reverse=True)

    results = []
    for doc, score in scored:
        if len(results) >= top_k:
            break
        requirement_id = "-"
        block_rows = RequirementBlock.query.filter_by(document_id=doc.id).all()
        if block_rows:
            best_req = None
            best_score = -1.0
            for row in block_rows:
                row_req = (row.requirement_id or "").strip()
                if not row_req:
                    continue
                row_score = lexical_score(row, query, active_terms, exact_terms)
                if row_score > best_score:
                    best_score = row_score
                    best_req = row_req
            if best_req:
                requirement_id = best_req
            else:
                for row in block_rows:
                    row_req = (row.requirement_id or "").strip()
                    if row_req:
                        requirement_id = row_req
                        break

        page = resolve_result_page(
            doc,
            0,
            doc.text_preview or "",
            active_terms,
            query,
            1,
        )

        # Better virtual-page estimate for docx/txt (since page_offsets_json is not meaningful there)
        if doc.extension.lower() in {"docx", "txt"}:
            try:
                full_text = get_document_full_text(doc)
                hit_pos = find_hit_position(full_text, active_terms, query)
                chunks = split_text_into_virtual_pages(full_text)
                if hit_pos >= 0 and chunks:
                    cursor = 0
                    for idx, chunk in enumerate(chunks, start=1):
                        cursor += len(chunk) + 2
                        if hit_pos < cursor:
                            page = idx
                            break
            except Exception as exc:
                print(f"Virtual page estimate error for doc {doc.id}: {exc}")

        image_pages = []
        image_indexes = []
        if doc.extension.lower() == "pdf":
            image_pages = select_pdf_drawing_pages(
                doc,
                query=query,
                active_terms=active_terms,
                center_page=page or 1,
                max_images=4,
            )
        if doc.extension.lower() == "docx" and storage.document_exists(doc.stored_filename, DOC_FOLDER):
            try:
                with storage.open_document_local_path(doc.stored_filename, DOC_FOLDER) as file_path:
                    image_indexes = select_docx_drawing_images(
                        file_path,
                        query=query,
                        active_terms=active_terms,
                        max_images=4,
                    )
            except Exception as exc:
                print(f"DOCX image preview error for doc {doc.id}: {exc}")

        has_drawings = bool(image_pages or image_indexes)
        if not has_drawings:
            continue

        results.append(
            {
                "document_id": doc.id,
                "filename": doc.original_filename,
                "page": page,
                "requirement_id": requirement_id,
                "snippet": get_text_snippet(doc.text_preview or "", active_terms, query),
                "relevance": round(score, 2),
                "is_pdf": doc.extension.lower() == "pdf",
                "is_docx": doc.extension.lower() == "docx",
                "is_txt": doc.extension.lower() == "txt",
                "is_image": doc.extension.lower() in {"png", "jpg", "jpeg"},
                "image_pages": image_pages,
                "image_indexes": image_indexes,
            }
        )
    return results


# =========================================================
# Templates
# =========================================================
BASE_CSS = """
<style>
body { font-family: Arial, sans-serif; margin: 24px; background:#f5f7fb; color:#222; }
.container { max-width: 1280px; margin:auto; background:#fff; padding:24px; border-radius:12px; box-shadow:0 2px 8px rgba(0,0,0,0.08); }
.topbar { display:flex; justify-content:space-between; align-items:flex-start; gap:16px; flex-wrap:nowrap; }
.topbar-brand { flex:1 1 auto; min-width:0; }
.topbar-brand h1 { margin:0 0 6px 0; line-height:1.2; }
.actions { display:flex; gap:8px; flex-wrap:wrap; justify-content:flex-end; align-items:flex-start; flex:0 0 auto; margin-left:auto; }
input[type=text], input[type=password], input[type=file], select { padding:10px; font-size:15px; }
input[type=text] { width:68%; min-width:280px; }
button, .btn { padding:10px 16px; border:none; border-radius:6px; text-decoration:none; color:white; cursor:pointer; display:inline-block; }
button { background:#0069d9; }
.btn-green { background:#28a745; }
.btn-purple { background:#6f42c1; }
.btn-orange { background:#fd7e14; }
.btn-gray { background:#6c757d; }
.notice, .flash, .warning, .info { padding:12px; border-radius:6px; margin:12px 0; }
.notice { background:#fff3cd; border-left:4px solid #ffc107; }
.flash { background:#e9f7ef; border-left:4px solid #28a745; }
.flash-success { background:#e9f7ef; border-left:4px solid #28a745; }
.flash-error { background:#fdecea; border-left:4px solid #dc3545; }
.flash-warning { background:#fff3cd; border-left:4px solid #ffc107; }
.flash-info, .flash-message { background:#eef6ff; border-left:4px solid #339af0; }
.admin-feedback-panel { margin:16px 0; }
.admin-feedback-panel .flash { font-size:15px; font-weight:500; }
.dict-search-panel { margin:16px 0; padding:16px 18px; background:#f8fafc; border:1px solid #e2e8f0; border-radius:8px; }
.dict-result-card { margin-top:16px; padding:16px 18px; border:1px solid #e2e8f0; border-left:4px solid #6f42c1; border-radius:8px; background:#fff; }
.dict-result-kind { display:inline-block; font-size:12px; font-weight:700; text-transform:uppercase; letter-spacing:0.04em; color:#6f42c1; margin-bottom:8px; }
.dict-result-term { font-size:22px; font-weight:800; color:#1f2937; margin-bottom:8px; }
.dict-result-content { font-size:16px; line-height:1.6; color:#334155; white-space:pre-wrap; overflow-wrap:anywhere; }
.dict-result-lang { margin-top:8px; font-size:16px; line-height:1.6; color:#334155; white-space:pre-wrap; overflow-wrap:anywhere; }
.dict-result-lang-label { display:inline-block; min-width:2.6em; font-weight:800; color:#475569; }
.dict-result-meta { margin-top:10px; font-size:13px; color:#64748b; }
.dict-empty { margin-top:16px; padding:14px 16px; background:#fff3cd; border-left:4px solid #ffc107; border-radius:6px; }
.warning { background:#fdecea; border-left:4px solid #dc3545; }
.info { background:#eef6ff; border-left:4px solid #339af0; }
.result-item { margin-top:16px; padding:16px; border-left:4px solid #0069d9; background:#fafafa; border-radius:6px; }
.filename { font-size:18px; font-weight:bold; color:#b00020; }
.meta { color:#555; margin:8px 0; font-size:14px; }
.meta-prominent {
  display: flex;
  flex-wrap: wrap;
  gap: 8px 16px;
  margin: 10px 0 12px 0;
}
.meta-prominent .meta-item {
  font-size: 17px;
  color: #1f2937;
}
.meta-prominent .meta-item strong {
  font-weight: 800;
}
.meta-secondary {
  font-size: 14px;
  color: #475569;
  margin-top: 6px;
  line-height: 1.5;
}
.summary-box { margin-top:16px; padding:12px; background:#f8f9fa; border:1px solid #ececec; border-radius:6px; }
.stats-panel, .search-meta-panel, .search-filter-panel {
  margin-top:16px;
  padding:16px 18px;
  background:#eef1f5;
  border:1px solid #d8dee6;
  border-radius:10px;
}
.panel-title {
  font-size:15px;
  font-weight:700;
  color:#334155;
  margin:0 0 12px 0;
  padding-bottom:8px;
  border-bottom:1px solid #d8dee6;
}
.panel-section { margin-top:14px; }
.panel-section:first-of-type { margin-top:0; }
.panel-label {
  font-size:13px;
  font-weight:700;
  color:#475569;
  margin-bottom:8px;
  letter-spacing:0.02em;
  text-transform:uppercase;
}
.stats-grid {
  display:flex;
  flex-wrap:wrap;
  gap:12px;
}
.stat-item {
  flex:1 1 180px;
  background:#fff;
  border:1px solid #dbe1ea;
  border-radius:8px;
  padding:12px 14px;
}
.stat-label { display:block; font-size:12px; color:#64748b; margin-bottom:4px; }
.stat-value { display:block; font-size:22px; font-weight:700; color:#0f172a; }
.storage-usage-panel {
  margin:12px 0;
  padding:14px 16px;
  background:#eef6ff;
  border:1px solid #bfdbfe;
  border-radius:8px;
}
.storage-usage-panel.storage-usage-warn {
  background:#fff7ed;
  border-color:#fdba74;
}
.storage-usage-panel.storage-usage-danger {
  background:#fef2f2;
  border-color:#fca5a5;
}
.storage-usage-summary {
  margin-bottom:8px;
  font-size:15px;
  color:#1e293b;
}
.storage-usage-bar {
  height:10px;
  background:#dbeafe;
  border-radius:999px;
  overflow:hidden;
  margin:8px 0;
}
.storage-usage-warn .storage-usage-bar { background:#ffedd5; }
.storage-usage-danger .storage-usage-bar { background:#fee2e2; }
.storage-usage-fill {
  height:100%;
  background:#0069d9;
  border-radius:999px;
  max-width:100%;
}
.storage-usage-warn .storage-usage-fill { background:#fd7e14; }
.storage-usage-danger .storage-usage-fill { background:#dc3545; }
.storage-usage-note { color:#64748b; font-size:13px; }
.term-chip-list { display:flex; flex-wrap:wrap; gap:8px; }
.term-chip {
  display:inline-block;
  background:#fff;
  border:1px solid #cbd5e1;
  color:#334155;
  padding:5px 10px;
  border-radius:999px;
  font-size:13px;
}
.term-chip.active {
  background:#dbeafe;
  border-color:#93c5fd;
  color:#1d4ed8;
  font-weight:700;
}
.term-chip.exact {
  border-color:#86efac;
  background:#ecfdf5;
}
.meta-line { margin-top:6px; line-height:1.6; color:#334155; }
.btn-small { padding:7px 12px; font-size:13px; }
.search-filter-panel .term-filter-row {
  display:flex;
  flex-wrap:wrap;
  gap:10px 16px;
  margin-bottom:12px;
}
.search-filter-panel label {
  display:inline-flex;
  align-items:center;
  gap:6px;
  cursor:pointer;
  background:#fff;
  border:1px solid #dbe1ea;
  border-radius:999px;
  padding:6px 12px;
}
.search-filter-panel label:has(input:checked) {
  background:#dbeafe;
  border-color:#93c5fd;
}
.search-filter-actions { display:flex; gap:8px; flex-wrap:wrap; margin-top:4px; }
.snippet, .full-block { background:#fff; padding:12px; white-space:pre-wrap; border:1px solid #eee; border-radius:6px; font-family:Consolas, monospace; }
.snippet-columns {
  display:grid;
  grid-template-columns: minmax(0, 1fr) minmax(0, 1fr);
  gap:12px;
  align-items:start;
  margin-top:10px;
}
.snippet-panel {
  min-width:0;
}
.snippet-panel .snippet {
  margin-top:8px;
  height:auto;
  box-sizing:border-box;
}
.snippet-panel-original .snippet {
  border-left:4px solid #0069d9;
}
.snippet-panel-translation .snippet {
  border-left:4px solid #28a745;
}
@media (max-width: 900px) {
  .snippet-columns {
    grid-template-columns: 1fr;
  }
}
.highlight { color:green; font-weight:bold; background:#eaf7ea; padding:1px 2px; border-radius:2px; }
.grid { display:grid; grid-template-columns: 1fr 1fr; gap:16px; }
.data-table table, table.data-table { width:100%; border-collapse:collapse; }
.data-table th, .data-table td, table.data-table th, table.data-table td { border:1px solid #ddd; padding:6px 8px; font-size:14px; vertical-align:top; }
.table-result-columns {
  display:flex;
  flex-direction:column;
  gap:12px;
  margin-top:6px;
}
.table-result-item {
  margin-top:8px;
  padding:10px 12px;
}
.table-result-item .meta-prominent {
  margin:6px 0 8px 0;
}
.table-result-item .meta {
  margin-top:8px;
}
.search-results-text {
  margin-top:12px;
}
.search-results-text h3 {
  margin:0 0 8px 0;
}
.search-empty-notice {
  margin-top:12px;
  padding:12px 14px;
}
.search-empty-notice .panel-title {
  margin-bottom:8px;
  padding-bottom:6px;
}
.compact-warning {
  margin:0;
  padding:10px 12px;
}
.table-scroll-wrap {
  overflow-x:auto;
  overflow-y:visible;
  max-width:100%;
  -webkit-overflow-scrolling:touch;
  border:1px solid #e5e7eb;
  border-radius:6px;
  background:#fff;
}
.requirement-bilingual-table td {
  vertical-align: top;
  white-space: pre-wrap;
  line-height: 1.45;
}
.requirement-bilingual-table .req-number-cell {
  font-weight: 700;
  background: #f8fafc;
}
.table-scroll-wrap .data-table {
  display: inline-block;
  width: auto;
  max-width: 100%;
}
.table-scroll-wrap .data-table table,
.table-scroll-wrap .data-table table.compact-table {
  width: max-content;
  max-width: none;
  table-layout: auto;
  margin: 0;
}
.table-scroll-wrap .data-table th,
.table-scroll-wrap .data-table td {
  width: auto;
  white-space: normal;
  overflow-wrap: anywhere;
  word-break: break-word;
  padding: 4px 10px;
  vertical-align: top;
}
.table-scroll-wrap .table-footnote {
  margin-top: 8px;
  padding: 8px 10px;
  font-size: 12px;
  line-height: 1.45;
  color: #475569;
  background: #f8fafc;
  border: 1px solid #e2e8f0;
  border-radius: 4px;
  white-space: normal;
}
.table-result-panel {
  min-width:0;
  width:100%;
}
.table-result-panel > strong {
  display:block;
  margin-bottom:6px;
}
a { color:#0069d9; text-decoration:none; }
a:hover { text-decoration:underline; }
.small { font-size:13px; color:#666; }
.badge { display:inline-block; background:#eef; color:#334; padding:3px 8px; border-radius:12px; font-size:12px; margin-right:6px; }
.page-img-preview { max-width:200px; max-height:200px; border:1px solid #ccc; margin-right:16px; float:left; cursor:pointer; transition: transform 0.2s; }
.page-img-preview:hover { transform: scale(1.02); }
.result-page-img img {
  max-width: 240px;
  max-height: 240px;
  border: 1px solid #ccc;
  border-radius: 6px;
}
.result-page-img {
  margin: 10px 0 12px 0;
}
.document-match-grid {
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(280px, 1fr));
  gap: 16px;
  margin-top: 12px;
}
.document-match-card {
  background: #fff;
  border: 1px solid #dbe1ea;
  border-radius: 10px;
  padding: 12px;
}
.document-match-card img {
  width: 100%;
  max-height: 260px;
  object-fit: contain;
  border: 1px solid #ddd;
  border-radius: 6px;
  background: #fff;
}
.document-match-card .doc-title {
  font-weight: 700;
  margin-bottom: 8px;
  line-height: 1.4;
}
.document-match-card .doc-snippet {
  margin-top: 10px;
  font-size: 13px;
}
.pagination-bar {
  margin-top: 14px;
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 10px;
  flex-wrap: wrap;
  padding-right: 132px;
  padding-bottom: 6px;
}
.pagination-buttons {
  display: flex;
  align-items: center;
  gap: 8px;
}
.back-to-top {
  position: fixed;
  right: 22px;
  bottom: 22px;
  z-index: 999;
  display: none;
  background: #0ea5a4;
  color: #fff;
  border: 1px solid #0b8d8c;
  box-shadow: 0 6px 14px rgba(14, 165, 164, 0.28);
  border-radius: 999px;
  font-weight: 600;
  padding: 10px 14px;
  transition: transform 0.15s ease, background 0.15s ease;
}
.back-to-top:hover {
  background: #0b8d8c;
  transform: translateY(-1px);
}
@media (max-width: 900px) {
  .pagination-bar {
    padding-right: 0;
    padding-bottom: 52px;
  }
  .back-to-top {
    right: 14px;
    bottom: 14px;
  }
  .result-page-img img {
    max-width: 100%;
    height: auto;
  }
}
.clearfix::after { content:""; display:table; clear:both; }
@media (max-width: 900px) {
  .grid { grid-template-columns: 1fr; }
  input[type=text] { width:100%; }
  .topbar { flex-wrap:wrap; }
  .actions { width:100%; justify-content:flex-end; }
}
</style>
"""

HOME_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Technical Standards Search Portal</title>
""" + BASE_CSS + """
</head>
<body>
<div class="container">
    <div class="topbar">
        <div class="topbar-brand">
            <h1>{% if search_mode == 'general' %}Normal Search{% else %}Requirement Browser{% endif %}</h1>
            <div class="small">Requirement-level search for PDF, TXT, CSV, XLSX, DOCX</div>
        </div>
        <div class="actions">
            <a class="btn btn-gray" href="{{ url_for('home') }}">Home page</a>
            <a class="btn btn-purple" href="{{ url_for('admin_documents') }}">Documents &amp; upload</a>
            <a class="btn btn-gray" href="{{ url_for('requirement_browser') }}">Requirement Browser</a>
            <a class="btn btn-purple" href="{{ url_for('dictionary_lookup') }}">Definitions</a>
        </div>
    </div>

    {% with messages = get_flashed_messages() %}
      {% if messages %}
        {% for msg in messages %}
          <div class="flash">{{ msg }}</div>
        {% endfor %}
      {% endif %}
    {% endwith %}

    <div class="notice">
        Examples: <strong>aarding</strong>, <strong>grounding</strong>, <strong>earthing resistance</strong>, <strong>AM-Req-6165</strong>.
        For abbreviations and reference text, use <a href="{{ url_for('dictionary_lookup') }}"><strong>Definitions</strong></a>
        (e.g. <strong>AC</strong>, <strong>declaration of performance</strong>).
    </div>

    {% if search_mode != 'requirement' %}
    <form method="GET" action="/" id="search-form">
        <input type="text" name="q" value="{{ query }}" placeholder="Search requirement ID, Dutch/English keyword, section, definition">
        <button type="submit">Search</button>

        {% if query and all_expanded_terms %}
        <div class="search-filter-panel">
            <div class="panel-title">Term filters</div>
            <div class="panel-section">
                <div class="panel-label">Select search terms</div>
                <div class="small" style="margin-bottom:10px;">Exact search word is checked by default. Choose additional related terms if needed.</div>
                <div class="term-filter-row">
                    {% for term in all_expanded_terms %}
                    <label>
                        <input type="checkbox" name="term" value="{{ term }}"
                            {% if term in selected_terms %}checked{% endif %}>
                        <span{% if term in exact_terms %} style="font-weight:700;"{% endif %}>{{ term }}</span>
                    </label>
                    {% endfor %}
                </div>
                <div class="panel-label" style="margin-top:14px;">SELECT SEARCH TYPES</div>
                <div class="term-filter-row">
                    <label>
                        <input type="checkbox" name="result_type" value="drawings"
                            {% if 'drawings' in result_types %}checked{% endif %}>
                        <span>related drawings</span>
                    </label>
                    <label>
                        <input type="checkbox" name="result_type" value="text"
                            {% if 'text' in result_types %}checked{% endif %}>
                        <span>Text</span>
                    </label>
                    <label>
                        <input type="checkbox" name="result_type" value="tables"
                            {% if 'tables' in result_types %}checked{% endif %}>
                        <span>excel table</span>
                    </label>
                </div>
                <div class="search-filter-actions">
                    <button type="submit" class="btn btn-green btn-small">Apply term filters</button>
                    <button type="button" class="btn btn-gray btn-small" onclick="selectAllTerms(true)">Select all</button>
                </div>
                {% if selected_terms %}
                <div class="meta-line small" style="margin-top:10px;">
                    <strong>Active terms:</strong> {{ selected_terms|join(', ') }}
                </div>
                {% endif %}
            </div>
        </div>
        {% endif %}
    </form>
    {% endif %}

    {% if search_mode != 'general' %}
    <div class="search-meta-panel" style="margin-top:16px;">
        <div class="panel-title">Requirement Browser</div>
        <form method="GET" action="{{ requirement_form_action }}" id="requirement-browser-form">
            <input type="text" name="req_lookup" value="{{ req_lookup_query }}" placeholder="Find one requirement, e.g. AM-Req-0286.06">
            <button type="submit">Browse</button>
        </form>
        <div class="small" style="margin-top:8px;">
            Browse all indexed information for one requirement, including related Excel tables.
        </div>

        {% if req_lookup_query %}
            {% if req_lookup and req_lookup.error_message %}
                <div class="warning" style="margin-top:12px;">
                    Requirement Browser error: <strong>{{ req_lookup.error_message }}</strong>
                </div>
            {% endif %}
            {% if req_lookup and req_lookup.found %}
                <div class="meta" style="margin-top:10px;">
                    Requirement: <strong>{{ req_lookup.normalized_id or req_lookup_query }}</strong>
                    · Requirement rows: <strong>{{ req_lookup.requirement_row_count }}</strong>
                    · Records: <strong>{{ req_lookup.block_count }}</strong>
                    · Related documents: <strong>{{ req_lookup.document_count }}</strong>
                    · Tables: <strong>{{ req_lookup.table_count }}</strong>
                    · Scan: <strong>{{ req_lookup.scan_elapsed_ms }} ms</strong>
                    · Files: <strong>{{ req_lookup.scan_documents }}</strong>
                    · Sheets: <strong>{{ req_lookup.scan_sheets }}</strong>
                </div>

                {% if req_lookup.requirement_rows %}
                <div class="result-item" style="margin-top:12px;">
                    <strong>Requirement information (from Excel rows)</strong>
                    {% for row in req_lookup.requirement_rows %}
                    <div class="summary-box" style="margin-top:10px;">
                        <div class="meta">
                            <strong>{{ row.document_name }}</strong>
                            · Sheet: {{ row.sheet_name }}
                            · Row: {{ row.row_number }}
                            · <a href="{{ url_for('table_preview', document_id=row.document_id) }}" target="_blank">Open table</a>
                        </div>
                        <div class="table-scroll-wrap" style="margin-top:8px;">
                            <table class="data-table" style="width:max-content; min-width:100%;">
                                <tr>
                                    {% for cell in row.cells %}
                                    <th>{{ cell.column }}</th>
                                    {% endfor %}
                                </tr>
                                <tr>
                                    {% for cell in row.cells %}
                                    <td style="white-space:pre-wrap; min-width:160px;">{{ cell.value }}</td>
                                    {% endfor %}
                                </tr>
                            </table>
                        </div>
                    </div>
                    {% endfor %}
                </div>
                {% endif %}

                {% if req_lookup.blocks %}
                <div class="result-item" style="margin-top:12px;">
                    <strong>Related document content</strong>
                    {% for item in req_lookup.blocks %}
                    <div class="summary-box" style="margin-top:10px;">
                        <div class="meta">
                            <strong>{{ item.document_name }}</strong>
                            {% if item.section %}
                            · Section: {{ item.section }}
                            {% endif %}
                            {% if item.page %}
                            · Page: {{ item.page }}
                            {% endif %}
                        </div>
                        <div class="table-scroll-wrap" style="margin-top:8px;">
                            <table class="data-table requirement-bilingual-table" style="width:100%; min-width:720px;">
                                <tr>
                                    <th style="width:50%;">Dutch (Original)</th>
                                    <th style="width:50%;">English (Translation)</th>
                                </tr>
                                <tr>
                                    <td class="req-number-cell">
                                        {% if item.requirement_id %}
                                        <strong>Requirement number: {{ item.requirement_id }}</strong>
                                        {% else %}
                                        -
                                        {% endif %}
                                    </td>
                                    <td class="req-number-cell">
                                        {% if item.requirement_id %}
                                        <strong>Requirement number: {{ item.requirement_id }}</strong>
                                        {% else %}
                                        -
                                        {% endif %}
                                    </td>
                                </tr>
                                <tr>
                                    <td>
                                        {% if item.title_nl %}
                                        <strong>Title:</strong> {{ item.title_nl }}
                                        {% else %}
                                        -
                                        {% endif %}
                                    </td>
                                    <td>
                                        {% if item.title_en %}
                                        {{ item.title_en }}
                                        {% else %}
                                        -
                                        {% endif %}
                                    </td>
                                </tr>
                                <tr>
                                    <td>
                                        {% if item.applicable_nl %}
                                        Applicable Discriminators Distance security: {{ item.applicable_nl }}
                                        {% else %}
                                        -
                                        {% endif %}
                                    </td>
                                    <td>
                                        {% if item.applicable_en %}
                                        {{ item.applicable_en }}
                                        {% else %}
                                        -
                                        {% endif %}
                                    </td>
                                </tr>
                            </table>
                        </div>
                        <div class="meta" style="margin-top:8px;">
                            {% if item.block_id %}
                            <a href="{{ url_for('requirement_detail', block_id=item.block_id) }}" target="_blank">Open requirement page</a>
                            {% endif %}
                            {% if item.is_pdf %}
                            · <a href="{{ url_for('pdf_viewer', document_id=item.document_id, page=item.page or 1) }}" target="_blank">Open PDF</a>
                            {% elif item.is_docx %}
                            · <a href="{{ url_for('docx_viewer', document_id=item.document_id, page=item.page or 1) }}" target="_blank">Open DOCX</a>
                            {% elif item.is_txt %}
                            · <a href="{{ url_for('document_view', document_id=item.document_id, page=item.page or 1) }}" target="_blank">Open TXT</a>
                            {% endif %}
                        </div>
                    </div>
                    {% endfor %}
                </div>
                {% endif %}

                {% if req_lookup.tables %}
                <div class="result-item" style="margin-top:12px;">
                    <strong>Related tables</strong>
                    {% for t in req_lookup.tables %}
                    <div class="summary-box" style="margin-top:10px;">
                        <div><strong>{{ t.document_name }}</strong></div>
                        <div class="meta">Sheet: {{ t.sheet_name }} · Format: {{ t.table_format }}</div>
                        <div style="margin-top:8px;">
                            <strong>Table 1 / Dutch (Original)</strong>
                            <div class="table-scroll-wrap" style="margin-top:6px;">
                                <div class="data-table">{{ t.html_table|safe }}</div>
                            </div>
                        </div>
                        {% if t.html_table_en %}
                        <div style="margin-top:10px;">
                            <strong>Table 2 / English (Translation)</strong>
                            <div class="table-scroll-wrap" style="margin-top:6px;">
                                <div class="data-table">{{ t.html_table_en|safe }}</div>
                            </div>
                        </div>
                        {% endif %}
                        <a href="{{ url_for('table_preview', document_id=t.document_id) }}" target="_blank">Open table</a>
                    </div>
                    {% endfor %}
                </div>
                {% endif %}
            {% else %}
                <div class="warning" style="margin-top:12px;">
                    No matching requirement data found for <strong>{{ req_lookup_query }}</strong>.
                </div>
                {% if req_lookup %}
                <div class="meta" style="margin-top:8px;">
                    Scan: <strong>{{ req_lookup.scan_elapsed_ms }} ms</strong>
                    · Files: <strong>{{ req_lookup.scan_documents }}</strong>
                    · Sheets: <strong>{{ req_lookup.scan_sheets }}</strong>
                </div>
                {% endif %}
            {% endif %}
        {% endif %}
    </div>
    {% endif %}

    <script>
    function selectAllTerms(checked) {
        document.querySelectorAll('#search-form input[name="term"]').forEach(function(el) {
            el.checked = checked;
        });
    }
    function scrollToTopSmooth() {
        window.scrollTo({ top: 0, behavior: "smooth" });
    }
    window.addEventListener("scroll", function() {
        const topBtn = document.getElementById("back-to-top");
        if (!topBtn) return;
        topBtn.style.display = window.scrollY > 300 ? "inline-block" : "none";
    });
    </script>

    {% if doc_count == 0 %}
        <div class="warning">
            No documents are indexed yet. Log in as admin, upload PDF files, then use <strong>Reindex</strong> in Documents.
        </div>
    {% endif %}

    {% if search_mode == 'general' and query %}
        {% if document_matches and current_page == 1 and 'drawings' in result_types %}
        <div class="search-meta-panel">
            <div class="panel-title">related drawings</div>
            <div class="document-match-grid">
                {% for doc in document_matches %}
                <div class="document-match-card">
                    <div class="doc-title">
                        {% if doc.is_pdf %}
                        <a href="{{ url_for('pdf_viewer', document_id=doc.document_id, page=doc.page or 1) }}" target="_blank">
                            {{ doc.filename }}
                        </a>
                        {% elif doc.is_docx %}
                        <a href="{{ url_for('docx_viewer', document_id=doc.document_id, page=doc.page or 1) }}" target="_blank">
                            {{ doc.filename }}
                        </a>
                        {% elif doc.is_txt %}
                        <a href="{{ url_for('document_view', document_id=doc.document_id, page=doc.page or 1) }}" target="_blank">
                            {{ doc.filename }}
                        </a>
                        {% else %}
                            {{ doc.filename }}
                        {% endif %}
                    </div>
                    <div class="meta">
                        Filename: <strong>{{ doc.filename }}</strong>
                        |
                        Requirement number: <strong>{{ doc.requirement_id or '-' }}</strong>
                        | Page: <strong>{{ doc.page or 1 }}</strong>
                    </div>
                    {% if doc.is_pdf and doc.image_pages %}
                    <div style="display:grid; grid-template-columns:repeat(auto-fit, minmax(160px, 1fr)); gap:8px;">
                        {% for page_num in doc.image_pages %}
                        <a href="{{ url_for('pdf_viewer', document_id=doc.document_id, page=page_num) }}" target="_blank">
                            <img
                                src="{{ url_for('page_image', document_id=doc.document_id, page=page_num) }}"
                                loading="lazy"
                                alt="PDF page {{ page_num }} preview"
                                onerror="this.style.display='none';"
                            />
                        </a>
                        {% endfor %}
                    </div>
                    {% endif %}
                    {% if doc.is_docx and doc.image_indexes %}
                    <div style="display:grid; grid-template-columns:repeat(auto-fit, minmax(160px, 1fr)); gap:8px;">
                        {% for image_index in doc.image_indexes %}
                        <a href="{{ url_for('docx_viewer', document_id=doc.document_id, page=doc.page or 1, image_index=image_index) }}" target="_blank">
                            <img
                                src="{{ url_for('docx_image', document_id=doc.document_id, image_index=image_index) }}"
                                loading="lazy"
                                alt="DOCX drawing {{ image_index + 1 }}"
                                onerror="this.style.display='none';"
                            />
                        </a>
                        {% endfor %}
                    </div>
                    {% endif %}
                </div>
                {% endfor %}
            </div>
        </div>
        {% elif 'drawings' in result_types and current_page == 1 %}
        <div class="search-meta-panel search-empty-notice">
            <div class="panel-title">related drawings</div>
            <div class="warning compact-warning">No matching related drawings found.</div>
        </div>
        {% endif %}

        {% if table_results and current_page == 1 and 'tables' in result_types %}
        <div class="search-meta-panel">
            <div class="panel-title">excel table</div>
            {% for t in table_results %}
            <div class="result-item table-result-item">
                <div class="filename">
                    <a href="{{ url_for('table_preview', document_id=t.document_id) }}" target="_blank">
                        {{ t.preview_title }}
                    </a>
                </div>
                <div class="meta-prominent">
                    <div class="meta-item">Filename: <strong>{{ t.filename }}</strong></div>
                    <div class="meta-item">Sheet: <strong>{{ t.sheet_name }}</strong></div>
                    <div class="meta-item">Format: <strong>{{ t.table_format }}</strong></div>
                </div>
                <div class="table-result-columns">
                    <div class="table-result-panel snippet-panel-original">
                        <strong>Table 1 / Dutch (Original)</strong>
                        <div class="table-scroll-wrap">
                            <div class="data-table">{{ t.html_table|safe }}</div>
                        </div>
                    </div>
                    {% if t.html_table_en %}
                    <div class="table-result-panel snippet-panel-translation">
                        <strong>Table 2 / English (Translation)</strong>
                        <div class="table-scroll-wrap">
                            <div class="data-table">{{ t.html_table_en|safe }}</div>
                        </div>
                    </div>
                    {% endif %}
                </div>
                <div class="meta">
                    <a href="{{ url_for('table_preview', document_id=t.document_id) }}" target="_blank">Open full table preview</a>
                </div>
            </div>
            {% endfor %}
        </div>
        {% elif 'tables' in result_types and current_page == 1 %}
        <div class="search-meta-panel search-empty-notice">
            <div class="panel-title">excel table</div>
            <div class="warning compact-warning">No matching excel table results found.</div>
        </div>
        {% endif %}

        {% if 'text' in result_types %}
        <div class="search-results-text">
        <h3>Text results for "{{ query }}"</h3>

        {% if results %}
            {% for res in results %}
                <div class="result-item">
                    <div class="filename">
                        {% if res.is_pdf %}
                            <a href="{{ url_for('pdf_viewer', document_id=res.document_id, page=res.page or 1) }}" target="_blank">
                                {{ res.requirement_id or '-' }} - {{ res.title or res.filename }}
                            </a>
                        {% elif res.is_docx %}
                            <a href="{{ url_for('docx_viewer', document_id=res.document_id, page=res.page or 1) }}" target="_blank">
                                {{ res.requirement_id or '-' }} - {{ res.title or res.filename }}
                            </a>
                        {% elif res.is_txt %}
                            <a href="{{ url_for('document_view', document_id=res.document_id, page=res.page or 1) }}" target="_blank">
                                {{ res.requirement_id or '-' }} - {{ res.title or res.filename }}
                            </a>
                        {% else %}
                            {{ res.requirement_id or '-' }} - {{ res.title or res.filename }}
                        {% endif %}
                    </div>

                    <div class="meta-prominent">
                        <div class="meta-item">Filename: <strong>{{ res.filename }}</strong></div>
                        <div class="meta-item">Page: <strong>{{ res.page or '?' }}</strong></div>
                        {% if res.major_section or res.section %}
                        <div class="meta-item">Section: <strong>{{ res.major_section or res.section }}</strong></div>
                        {% endif %}
                    </div>
                    {% if res.ocr_used %}
                    <div class="meta-secondary">OCR used</div>
                    {% endif %}

                    {% if res.is_pdf %}
                    <div class="result-page-img">
                        <a href="{{ url_for('pdf_viewer', document_id=res.document_id, page=res.page or 1) }}" target="_blank">
                            <img
                                src="{{ url_for('page_image', document_id=res.document_id, page=res.page or 1) }}"
                                loading="lazy"
                                alt="Page {{ res.page or 1 }} preview"
                                onerror="this.style.display='none';"
                            />
                        </a>
                    </div>
                    {% endif %}

                    {% if res.definition %}
                        <div class="meta-secondary">Definition (NL): ({{ res.definition }})</div>
                    {% endif %}
                    {% if res.summary %}
                        <div class="meta-secondary">Summary (NL): ({{ res.summary }})</div>
                    {% endif %}
                    {% if res.section and res.major_section and res.section != res.major_section %}
                        <div class="meta-secondary">Subsection (NL): ({{ res.section }})</div>
                    {% elif res.section and not res.major_section %}
                        <div class="meta-secondary">Section (NL): ({{ res.section }})</div>
                    {% endif %}

                    <div class="snippet-columns">
                        <div class="snippet-panel snippet-panel-original">
                            <strong>Table 1 / Dutch (Original)</strong>
                            <div class="snippet">({{ res.snippet|safe }})</div>
                        </div>
                        {% if res.snippet_en %}
                        <div class="snippet-panel snippet-panel-translation">
                            <strong>Table 2 / English (Translation)</strong>
                            <div class="snippet">{{ res.snippet_en|safe }}</div>
                        </div>
                        {% endif %}
                    </div>

                    <div class="meta">
                        {% if res.is_document_fallback %}
                            {% if res.is_pdf %}
                            <a href="{{ url_for('pdf_viewer', document_id=res.document_id, page=res.page or 1) }}" target="_blank">Open PDF at page {{ res.page or 1 }}</a>
                            {% elif res.is_docx %}
                            <a href="{{ url_for('docx_viewer', document_id=res.document_id, page=res.page or 1) }}" target="_blank">Open DOCX translation</a>
                            {% elif res.is_txt %}
                            <a href="{{ url_for('document_view', document_id=res.document_id, page=res.page or 1) }}" target="_blank">Open TXT translation</a>
                            {% endif %}
                        {% else %}
                            <a href="{{ url_for('requirement_detail', block_id=res.block_id) }}">Open full requirement</a>
                            {% if res.is_pdf %}
                                | <a href="{{ url_for('pdf_viewer', document_id=res.document_id, page=res.page or 1) }}" target="_blank">Open PDF at page {{ res.page or 1 }}</a>
                            {% elif res.is_docx %}
                                | <a href="{{ url_for('docx_viewer', document_id=res.document_id, page=res.page or 1) }}" target="_blank">Open DOCX translation</a>
                            {% elif res.is_txt %}
                                | <a href="{{ url_for('document_view', document_id=res.document_id, page=res.page or 1) }}" target="_blank">Open TXT translation</a>
                            {% endif %}
                            {% if res.has_table %}
                                | <a href="{{ url_for('table_preview', document_id=res.document_id) }}">Open table preview</a>
                            {% endif %}
                        {% endif %}
                    </div>
                </div>
            {% endfor %}
            {% if total_pages > 1 %}
            <div class="pagination-bar">
                <div class="small">
                    Showing {{ ((current_page - 1) * results_per_page) + 1 }}-{{ ((current_page - 1) * results_per_page) + (results|length) }}
                    of {{ total_results }} results
                </div>
                <div class="pagination-buttons">
                    {% if current_page > 1 %}
                    <a class="btn btn-gray btn-small" href="{{ url_for('home', q=query, page=current_page-1, term=selected_terms, result_type=result_types) }}">Previous</a>
                    {% endif %}
                    <span class="small">Page {{ current_page }} / {{ total_pages }}</span>
                    {% if current_page < total_pages %}
                    <a class="btn btn-gray btn-small" href="{{ url_for('home', q=query, page=current_page+1, term=selected_terms, result_type=result_types) }}">Next</a>
                    {% endif %}
                </div>
            </div>
            {% endif %}
        {% else %}
            <div class="warning compact-warning">No matching Text results found.</div>
        {% endif %}
        </div>
        {% endif %}
    {% endif %}
</div>
<button id="back-to-top" class="btn btn-gray back-to-top" onclick="scrollToTopSmooth()">Back to top</button>
</body>
</html>
"""

# Requirement detail template also gets image preview
REQUIREMENT_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Requirement detail</title>
""" + BASE_CSS + """
</head>
<body>
<div class="container">
    <h2>{{ block.requirement_id or 'Requirement block' }}</h2>

    <div class="summary-box">
        <div><strong>Title:</strong> {{ block.title or '-' }}</div>
        <div><strong>Document:</strong> {{ block.document.original_filename }}</div>
        <div><strong>Page:</strong> {{ block.page }}</div>
        <div><strong>Category:</strong> {{ block.category }}</div>
        {% if block.major_section %}<div><strong>Section:</strong> {{ block.major_section }}</div>{% endif %}
        {% if block.section and block.section != block.major_section %}<div><strong>Subsection:</strong> {{ block.section }}</div>{% endif %}
        {% if block.definition %}<div><strong>Definition:</strong> {{ block.definition }}</div>{% endif %}
        {% if block.summary %}<div><strong>Summary:</strong> {{ block.summary }}</div>{% endif %}
    </div>

    <!-- NEW: page image preview for PDFs in requirement detail -->
    {% if block.document.extension == 'pdf' and pdf2image_available %}
        <div style="margin: 16px 0;">
            <a href="{{ url_for('pdf_viewer', document_id=block.document.id, page=block.page or 1) }}" target="_blank">
                <img src="{{ url_for('page_image', document_id=block.document.id, page=block.page or 1) }}"
                     style="max-width:100%; max-height:500px; border:1px solid #ccc;"
                     alt="Page {{ block.page or 1 }} preview"
                     onerror="this.style.display='none'">
            </a>
        </div>
    {% endif %}

    <h3>Full requirement text</h3>
    <div class="full-block">{{ block.full_text }}</div>

    <div style="margin-top:16px;">
        <a class="btn btn-gray" href="{{ url_for('home') }}">Back</a>
        {% if block.document.extension == 'pdf' %}
            <a class="btn btn-purple" href="{{ url_for('pdf_viewer', document_id=block.document.id, page=block.page or 1) }}" target="_blank">Open PDF at page {{ block.page or 1 }}</a>
        {% elif block.document.extension == 'docx' %}
            <a class="btn btn-purple" href="{{ url_for('docx_viewer', document_id=block.document.id, page=1) }}" target="_blank">Open DOCX translation</a>
        {% elif block.document.extension == 'txt' %}
            <a class="btn btn-purple" href="{{ url_for('document_view', document_id=block.document.id, page=1) }}" target="_blank">Open TXT translation</a>
        {% endif %}
    </div>
</div>
</body>
</html>
"""

#LOGIN_TEMPLATE
LOGIN_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Login</title>
""" + BASE_CSS + """
</head>
<body>
<div class="container" style="max-width: 520px;">
    <h2>Login</h2>

    {% with messages = get_flashed_messages() %}
      {% if messages %}
        {% for msg in messages %}
          <div class="warning">{{ msg }}</div>
        {% endfor %}
      {% endif %}
    {% endwith %}

    <form method="POST">
        <div style="margin-bottom:12px;">
            <label>Username</label><br>
            <input type="text" name="username" required style="width:100%;">
        </div>
        <div style="margin-bottom:12px;">
            <label>Password</label><br>
            <input type="password" name="password" required style="width:100%;">
        </div>
        <button type="submit">Login</button>
        <a class="btn btn-gray" href="{{ url_for('home') }}">Back</a>
    </form>

    <div class="notice" style="margin-top:16px;">
        Create the first admin user by setting <code>ADMIN_USERNAME</code> and
        <code>ADMIN_PASSWORD</code> before the first start.
    </div>
</div>
</body>
</html>
"""

STORAGE_USAGE_PANEL = """
    <div class="storage-usage-panel{% if storage_usage.near_limit %} storage-usage-warn{% endif %}{% if storage_usage.over_limit %} storage-usage-danger{% endif %}">
        <div class="panel-title">Storage usage</div>
        <div class="storage-usage-summary">
            <strong>{{ storage_usage.used_label }}</strong> / {{ storage_usage.quota_label }}
            <span class="small">({{ storage_usage.percent }}% · {{ storage_usage.doc_count }} documents{% if storage_usage.dict_count %} · {{ storage_usage.dict_count }} definition files{% endif %})</span>
        </div>
        <div class="storage-usage-bar" aria-hidden="true">
            <div class="storage-usage-fill" style="width: {% if storage_usage.percent > 100 %}100{% else %}{{ storage_usage.percent }}{% endif %}%;"></div>
        </div>
        <div class="storage-usage-note">{{ storage_usage.remaining_label }} remaining · documents and definition files</div>
    </div>
"""

UPLOAD_PANEL = """
    <div class="stats-panel" id="upload-panel">
        <div class="panel-title">Upload files</div>
        """ + STORAGE_USAGE_PANEL + """
        <div class="info" style="margin-bottom:12px;">
            Allowed types: {{ allowed_extensions }}<br>
            Max size: 100 MB per request<br>
            Storage backend: <strong>{{ storage_backend }}</strong>
            {% if storage_status.persistent %}
            <br>Bucket: <strong>{{ storage_status.bucket }}</strong>
            {% if storage_status.key_prefix and storage_status.key_prefix != '(root)' %}
            <br>Key prefix: <strong>{{ storage_status.key_prefix }}</strong>
            {% endif %}
            {% endif %}
            {% if not storage_persistent %}
            <br><span style="color:#b45309;"><strong>Warning:</strong> Files are stored on temporary disk and will be lost after a Render restart. Configure Supabase S3 for persistent PDF storage.</span>
            {% endif %}
        </div>
        {% if storage_status.persistent and not storage_status.ok %}
        <div class="flash" style="background:#fdecea; border-left-color:#dc3545;">
            <strong>Storage not ready:</strong> {{ storage_status.message }}<br>
            {{ storage_status.hint }}
        </div>
        {% endif %}
        <form method="POST" action="{{ url_for('upload_files') }}" enctype="multipart/form-data" id="upload-form"{% if storage_status.persistent and not storage_status.ok %} onsubmit="alert('Fix Supabase storage configuration before uploading.'); return false;"{% else %} onsubmit="return handleUploadSubmit(this);"{% endif %}>
            <input type="file" name="files" multiple required>
            <div style="margin-top:12px;">
                <button type="submit" id="upload-submit-btn" class="btn btn-green">Upload and index</button>
            </div>
        </form>
    </div>
"""

DOCS_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Documents</title>
""" + BASE_CSS + """
</head>
<body>
<div class="container">
    <h2>Documents &amp; upload</h2>

    <div style="margin-bottom:16px;">
        <a class="btn btn-gray" href="{{ url_for('home') }}">Back</a>
    </div>

    """ + UPLOAD_PANEL + """

    {% with messages = get_flashed_messages(with_categories=true) %}
      {% if messages %}
        <div id="admin-feedback" class="admin-feedback-panel">
        {% for category, msg in messages %}
          <div class="flash flash-{{ category if category != 'message' else 'info' }}">{{ msg }}</div>
        {% endfor %}
        </div>
      {% else %}
        <div id="admin-feedback"></div>
      {% endif %}
    {% endwith %}

    <div class="stats-panel">
        <div class="panel-title">Index statistics</div>
        <div class="stats-grid">
            <div class="stat-item">
                <span class="stat-label">Indexed documents</span>
                <span class="stat-value">{{ doc_count }}</span>
            </div>
            <div class="stat-item">
                <span class="stat-label">Requirement blocks</span>
                <span class="stat-value">{{ block_count }}</span>
            </div>
            <div class="stat-item">
                <span class="stat-label">Tables</span>
                <span class="stat-value">{{ table_count }}</span>
            </div>
        </div>
    </div>

    <form method="POST" action="{{ url_for('bulk_delete_documents') }}">
        <div class="panel-title" style="margin-top:8px;">Indexed documents</div>
        <div style="margin-bottom:12px; display:flex; gap:8px; flex-wrap:wrap; align-items:center;">
            <button type="submit" formaction="{{ url_for('bulk_reindex_documents') }}" class="btn btn-green" onclick="return confirmBulkReindex();">Reindex Selected</button>
            <button type="submit" class="btn btn-orange" onclick="return confirmBulkDelete();">Delete Selected</button>
            <button type="button" class="btn btn-gray" onclick="toggleAll(true)">Select All</button>
            <button type="button" class="btn btn-gray" onclick="toggleAll(false)">Clear All</button>
            <span class="small">Select multiple documents to reindex or delete in one action.</span>
        </div>

        <table class="data-table" style="margin-top:16px; width:100%;">
            <tr>
                <th style="width:40px;">
                    <input type="checkbox" onclick="toggleFromHeader(this)">
                </th>
                <th>ID</th>
                <th>Filename</th>
                <th>Type</th>
                <th>Size</th>
                <th>OCR</th>
                <th>Status</th>
                <th>File</th>
                <th>Uploaded</th>
                <th>Actions</th>
            </tr>
            {% for row in doc_rows %}
            {% set d = row.record %}
            <tr>
                <td>
                    <input type="checkbox" name="document_ids" value="{{ d.id }}" class="doc-checkbox">
                </td>
                <td>{{ d.id }}</td>
                <td>{{ d.original_filename }}</td>
                <td>{{ d.extension }}</td>
                <td>{{ d.size_bytes }}</td>
                <td>{{ 'Yes' if d.is_ocr else 'No' }}</td>
                <td>{{ d.status }}</td>
                <td>
                    {% if row.file_available %}
                    <span style="color:#198754; font-weight:700;">OK</span>
                    {% else %}
                    <span style="color:#dc3545; font-weight:700;">Missing</span>
                    {% endif %}
                </td>
                <td>{{ d.uploaded_at }}</td>
                <td>
                    <a href="{{ url_for('reindex_document', document_id=d.id) }}">Reindex</a>
                    |
                    <a href="{{ url_for('delete_document', document_id=d.id) }}" onclick="return confirm('Delete this document?');">Delete</a>
                    {% if d.extension == 'pdf' %}
                    | <a href="{{ url_for('pdf_viewer', document_id=d.id, page=1) }}" target="_blank">View PDF</a>
                    | <a href="{{ url_for('export_translation_pdf', document_id=d.id) }}">Export EN PDF</a>
                    {% elif d.extension == 'docx' %}
                    | <a href="{{ url_for('docx_viewer', document_id=d.id, page=1) }}" target="_blank">View DOCX</a>
                    | <a href="{{ url_for('export_translation_pdf', document_id=d.id) }}">Export EN PDF</a>
                    {% elif d.extension == 'txt' %}
                    | <a href="{{ url_for('document_view', document_id=d.id, page=1) }}" target="_blank">View TXT</a>
                    | <a href="{{ url_for('export_translation_pdf', document_id=d.id) }}">Export EN PDF</a>
                    {% endif %}
                    {% if d.extension in ['csv','xlsx'] %}
                    | <a href="{{ url_for('table_preview', document_id=d.id) }}">Preview table</a>
                    {% endif %}
                </td>
            </tr>
            {% endfor %}
        </table>
    </form>
</div>

<script>
function toggleAll(checked) {
    const boxes = document.querySelectorAll('.doc-checkbox');
    boxes.forEach(box => box.checked = checked);
}

function toggleFromHeader(headerCheckbox) {
    toggleAll(headerCheckbox.checked);
}

function confirmBulkDelete() {
    const checked = document.querySelectorAll('.doc-checkbox:checked');
    if (checked.length === 0) {
        alert('Please select at least one document to delete.');
        return false;
    }
    return confirm('Delete ' + checked.length + ' selected document(s)? This cannot be undone.');
}

function confirmBulkReindex() {
    const checked = document.querySelectorAll('.doc-checkbox:checked');
    if (checked.length === 0) {
        alert('Please select at least one document to reindex.');
        return false;
    }
    return confirm('Reindex ' + checked.length + ' selected document(s)? This may take a while.');
}

function handleUploadSubmit(form) {
    const btn = form.querySelector('#upload-submit-btn');
    if (!btn || btn.disabled) {
        return false;
    }
    btn.disabled = true;
    btn.textContent = 'Uploading and indexing…';
    return true;
}

(function () {
    const feedback = document.getElementById('admin-feedback');
    if (feedback && feedback.querySelector('.flash')) {
        feedback.scrollIntoView({ behavior: 'smooth', block: 'nearest' });
    } else if (location.hash === '#admin-feedback' || location.hash === '#upload-panel') {
        const target = document.querySelector(location.hash);
        if (target) {
            target.scrollIntoView({ behavior: 'smooth', block: 'nearest' });
        }
    }
})();
</script>
</body>
</html>
"""

   

USERS_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Users</title>
""" + BASE_CSS + """
</head>
<body>
<div class="container">
    <h2>User management</h2>

    <form method="POST" style="margin-bottom:24px;">
        <div class="grid">
            <div>
                <label>Username</label><br>
                <input type="text" name="username" required style="width:100%;">
            </div>
            <div>
                <label>Password</label><br>
                <input type="password" name="password" required style="width:100%;">
            </div>
        </div>
        <div style="margin-top:12px;">
            <label>Role</label><br>
            <select name="role">
                <option value="user">user</option>
                <option value="admin">admin</option>
            </select>
        </div>
        <br>
        <button type="submit">Create user</button>
        <a class="btn btn-gray" href="{{ url_for('home') }}">Back</a>
    </form>

    <table class="data-table">
        <tr>
            <th>ID</th>
            <th>Username</th>
            <th>Role</th>
            <th>Created</th>
        </tr>
        {% for u in users %}
        <tr>
            <td>{{ u.id }}</td>
            <td>{{ u.username }}</td>
            <td>{{ u.role }}</td>
            <td>{{ u.created_at }}</td>
        </tr>
        {% endfor %}
    </table>
</div>
</body>
</html>
"""

REQUIREMENT_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Requirement detail</title>
""" + BASE_CSS + """
</head>
<body>
<div class="container">
    <h2>{{ block.requirement_id or 'Requirement block' }}</h2>

    <div class="summary-box">
        <div><strong>Title:</strong> {{ block.title or '-' }}</div>
        <div><strong>Document:</strong> {{ block.document.original_filename }}</div>
        <div><strong>Page:</strong> {{ block.page }}</div>
        <div><strong>Category:</strong> {{ block.category }}</div>
        {% if block.major_section %}<div><strong>Section:</strong> {{ block.major_section }}</div>{% endif %}
        {% if block.section and block.section != block.major_section %}<div><strong>Subsection:</strong> {{ block.section }}</div>{% endif %}
        {% if block.definition %}<div><strong>Definition:</strong> {{ block.definition }}</div>{% endif %}
        {% if block.summary %}<div><strong>Summary:</strong> {{ block.summary }}</div>{% endif %}
    </div>

    <h3>Full requirement text</h3>
    <div class="summary-box"><strong>Dutch / Original</strong></div>
    <div class="full-block">{{ block.full_text }}</div>
    {% if block.full_text_en %}
    <div class="summary-box" style="margin-top:16px;"><strong>English (Translation)</strong></div>
    <div class="full-block">{{ block.full_text_en }}</div>
    {% endif %}

    <div style="margin-top:16px;">
        <a class="btn btn-gray" href="{{ url_for('home') }}">Back</a>
        {% if block.document.extension == 'pdf' %}
            <a class="btn btn-purple" href="{{ url_for('pdf_viewer', document_id=block.document.id, page=block.page or 1) }}" target="_blank">Open PDF at page {{ block.page or 1 }}</a>
        {% elif block.document.extension == 'docx' %}
            <a class="btn btn-purple" href="{{ url_for('docx_viewer', document_id=block.document.id, page=1) }}" target="_blank">Open DOCX translation</a>
        {% elif block.document.extension == 'txt' %}
            <a class="btn btn-purple" href="{{ url_for('document_view', document_id=block.document.id, page=1) }}" target="_blank">Open TXT translation</a>
        {% endif %}
    </div>
</div>
</body>
</html>
"""

TRANSLATABLE_EXTENSIONS = {"pdf", "docx", "txt"}


def build_document_view_context(doc, page=1, highlight_image_index=None):
    ext = doc.extension.lower()
    if ext not in TRANSLATABLE_EXTENSIONS:
        return None

    file_available = storage.document_exists(doc.stored_filename, DOC_FOLDER)
    total_pages = count_document_pages(doc)
    if total_pages <= 0:
        total_pages = 1
    page = max(1, min(page, total_pages))

    original_text = ""
    docx_image_indexes = []
    if ext in {"docx", "txt"}:
        chunks = split_text_into_virtual_pages(get_document_full_text(doc))
        original_text = chunks[page - 1] if chunks else ""
    elif ext == "pdf" and not file_available:
        original_text = extract_page_text_from_preview(doc, page)

    if ext == "docx" and file_available:
        try:
            with storage.open_document_local_path(doc.stored_filename, DOC_FOLDER) as file_path:
                docx_image_indexes = get_docx_viewer_image_indexes(
                    file_path,
                    page=page,
                    total_pages=total_pages,
                    highlight_index=highlight_image_index,
                )
        except Exception as exc:
            print(f"DOCX viewer image load error for doc {doc.id}: {exc}")

    return {
        "document_id": doc.id,
        "filename": doc.original_filename,
        "extension": ext,
        "page": page,
        "total_pages": total_pages,
        "view_mode": "pdf" if ext == "pdf" else "text",
        "file_available": file_available,
        "pdf_available": file_available and ext == "pdf",
        "pdf_file_missing": pdf_source_file_missing(doc) if ext == "pdf" else False,
        "pdf_url": url_for("serve_document", document_id=doc.id) if ext == "pdf" and file_available else "",
        "download_url": url_for("serve_document", document_id=doc.id) if file_available else "",
        "original_text": original_text,
        "docx_image_indexes": docx_image_indexes,
        "highlight_image_index": highlight_image_index,
        "translation_enabled": translation.translation_enabled(),
        "page_label": "Page" if ext in {"pdf", "docx"} else "Section",
        "supports_page_jump": total_pages > 1 and ext in {"pdf", "docx", "txt"},
        "viewer_route": "document_view",
    }


DOCUMENT_VIEWER_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Document Viewer</title>
""" + BASE_CSS + """
<style>
body { margin:0; }
.toolbar { padding:12px 16px; background:white; border-bottom:1px solid #ddd; display:flex; justify-content:space-between; align-items:center; gap:12px; flex-wrap:wrap; }
.viewer-layout { display:flex; height: calc(100vh - 60px); }
.original-pane { flex: 1 1 58%; min-width: 320px; border-right:1px solid #ddd; position:relative; overflow:auto; }
.translation-pane { flex: 1 1 42%; min-width: 280px; overflow:auto; background:#f8fafc; padding:16px; }
iframe { width:100%; height:100%; border:none; }
.file-missing { padding:24px; color:#842029; background:#f8d7da; border:1px solid #f5c2c7; margin:16px; border-radius:8px; line-height:1.6; }
.original-text-box { margin:16px; padding:16px; background:white; border:1px solid #e2e8f0; border-radius:8px; white-space:pre-wrap; line-height:1.55; }
.translation-box { background:white; border:1px solid #e2e8f0; border-left:4px solid #28a745; border-radius:8px; padding:14px; line-height:1.55; margin-bottom:14px; }
.translation-box h4 { margin:0 0 8px 0; color:#334155; }
#translated-text { white-space:pre-wrap; line-height:1.55; }
.translation-layout-flow { display:flex; flex-direction:column; gap:6px; }
.translation-flow-row {
  display:flex;
  gap:8px;
  align-items:flex-start;
  flex-wrap:wrap;
  width:100%;
}
.translation-flow-cell {
  font-size:12px;
  line-height:1.4;
  white-space:pre-wrap;
  word-break:break-word;
  padding:5px 8px;
  background:#f8fafc;
  border:1px solid #e2e8f0;
  border-radius:4px;
  box-sizing:border-box;
}
.translation-flow-cell-label {
  flex:0 0 auto;
  max-width:34%;
  min-width:72px;
}
.translation-flow-cell-narrow {
  flex:0 1 auto;
  max-width:42%;
}
.translation-flow-cell-body {
  flex:1 1 220px;
  min-width:0;
}
.translation-flow-cell-full {
  flex:1 1 100%;
  width:100%;
}
.translation-flow-cell-table {
  flex:0 1 auto;
  max-width:24%;
  min-width:56px;
  font-size:11px;
  padding:4px 6px;
}
.translation-mini-table {
  width:100%;
  border-collapse:collapse;
  margin:2px 0 4px 0;
  font-size:11px;
}
.translation-mini-table td {
  border:1px solid #dbe1ea;
  padding:4px 6px;
  vertical-align:top;
  white-space:nowrap;
  width:1%;
}
.translation-status { color:#64748b; font-size:13px; margin-bottom:12px; }
.btn-small { padding:6px 10px; font-size:13px; }
.page-nav { display:flex; gap:6px; align-items:center; flex-wrap:wrap; }
.jump-form { display:flex; gap:6px; align-items:center; }
.jump-input {
  width:72px;
  padding:6px 8px;
  border:1px solid #cbd5e1;
  border-radius:6px;
  font-size:13px;
}
.docx-drawing-panel {
  margin:16px;
  padding:14px;
  background:#fff;
  border:1px solid #e2e8f0;
  border-left:4px solid #0069d9;
  border-radius:8px;
}
.docx-drawing-panel h3 {
  margin:0 0 4px 0;
  color:#334155;
  font-size:16px;
}
.docx-drawing-hint {
  margin:0 0 10px 0;
  color:#64748b;
  font-size:12px;
}
.docx-drawing-grid {
  display:grid;
  grid-template-columns:repeat(auto-fit, minmax(220px, 1fr));
  gap:10px;
}
.docx-drawing-thumb {
  display:block;
  width:100%;
  padding:0;
  border:none;
  background:transparent;
  cursor:zoom-in;
}
.docx-drawing-thumb img {
  width:100%;
  height:auto;
  border:1px solid #dbe1ea;
  border-radius:6px;
  background:#fff;
  transition:box-shadow 0.15s ease, transform 0.15s ease;
}
.docx-drawing-thumb:hover img {
  box-shadow:0 4px 14px rgba(15,23,42,0.12);
  transform:translateY(-1px);
}
.docx-drawing-thumb img.highlighted {
  border:2px solid #0069d9;
  box-shadow:0 0 0 2px rgba(0,105,217,0.15);
}
.docx-lightbox {
  position:fixed;
  inset:0;
  z-index:9999;
  background:rgba(15,23,42,0.82);
  display:flex;
  align-items:center;
  justify-content:center;
  padding:24px;
  box-sizing:border-box;
}
.docx-lightbox.hidden {
  display:none;
}
.docx-lightbox-content {
  position:relative;
  max-width:96vw;
  max-height:92vh;
  display:flex;
  flex-direction:column;
  align-items:center;
  gap:12px;
}
.docx-lightbox-content img {
  max-width:96vw;
  max-height:82vh;
  width:auto;
  height:auto;
  border-radius:8px;
  background:#fff;
  box-shadow:0 10px 30px rgba(0,0,0,0.35);
}
.docx-lightbox-toolbar {
  display:flex;
  gap:8px;
  align-items:center;
}
.docx-lightbox-close {
  position:absolute;
  top:-12px;
  right:-12px;
  width:36px;
  height:36px;
  border:none;
  border-radius:999px;
  background:#fff;
  color:#334155;
  font-size:22px;
  line-height:1;
  cursor:pointer;
  box-shadow:0 2px 8px rgba(0,0,0,0.2);
}
</style>
</head>
<body>
    <div class="toolbar">
        <div>
            <strong>{{ filename }}</strong>
            | {{ page_label }} {{ page }} / {{ total_pages }}
        </div>
        <div style="display:flex; gap:8px; align-items:center; flex-wrap:wrap;">
            <div class="page-nav">
                {% if page > 1 %}
                <a class="btn btn-gray btn-small" href="{{ url_for(viewer_route, document_id=document_id, page=page-1) }}">Prev</a>
                {% endif %}
                {% if page < total_pages %}
                <a class="btn btn-gray btn-small" href="{{ url_for(viewer_route, document_id=document_id, page=page+1) }}">Next</a>
                {% endif %}
                {% if supports_page_jump %}
                <form class="jump-form" method="get" action="{{ url_for(viewer_route, document_id=document_id) }}">
                    <input
                        class="jump-input"
                        type="number"
                        name="page"
                        min="1"
                        max="{{ total_pages }}"
                        value="{{ page }}"
                        aria-label="Jump to {{ page_label|lower }}"
                        title="Jump to {{ page_label|lower }} (1-{{ total_pages }})"
                    >
                    <button class="btn btn-gray btn-small" type="submit">Go</button>
                </form>
                {% endif %}
            </div>
            {% if download_url %}
            <a class="btn btn-gray btn-small" href="{{ download_url }}">Download original</a>
            {% endif %}
            {% if translation_enabled %}
            <button class="btn btn-green btn-small" id="reload-translation">Refresh</button>
            <button class="btn btn-purple btn-small" id="export-translation" disabled>Export this page (.txt)</button>
            <a class="btn btn-orange btn-small" href="{{ url_for('export_translation_pdf', document_id=document_id, lang='en') }}">Export full EN PDF</a>
            {% endif %}
            <a href="{{ url_for('home') }}">Back</a>
        </div>
    </div>
    <div class="viewer-layout">
        <div class="original-pane">
            {% if view_mode == 'pdf' %}
                {% if pdf_available %}
                <iframe src="{{ pdf_url }}#page={{ page }}"></iframe>
                {% elif original_text %}
                <div style="padding:16px;">
                    <h3 style="margin-top:0;">Original text — {{ page_label }} {{ page }}</h3>
                    <div class="notice" style="margin-bottom:12px;">
                        <strong>PDF file not found on server.</strong>
                        Re-upload in <strong>Documents</strong> and click <strong>Reindex</strong> to restore the PDF preview.
                        Showing indexed text below.
                    </div>
                    <div class="original-text-box">{{ original_text }}</div>
                </div>
                {% else %}
                <div class="file-missing">
                    <strong>Original PDF not found on server.</strong><br>
                    Re-upload the file in <strong>Documents</strong>, then click <strong>Reindex</strong>.
                </div>
                {% endif %}
            {% else %}
                <div style="padding:16px;">
                    {% if extension == 'docx' and docx_image_indexes %}
                    <div class="docx-drawing-panel">
                        <h3>Related drawings</h3>
                        <p class="docx-drawing-hint">Click any drawing to enlarge. Press Esc to close.</p>
                        <div class="docx-drawing-grid">
                            {% for image_index in docx_image_indexes %}
                            {% set image_url = url_for('docx_image', document_id=document_id, image_index=image_index) %}
                            <button
                                type="button"
                                class="docx-drawing-thumb"
                                data-image-src="{{ image_url }}"
                                aria-label="Open drawing {{ image_index + 1 }}"
                            >
                                <img
                                    src="{{ image_url }}"
                                    alt="DOCX drawing {{ image_index + 1 }}"
                                    {% if highlight_image_index is not none and highlight_image_index == image_index %}class="highlighted"{% endif %}
                                    loading="lazy"
                                />
                            </button>
                            {% endfor %}
                        </div>
                    </div>
                    {% endif %}
                    <h3 style="margin-top:0;">Original (Dutch)</h3>
                    {% if original_text %}
                    <div class="original-text-box">{{ original_text }}</div>
                    {% else %}
                    <div class="file-missing">
                        <strong>Original text not available.</strong><br>
                        Re-upload the file and click <strong>Reindex</strong>.
                    </div>
                    {% endif %}
                </div>
            {% endif %}
        </div>
        {% if translation_enabled %}
        <div class="translation-pane">
            <div class="translation-status" id="translation-status">Loading translation for {{ page_label|lower }} {{ page }}...</div>
            <div class="translation-box">
                <h4>English translation — {{ page_label }} {{ page }} / {{ total_pages }}</h4>
                <div id="translated-text">...</div>
            </div>
        </div>
        {% endif %}
    </div>
    {% if extension == 'docx' and docx_image_indexes %}
    <div id="docx-lightbox" class="docx-lightbox hidden" role="dialog" aria-modal="true" aria-label="Drawing preview">
        <div class="docx-lightbox-content">
            <button type="button" class="docx-lightbox-close" id="docx-lightbox-close" aria-label="Close">×</button>
            <img id="docx-lightbox-img" src="" alt="Expanded drawing" />
            <div class="docx-lightbox-toolbar">
                <a id="docx-lightbox-open" class="btn btn-gray btn-small" href="#" target="_blank" rel="noopener">Open full size</a>
            </div>
        </div>
    </div>
    <script>
    (function() {
        const lightbox = document.getElementById("docx-lightbox");
        const lightboxImg = document.getElementById("docx-lightbox-img");
        const lightboxOpen = document.getElementById("docx-lightbox-open");
        const closeBtn = document.getElementById("docx-lightbox-close");
        if (!lightbox || !lightboxImg) return;

        function openLightbox(src) {
            lightboxImg.src = src;
            lightboxOpen.href = src;
            lightbox.classList.remove("hidden");
            document.body.style.overflow = "hidden";
        }

        function closeLightbox() {
            lightbox.classList.add("hidden");
            lightboxImg.src = "";
            document.body.style.overflow = "";
        }

        document.querySelectorAll(".docx-drawing-thumb").forEach(function(btn) {
            btn.addEventListener("click", function() {
                const src = btn.getAttribute("data-image-src");
                if (src) openLightbox(src);
            });
        });

        closeBtn.addEventListener("click", closeLightbox);
        lightbox.addEventListener("click", function(event) {
            if (event.target === lightbox) closeLightbox();
        });
        document.addEventListener("keydown", function(event) {
            if (event.key === "Escape" && !lightbox.classList.contains("hidden")) {
                closeLightbox();
            }
        });
    })();
    </script>
    {% endif %}
    {% if translation_enabled %}
    <script>
    const documentId = {{ document_id }};
    const pageNum = {{ page }};
    const totalPages = {{ total_pages }};
    const pageLabel = {{ page_label|tojson }};
    const fileBase = {{ filename|tojson }};
    const statusEl = document.getElementById("translation-status");
    const translatedEl = document.getElementById("translated-text");
    const exportBtn = document.getElementById("export-translation");
    let currentTranslation = "";

    async function readJsonResponse(resp) {
        const raw = await resp.text();
        if (!raw.trim()) {
            throw new Error(`Empty server response (${resp.status}).`);
        }
        try {
            return JSON.parse(raw);
        } catch (parseErr) {
            const preview = raw.replace(/\\s+/g, " ").slice(0, 160);
            throw new Error(
                preview.startsWith("<")
                    ? `Server returned HTML instead of JSON (${resp.status}). Try logging in again or redeploying the app.`
                    : `Invalid server response (${resp.status}).`
            );
        }
    }

    async function loadCurrentTranslation() {
        statusEl.textContent = `Translating ${pageLabel.toLowerCase()} ${pageNum} of ${totalPages}...`;
        translatedEl.textContent = "...";
        translatedEl.innerHTML = "";
        currentTranslation = "";
        exportBtn.disabled = true;

        try {
            const resp = await fetch(`/api/document/${documentId}/translate-page?page=${pageNum}`);
            const data = await readJsonResponse(resp);
            if (!resp.ok) {
                throw new Error(data.error || "Translation request failed");
            }

            currentTranslation = data.translation || "";
            if (data.translation_html) {
                translatedEl.innerHTML = data.translation_html;
            } else {
                translatedEl.textContent = currentTranslation || "(Translation unavailable)";
            }
            const provider = data.provider || "unknown";
            statusEl.textContent = data.cached
                ? `${pageLabel} ${pageNum} — cached (${provider}). Use Next to translate the following page.`
                : `${pageLabel} ${pageNum} — translated with ${provider}. Use Next for the next page.`;
            if (data.error) {
                statusEl.textContent = data.error;
            }
            exportBtn.disabled = !currentTranslation.trim();
        } catch (err) {
            statusEl.textContent = "Translation error: " + err.message;
            translatedEl.textContent = "";
            translatedEl.innerHTML = "";
            exportBtn.disabled = true;
        }
    }

    function exportTranslation() {
        if (!currentTranslation.trim()) {
            return;
        }
        const safeName = fileBase.replace(/\\.[^.]+$/, "") || "document";
        const blob = new Blob([currentTranslation], { type: "text/plain;charset=utf-8" });
        const url = URL.createObjectURL(blob);
        const link = document.createElement("a");
        link.href = url;
        link.download = `${safeName}_${pageLabel.toLowerCase()}_${pageNum}_en.txt`;
        document.body.appendChild(link);
        link.click();
        link.remove();
        URL.revokeObjectURL(url);
    }

    document.getElementById("reload-translation").addEventListener("click", loadCurrentTranslation);
    exportBtn.addEventListener("click", exportTranslation);
    loadCurrentTranslation();
    </script>
    {% endif %}
</body>
</html>
"""

AI_TRANSLATE_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>AI Translation Tool</title>
""" + BASE_CSS + """
</head>
<body>
<div class="container">
    <h2>AI Translation Tool</h2>
    <p class="small">
        Provider: <strong>{{ provider_info.provider }}</strong>
        | Model: <strong>{{ provider_info.model }}</strong>
        | Status: <strong>{{ "Ready" if provider_info.configured else "Not configured" }}</strong>
    </p>

    <form method="post">
        <label><strong>Source text (Dutch / mixed)</strong></label>
        <textarea name="source_text" rows="10" style="width:100%; margin-top:8px;">{{ source_text }}</textarea>
        <div style="margin-top:12px;">
            <button class="btn btn-green" type="submit">Translate to English</button>
            <a class="btn btn-gray" href="{{ url_for('home') }}">Back</a>
        </div>
    </form>

    {% if translated_text %}
    <div class="summary-box" style="margin-top:18px;">
        <strong>English translation</strong>
        <div class="full-block" style="margin-top:10px;">{{ translated_text }}</div>
    </div>
    {% endif %}
</div>
</body>
</html>
"""

TABLE_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Table preview</title>
""" + BASE_CSS + """
</head>
<body>
<div class="container">
    <h2>Table preview - {{ doc.original_filename }}</h2>
    <div style="margin:12px 0; display:flex; gap:8px; flex-wrap:wrap;">
        <a class="btn btn-gray" href="{{ url_for('home') }}">Back</a>
        <a class="btn btn-green" href="{{ url_for('refresh_table_preview', document_id=doc.id) }}">Refresh table &amp; translate</a>
        {% if tables %}
        <a class="btn btn-purple" href="{{ url_for('export_table_xlsx', document_id=doc.id) }}">Export Excel</a>
        {% endif %}
    </div>

    {% if tables %}
        {% for t in tables %}
            <div class="result-item">
                <div class="filename">{{ t.preview_title }}</div>
                <div class="meta">Format: {{ t.table_format }} | Sheet: {{ t.sheet_name }}</div>
                <h3>Table 1 - Dutch (Original)</h3>
                <div class="table-scroll-wrap">
                    <div class="data-table">{{ t.html_table|safe }}</div>
                </div>
                {% if t.html_table_en %}
                <h3 style="margin-top:18px;">Table 2 - English (Translation)</h3>
                <div class="table-scroll-wrap">
                    <div class="data-table">{{ t.html_table_en|safe }}</div>
                </div>
                {% else %}
                <div class="notice" style="margin-top:12px;">English translation not available yet. Reindex the document to generate it.</div>
                {% endif %}
            </div>
        {% endfor %}
    {% else %}
        <div class="warning">No table preview available.</div>
        {% if error_msg %}
        <div class="notice" style="margin-top:12px;">{{ error_msg }}</div>
        {% else %}
        <div class="notice" style="margin-top:12px;">
            The table was not indexed yet, or the source file was lost after a server restart.
            Go to <a href="{{ url_for('admin_documents') }}">Documents</a> and click <strong>Reindex</strong> for this file.
        </div>
        {% endif %}
    {% endif %}


</div>
</body>
</html>
"""


DICTIONARY_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Definitions</title>
""" + BASE_CSS + """
</head>
<body>
<div class="container">
    <div class="topbar">
        <div class="topbar-brand">
            <h1>Definitions</h1>
            <div class="small">Abbreviations and reference text lookup</div>
        </div>
        <div class="actions">
            <a class="btn btn-gray" href="{{ url_for('home') }}">Back to search</a>
            <a class="btn btn-green" href="{{ url_for('admin_dictionaries') }}">Manage definitions</a>
        </div>
    </div>

    <div class="dict-search-panel">
        <form method="GET" action="{{ url_for('dictionary_lookup') }}">
            <input type="text" name="q" value="{{ query }}" placeholder="Abbreviation, cabinet code (e.g. 001), or keyword" style="width:72%; min-width:280px;">
            <button type="submit">Look up</button>
            <div class="search-filter-panel" style="margin-top:12px;">
                <div class="panel-title">Types</div>
                <div class="term-filter-row">
                    {% for opt in type_options %}
                    <label>
                        <input type="checkbox" name="type" value="{{ opt.id }}"
                            {% if opt.id in selected_types %}checked{% endif %}>
                        <span>{{ opt.label }}</span>
                    </label>
                    {% endfor %}
                </div>
            </div>
        </form>
        <div class="small" style="margin-top:10px;">
            Abbreviations and cabinet codes match first (e.g. <strong>AC</strong>, <strong>001</strong>, <strong>=__+A001</strong>).
            Results show Dutch and English where available.
        </div>
    </div>

    {% if query %}
        {% if results %}
            <div class="small" style="margin-top:8px;">{{ results|length }} result(s) for <strong>{{ query }}</strong></div>
            {% for entry in results %}
            <div class="dict-result-card">
                <div class="dict-result-kind">{{ entry.kind_label }}</div>
                <div class="dict-result-term">{{ entry.term }}</div>
                {% if entry.kind_label == 'Project Documents' %}
                    {% if entry.content_nl %}
                    <div class="dict-result-content">{{ entry.content_nl }}</div>
                    {% endif %}
                {% else %}
                    {% if entry.content_nl %}
                    <div class="dict-result-lang"><span class="dict-result-lang-label">NL:</span> {{ entry.content_nl }}</div>
                    {% endif %}
                    {% if entry.content_en %}
                    <div class="dict-result-lang"><span class="dict-result-lang-label">EN:</span> {{ entry.content_en }}</div>
                    {% endif %}
                {% endif %}
                {% if not entry.content_nl and not entry.content_en and entry.content %}
                <div class="dict-result-content">{{ entry.content }}</div>
                {% endif %}
                <div class="dict-result-meta">Source: {{ entry.source_name }}</div>
            </div>
            {% endfor %}
        {% else %}
            <div class="dict-empty">No definitions found for <strong>{{ query }}</strong>{% if not selected_types %} (no types selected){% endif %}.</div>
        {% endif %}
    {% else %}
        <div class="notice">
            Examples: <strong>AC</strong>, <strong>001</strong>, <strong>declaration of performance</strong>
        </div>
    {% endif %}
</div>
</body>
</html>
"""


ADMIN_DICTIONARIES_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Manage definitions</title>
""" + BASE_CSS + """
</head>
<body>
<div class="container">
    <h2>Manage definitions</h2>
    <div style="margin-bottom:16px;">
        <a class="btn btn-gray" href="{{ url_for('home') }}">Back</a>
        <a class="btn btn-purple" href="{{ url_for('dictionary_lookup') }}">Definitions</a>
    </div>

    {% with messages = get_flashed_messages(with_categories=true) %}
      {% if messages %}
        <div id="admin-feedback" class="admin-feedback-panel">
        {% for category, msg in messages %}
          <div class="flash flash-{{ category if category != 'message' else 'info' }}">{{ msg }}</div>
        {% endfor %}
        </div>
      {% endif %}
    {% endwith %}

    """ + STORAGE_USAGE_PANEL + """

    <div class="stats-panel" id="upload-panel">
        <div class="panel-title">Upload definitions Excel</div>
        <div class="info" style="margin-bottom:12px;">
            Use a two-column or three-column sheet:<br>
            <strong>Glossary (3 columns):</strong> Abreviation | Full Name | English Translation<br>
            <strong>Cabinet file:</strong> <code>=__+A001</code> | Dutch function (English auto-translated on upload)<br>
            <strong>Project documents:</strong> Document | Content<br>
            Multi-sheet workbooks are supported. Re-uploading the same filename replaces existing entries for that file.
        </div>
        <form method="POST" enctype="multipart/form-data" onsubmit="return handleDictUploadSubmit(this);">
            <input type="file" name="dictionary_file" accept=".xlsx,.xls" required>
            <div style="margin-top:12px;">
                <button type="submit" id="dict-upload-btn" class="btn btn-green">Upload</button>
            </div>
        </form>
    </div>

    <div class="stats-panel">
        <div class="panel-title">Uploaded definition files</div>
        {% if sources %}
        <table class="data-table" style="width:100%; margin-top:12px;">
            <tr>
                <th>Name</th>
                <th>Filename</th>
                <th>Sheet</th>
                <th>Type</th>
                <th>Entries</th>
                <th>Uploaded</th>
                <th>Actions</th>
            </tr>
            {% for row in sources %}
            <tr>
                <td>{{ row.name }}</td>
                <td>{{ row.original_filename }}</td>
                <td>{{ row.sheet_name or '-' }}</td>
                <td>{{ row.kind_label }}</td>
                <td>{{ row.entry_count }}</td>
                <td>{{ row.uploaded_at }}</td>
                <td>
                    <a class="btn btn-orange" href="{{ url_for('delete_dictionary_source', source_id=row.id) }}"
                       onclick="return confirm('Delete this definition file and all {{ row.entry_count }} entries?');">Delete</a>
                </td>
            </tr>
            {% endfor %}
        </table>
        {% else %}
        <div class="notice" style="margin-top:12px;">No definition files uploaded yet.</div>
        {% endif %}
    </div>
</div>
<script>
function handleDictUploadSubmit(form) {
    const btn = form.querySelector('#dict-upload-btn');
    if (!btn || btn.disabled) {
        return false;
    }
    btn.disabled = true;
    btn.textContent = 'Uploading…';
    return true;
}
</script>
</body>
</html>
"""


# =========================================================
# Routes
# =========================================================
def normalize_requirement_lookup_id(raw_query):
    query = (raw_query or "").strip()
    if not query:
        return ""
    match = REQ_ID_REGEX.search(query)
    if match:
        return match.group(1)
    return query


def requirement_lookup_variants(req_id):
    normalized = normalize_requirement_lookup_id(req_id)
    if not normalized:
        return []

    variants = [normalized]
    family_match = re.match(
        r"^([A-Z]{1,10}-Req-\d+)(?:\.\d+)?$",
        normalized,
        re.IGNORECASE,
    )
    if family_match:
        family_id = family_match.group(1)
        if family_id.lower() != normalized.lower():
            variants.append(family_id)
    return variants


def requirement_id_matches_text(req_id, text):
    req = (req_id or "").strip().lower()
    value = (text or "").strip().lower()
    if not req or not value:
        return False
    if req == value:
        return True
    if req in value:
        return True
    req_compact = re.sub(r"[^a-z0-9]+", "", req)
    value_compact = re.sub(r"[^a-z0-9]+", "", value)
    return bool(req_compact and req_compact in value_compact)


def requirement_row_spec_code(row):
    for cell in row.get("cells", []):
        col_key = re.sub(r"[^a-z0-9]+", "", str(cell.get("column", "")).lower())
        if col_key == "speccode":
            return normalize_requirement_lookup_id(cell.get("value", ""))
    return ""


def sort_requirement_rows_by_query(rows, query_id):
    key = normalize_requirement_lookup_id(query_id).lower()
    key_compact = re.sub(r"[^a-z0-9]+", "", key)

    def priority(row):
        spec = requirement_row_spec_code(row).lower()
        spec_compact = re.sub(r"[^a-z0-9]+", "", spec)
        if spec == key or spec_compact == key_compact:
            return 0
        return 1

    return sorted(
        rows,
        key=lambda row: (priority(row), row.get("row_number") or 0),
    )


def _collapse_requirement_text(text):
    cleaned = compact_display_text(text or "")
    return re.sub(r"\s+", " ", cleaned).strip()


def _clean_requirement_field_line(line):
    return re.sub(r"\s+", " ", (line or "").strip())


def _is_requirement_metadata_line(line):
    low = (line or "").strip().lower()
    skip_prefixes = (
        "applicable",
        "verification",
        "figure",
        "explanation",
        "phase",
        "method",
        "remark",
        "test",
        "berekening",
        "calculation",
        "publication date",
        "generated from",
        "beveiliging intelligent",
        "security intelligent",
        "category:",
    )
    return any(low.startswith(prefix) for prefix in skip_prefixes)


def sort_requirement_content_rows(rows, query_id):
    key = normalize_requirement_lookup_id(query_id).lower()
    key_compact = re.sub(r"[^a-z0-9]+", "", key)

    def priority(row):
        spec = (row.get("requirement_id") or "").lower()
        spec_compact = re.sub(r"[^a-z0-9]+", "", spec)
        if spec == key or spec_compact == key_compact:
            return 0
        return 1

    return sorted(
        rows,
        key=lambda row: (priority(row), row.get("page") or 0),
    )


def _looks_like_requirement_category(line):
    low = (line or "").lower()
    return bool(
        re.search(r"\b(systeem|system|aspect|betrouwbaarheid|reliability)\b", low)
    ) and "," in (line or "")


def _parse_structured_requirement_fields(requirement_id, full_text, lang="nl"):
    text = compact_display_text(full_text or "")
    if not text:
        return "", ""

    lines = [
        _clean_requirement_field_line(ln)
        for ln in text.splitlines()
        if _clean_requirement_field_line(ln)
    ]
    title = ""
    applicable = ""

    for line in lines:
        title_match = re.match(r"(?i)^Title:\s*(.+)$", line)
        if title_match:
            title = title_match.group(1).strip()
            break

    stop_pattern = re.compile(
        r"(?i)^(figure|explanation|verification|phase|method|remark|test|berekening|calculation|am-req-)\b"
    )
    disc_idx = -1
    for i, line in enumerate(lines):
        if re.search(
            r"(?i)applicable\s+discriminators|distance\s+security|distantie\s+beveiliging",
            line,
        ):
            disc_idx = i
            break

    if disc_idx >= 0:
        body_parts = []
        for line in lines[disc_idx + 1 :]:
            if stop_pattern.search(line):
                break
            if _is_requirement_metadata_line(line):
                continue
            if re.search(r"(?i)^(distance security|distantie beveiliging)$", line):
                continue
            body_parts.append(line)
        applicable = _collapse_requirement_text(" ".join(body_parts))

    if lang == "nl" and not applicable:
        body_match = re.search(
            r"(De\s+distantiefunctie\s*\(ANSI:21\).+?)"
            r"(?=(?:\s*(?:Figure|Explanation|Verification|Phase|Method|Remark|Test|Berekening)\b|AM-Req-|\Z))",
            text,
            re.IGNORECASE | re.DOTALL,
        )
        if body_match:
            applicable = _collapse_requirement_text(body_match.group(1))
    elif lang == "en" and not applicable:
        body_match = re.search(
            r"(The\s+distance\s+function\s*\(ANSI:21\).+?)"
            r"(?=(?:\s*(?:Figure|Explanation|Verification|Phase|Method|Remark|Test|Calculation)\b|AM-Req-|\Z))",
            text,
            re.IGNORECASE | re.DOTALL,
        )
        if body_match:
            applicable = _collapse_requirement_text(body_match.group(1))

    req_lower = (requirement_id or "").strip().lower()
    if not title and req_lower:
        req_idx = next((i for i, line in enumerate(lines) if req_lower in line.lower()), -1)
        if req_idx >= 0:
            for cand in lines[req_idx + 1 : req_idx + 6]:
                if _is_requirement_metadata_line(cand):
                    continue
                if req_lower in cand.lower():
                    continue
                low = cand.lower()
                if re.search(r"(?i)\b(zone|functie|function|instelling|component|beveiliging\s*-\s*)", cand):
                    title = cand
                    break
            if not title:
                for cand in lines[req_idx + 1 : req_idx + 4]:
                    if _is_requirement_metadata_line(cand) or req_lower in cand.lower():
                        continue
                    if _looks_like_requirement_category(cand):
                        continue
                    if len(cand) > 12:
                        title = cand
                        break

    if lang == "en" and not title:
        for cand in lines:
            if re.search(r"(?i)(distancing function|distance function|forward zone)", cand):
                title = cand
                break

    return title, applicable


def _get_requirement_page_translation_text(doc, page_num):
    page_translation = DocumentPageTranslation.query.filter_by(
        document_id=doc.id, page=page_num
    ).first()
    if page_translation and (page_translation.translated_text or "").strip():
        return page_translation.translated_text or ""

    if (doc.extension or "").lower() != "pdf":
        return ""

    try:
        payload = get_or_translate_page(doc, page_num, target_lang="en")
        return (payload or {}).get("translation") or ""
    except Exception as exc:
        print(f"Requirement page translation error for doc {doc.id} page {page_num}: {exc}")
        return ""


def _build_requirement_content_fields(requirement_id, full_text_nl, full_text_en=""):
    title_nl, applicable_nl = _parse_structured_requirement_fields(
        requirement_id, full_text_nl, lang="nl"
    )
    title_en, applicable_en = ("", "")
    if (full_text_en or "").strip():
        title_en, applicable_en = _parse_structured_requirement_fields(
            requirement_id, full_text_en, lang="en"
        )

    if not title_en and title_nl:
        title_en = (translate_to_english(title_nl) or "").strip()
    if not applicable_en and applicable_nl:
        applicable_en = (translate_to_english(applicable_nl) or "").strip()

    return {
        "title_nl": title_nl,
        "title_en": title_en,
        "applicable_nl": applicable_nl,
        "applicable_en": applicable_en,
    }


def _build_requirement_browser_content_row(block):
    doc = block.document
    if not doc:
        return None
    ext = (doc.extension or "").lower()
    page_num = block.page or 1

    full_text_nl = block.full_text or ""
    translated_text = _get_requirement_page_translation_text(doc, page_num)
    fields = _build_requirement_content_fields(
        block.requirement_id, full_text_nl, translated_text
    )

    return {
        "block_id": block.id,
        "document_id": doc.id,
        "document_name": doc.original_filename or f"Document {doc.id}",
        "requirement_id": block.requirement_id or "",
        "title": block.title or "",
        "section": block.major_section or block.section or "",
        "page": page_num,
        "snippet": compact_display_text(block.full_text or "")[:900].rstrip()
        + ("..." if len(compact_display_text(block.full_text or "")) > 900 else ""),
        **fields,
        "is_pdf": ext == "pdf",
        "is_docx": ext == "docx",
        "is_txt": ext == "txt",
    }


def _build_requirement_browser_content_from_document(doc, requirement_id):
    ext = (doc.extension or "").lower()
    if ext not in {"pdf", "docx", "txt"}:
        return None

    full_text_nl = compact_display_text((doc.text_preview or "").strip())
    if not full_text_nl:
        full_text_nl = compact_display_text(get_document_full_text(doc))
    if not full_text_nl:
        return None

    fields = _build_requirement_content_fields(requirement_id, full_text_nl)
    if not any(fields.values()):
        return None

    return {
        "block_id": None,
        "document_id": doc.id,
        "document_name": doc.original_filename or f"Document {doc.id}",
        "requirement_id": requirement_id or "",
        "title": "",
        "section": "",
        "page": 1,
        "snippet": "",
        **fields,
        "is_pdf": ext == "pdf",
        "is_docx": ext == "docx",
        "is_txt": ext == "txt",
    }


def _append_requirement_rows_from_dataframe(rows, seen, df, doc, sheet_name, variants, max_rows):
    columns = [str(col) for col in df.columns]
    for row_idx, row in df.iterrows():
        values = [cell_to_dictionary_text(row.get(col)) for col in df.columns]
        if not any(
            requirement_id_matches_text(variant, value)
            for variant in variants
            for value in values
        ):
            continue

        row_position = int(cast(Any, row_idx))
        row_key = (doc.id, str(sheet_name or ""), row_position)
        if row_key in seen:
            continue
        seen.add(row_key)

        cells = []
        for col_name, value in zip(columns, values):
            if not value:
                continue
            cells.append({"column": col_name, "value": value})
        if not cells:
            continue

        rows.append(
            {
                "document_id": doc.id,
                "document_name": doc.original_filename or f"Document {doc.id}",
                "sheet_name": str(sheet_name or "-"),
                "table_format": doc.extension or "-",
                "row_number": row_position + 1,
                "cells": cells,
            }
        )
        if len(rows) >= max_rows:
            return True
    return False


def _scan_requirement_master_xlsx_rows(xlsx_path, variants, max_collect=200):
    """
    Memory-safe scan: stream XLSX rows (read_only) and only materialize matched rows.
    """
    wb = load_workbook(filename=xlsx_path, read_only=True, data_only=True)
    matched = []
    sheets_scanned = 0

    try:
        for ws in cast(Any, wb.worksheets):
            # Find header row within first 60 rows (row with most non-empty cells)
            best_row = None
            best_count = 0
            for r_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=60, values_only=True), start=1):
                values = [normalize_cell_text(v) for v in (row or [])]
                count = sum(1 for v in values if v)
                if count > best_count:
                    best_count = count
                    best_row = (r_idx, values)
            if not best_row or best_count < 2:
                continue

            header_row_idx, header_values = best_row
            header_keys = [
                re.sub(r"[^a-z0-9]+", "", (h or "").strip().lower())
                for h in header_values
            ]
            if "speccode" not in header_keys or "sourcedocument" not in header_keys:
                continue

            sheets_scanned += 1
            spec_idx = header_keys.index("speccode")

            # Stream the body rows
            iter_rows_fn = getattr(ws, "iter_rows", None)
            if not callable(iter_rows_fn):
                continue
            for excel_row_num, body_row in enumerate(
                cast(Any, iter_rows_fn)(min_row=header_row_idx + 1, values_only=True),
                start=header_row_idx + 1,
            ):
                if not body_row:
                    continue
                cells = [normalize_cell_text(v) for v in body_row]
                if spec_idx >= len(cells):
                    continue
                spec_val = cells[spec_idx]
                if not spec_val:
                    continue
                if not any(requirement_id_matches_text(v, spec_val) for v in variants):
                    continue

                row_cells = []
                for col_name, value in zip(header_values, cells):
                    col = normalize_cell_text(col_name)
                    if not col:
                        continue
                    if value:
                        row_cells.append({"column": col, "value": value})
                if row_cells:
                    matched.append(
                        {
                            "sheet_name": ws.title,
                            "row_number": excel_row_num,
                            "cells": row_cells,
                        }
                    )
                if len(matched) >= max_collect:
                    return matched, sheets_scanned
    finally:
        try:
            wb.close()
        except Exception:
            pass

    return matched, sheets_scanned


def build_requirement_rows_from_tables(req_id, max_rows=10):
    if not req_id:
        return [], {"documents_scanned": 0, "sheets_scanned": 0, "elapsed_ms": 0}

    variants = requirement_lookup_variants(req_id)
    if not variants:
        return [], {"documents_scanned": 0, "sheets_scanned": 0, "elapsed_ms": 0}

    started = time.perf_counter()
    normalized_id = normalize_requirement_lookup_id(req_id)
    cache_key = (normalized_id.lower(), max_rows)
    cached = _requirement_master_result_cache.get(cache_key)
    now_ts = time.time()
    if cached and (now_ts - cached.get("ts", 0)) <= REQUIREMENT_MASTER_RESULT_CACHE_TTL_SECONDS:
        return cached.get("rows", []), cached.get("stats", {"documents_scanned": 0, "sheets_scanned": 0, "elapsed_ms": 0})

    rows = []
    seen = set()
    documents_scanned = 0
    sheets_scanned = 0
    try:
        tennet_docs = _db_recover_and_retry(
            lambda: (
                DocumentRecord.query.filter(
                    db.or_(
                        DocumentRecord.original_filename.ilike("%tennet%requirement%"),
                        DocumentRecord.stored_filename.ilike("%tennet%requirement%"),
                    )
                )
                .order_by(DocumentRecord.id.desc())
                .limit(REQUIREMENT_LOOKUP_MAX_MASTER_DOCS)
                .all()
            ),
            retries=1,
            label="tennet_docs",
        )
    except Exception as exc:
        print(f"Requirement master DB query failed: {exc}")
        return [], {
            "documents_scanned": 0,
            "sheets_scanned": 0,
            "elapsed_ms": int((time.perf_counter() - started) * 1000),
        }

    # Memory-safe path: stream scan source XLSX (avoid pandas/full table preview).
    for doc in tennet_docs:
        if (time.perf_counter() - started) > REQUIREMENT_LOOKUP_MAX_SCAN_SECONDS:
            break
        if not storage.document_exists(doc.stored_filename, DOC_FOLDER):
            continue
        documents_scanned += 1
        try:
            with storage.open_document_local_path(doc.stored_filename, DOC_FOLDER) as full_path:
                matched_rows, sheet_count = _scan_requirement_master_xlsx_rows(
                    full_path, variants, max_collect=max(max_rows * 20, 100)
                )
                sheets_scanned += sheet_count
        except Exception as exc:
            print(f"Requirement master scan error for {doc.original_filename}: {exc}")
            continue
        for m in matched_rows:
            row_key = (doc.id, m.get("sheet_name") or "", m.get("row_number") or 0)
            if row_key in seen:
                continue
            seen.add(row_key)
            rows.append(
                {
                    "document_id": doc.id,
                    "document_name": doc.original_filename or f"Document {doc.id}",
                    "sheet_name": m.get("sheet_name") or "-",
                    "table_format": doc.extension or "-",
                    "row_number": m.get("row_number") or 0,
                    "cells": m.get("cells") or [],
                }
            )

    rows = sort_requirement_rows_by_query(rows, normalized_id)[:max_rows]

    elapsed_ms = int((time.perf_counter() - started) * 1000)
    stats = {
        "documents_scanned": documents_scanned,
        "sheets_scanned": sheets_scanned,
        "elapsed_ms": elapsed_ms,
    }
    _requirement_master_result_cache[cache_key] = {"ts": now_ts, "rows": rows, "stats": stats}
    return rows, stats


def build_requirement_lookup_result(raw_query, max_blocks=8, max_tables=8):
    query = (raw_query or "").strip()
    normalized = normalize_requirement_lookup_id(query)
    if not normalized:
        return {
            "found": False,
            "normalized_id": "",
            "blocks": [],
            "tables": [],
            "block_count": 0,
            "table_count": 0,
            "document_count": 0,
            "scan_elapsed_ms": 0,
            "scan_documents": 0,
            "scan_sheets": 0,
        }

    requirement_rows, row_scan_stats = build_requirement_rows_from_tables(normalized, max_rows=max_tables)
    variants = requirement_lookup_variants(normalized)
    key = normalized.lower()

    exact_blocks = (
        RequirementBlock.query
        .filter(db.func.lower(RequirementBlock.requirement_id) == key)
        .order_by(RequirementBlock.id.desc())
        .all()
    )

    variant_filters = [
        RequirementBlock.requirement_id.ilike(f"%{variant}%")
        for variant in variants
    ]
    if not variant_filters:
        variant_filters = [RequirementBlock.requirement_id.ilike(f"%{normalized}%")]

    if exact_blocks:
        matched_blocks = exact_blocks
    else:
        matched_blocks = (
            RequirementBlock.query
            .filter(db.or_(*variant_filters))
            .order_by(RequirementBlock.id.desc())
            .all()
        )

    block_rows = []
    seen_block_docs = set()
    for block in matched_blocks:
        if len(block_rows) >= max_blocks:
            break
        doc = block.document
        if not doc:
            continue
        doc_key = (doc.id, block.requirement_id or "", block.page or 1)
        if doc_key in seen_block_docs:
            continue
        seen_block_docs.add(doc_key)
        formatted_row = _build_requirement_browser_content_row(block)
        if formatted_row:
            block_rows.append(formatted_row)

    matched_doc_ids = {row["document_id"] for row in requirement_rows if row.get("document_id")}
    matched_doc_ids.update(row["document_id"] for row in block_rows if row.get("document_id"))

    filename_filters = [
        DocumentRecord.original_filename.ilike(f"%{variant}%")
        for variant in variants
    ]
    if not filename_filters:
        filename_filters = [DocumentRecord.original_filename.ilike(f"%{normalized}%")]
    filename_docs = DocumentRecord.query.filter(db.or_(*filename_filters)).all()

    stored_name_filters = [
        DocumentRecord.stored_filename.ilike(f"%{variant}%")
        for variant in variants
    ]
    if not stored_name_filters:
        stored_name_filters = [DocumentRecord.stored_filename.ilike(f"%{normalized}%")]
    stored_name_docs = DocumentRecord.query.filter(db.or_(*stored_name_filters)).all()
    matched_doc_ids.update(d.id for d in filename_docs)
    matched_doc_ids.update(d.id for d in stored_name_docs)

    content_rows = list(block_rows)
    seen_content_keys = {
        (row.get("document_id"), row.get("requirement_id") or "")
        for row in content_rows
    }
    related_docs = {}
    for doc in filename_docs + stored_name_docs:
        if doc and doc.id:
            related_docs[doc.id] = doc
    for doc_id, doc in related_docs.items():
        doc_req_id = normalize_requirement_lookup_id(doc.original_filename or "") or normalized
        content_key = (doc_id, doc_req_id)
        if content_key in seen_content_keys:
            continue
        formatted_row = _build_requirement_browser_content_from_document(doc, doc_req_id)
        if not formatted_row:
            continue
        seen_content_keys.add(content_key)
        content_rows.append(formatted_row)
        if len(content_rows) >= max_blocks:
            break

    content_rows = sort_requirement_content_rows(content_rows, normalized)[:max_blocks]

    if not matched_doc_ids:
        table_doc_filters = []
        for variant in variants:
            table_doc_filters.extend(
                [
                    TablePreview.csv_text.ilike(f"%{variant}%"),
                    TablePreview.csv_text_en.ilike(f"%{variant}%"),
                    TablePreview.html_table.ilike(f"%{variant}%"),
                    TablePreview.html_table_en.ilike(f"%{variant}%"),
                ]
            )
        if table_doc_filters:
            related_table_docs = TablePreview.query.filter(db.or_(*table_doc_filters)).all()
            matched_doc_ids.update(row.document_id for row in related_table_docs if row.document_id)

    table_rows = []
    if matched_doc_ids:
        tables = (
            TablePreview.query.filter(TablePreview.document_id.in_(matched_doc_ids))
            .order_by(TablePreview.id.desc())
            .all()
        )
        seen_table_keys = set()
        for row in tables:
            key_tuple = (row.document_id, row.sheet_name or "", row.table_format or "")
            if key_tuple in seen_table_keys:
                continue
            seen_table_keys.add(key_tuple)
            doc = row.document
            table_rows.append(
                {
                    "document_id": row.document_id,
                    "document_name": doc.original_filename if doc else f"Document {row.document_id}",
                    "sheet_name": row.sheet_name or "-",
                    "table_format": row.table_format or "-",
                    "html_table": render_table_preview_html(row, "nl"),
                    "html_table_en": "" if should_skip_table_translation(doc) else render_table_preview_html(row, "en"),
                }
            )
            if len(table_rows) >= max_tables:
                break

    return {
        "found": bool(content_rows or table_rows or requirement_rows),
        "normalized_id": normalized,
        "requirement_rows": requirement_rows,
        "blocks": content_rows,
        "tables": table_rows,
        "block_count": len(content_rows),
        "requirement_row_count": len(requirement_rows),
        "table_count": len(table_rows),
        "document_count": len(matched_doc_ids),
        "scan_elapsed_ms": row_scan_stats.get("elapsed_ms", 0),
        "scan_documents": row_scan_stats.get("documents_scanned", 0),
        "scan_sheets": row_scan_stats.get("sheets_scanned", 0),
    }


@app.route("/")
def home():
    query = request.args.get("q", "").strip()
    search_mode = "general"
    req_lookup_query = ""
    req_lookup = None
    try:
        current_page = max(1, request.args.get("page", 1, type=int) or 1)
        total_pages = 1
        total_results = 0
        results = []
        all_expanded_terms = []
        selected_terms = []
        exact_terms = []
        summary = {"top_categories": [], "top_requirement_ids": []}
        document_matches = []
        table_results = []
        result_types = ["drawings", "text", "tables"]

        if query:
            all_expanded_terms = get_all_expanded_terms(query)
            exact_terms = default_exact_search_terms(query)
            if request.args.getlist("term"):
                selected_terms = resolve_active_search_terms(query, request.args.getlist("term"))
            else:
                selected_terms = exact_terms

            requested_result_types = [
                t for t in request.args.getlist("result_type") if t in {"drawings", "text", "tables"}
            ]
            if requested_result_types:
                result_types = requested_result_types

            if "drawings" in result_types:
                document_matches = search_documents(query, active_terms=selected_terms)
            if "tables" in result_types:
                table_results = search_table_results(query, active_terms=selected_terms)
            if "text" in result_types:
                all_results, _ = search_requirements(query, active_terms=selected_terms)
                if table_results:
                    table_doc_ids = {t["document_id"] for t in table_results}
                    all_results = [
                        r for r in all_results
                        if r.get("document_id") not in table_doc_ids
                    ]
                summary = grouped_search_summary(all_results)
                total_results = len(all_results)
                total_pages = max(1, math.ceil(total_results / RESULTS_PER_PAGE))
                current_page = min(current_page, total_pages)
                start_idx = (current_page - 1) * RESULTS_PER_PAGE
                end_idx = start_idx + RESULTS_PER_PAGE
                results = all_results[start_idx:end_idx]

        return _db_recover_and_retry(
            lambda: render_template_string(
            HOME_TEMPLATE,
            query=query,
            search_mode=search_mode,
            requirement_form_action=url_for("requirement_browser"),
            results=results,
            document_matches=document_matches,
            table_results=table_results,
            all_expanded_terms=all_expanded_terms,
            selected_terms=selected_terms,
            exact_terms=exact_terms,
            result_types=result_types,
            summary=summary,
            current_page=current_page,
            total_pages=total_pages,
            total_results=total_results,
            results_per_page=RESULTS_PER_PAGE,
            req_lookup_query=req_lookup_query,
            req_lookup=req_lookup,
            doc_count=DocumentRecord.query.count(),
            block_count=RequirementBlock.query.count(),
            table_count=TablePreview.query.count(),
            ),
            retries=1,
            label="home_render",
        )
    except Exception as exc:
        db.session.rollback()
        print(f"Home render error: {exc}")
        safe_lookup = req_lookup or {
            "found": False,
            "normalized_id": req_lookup_query,
            "requirement_rows": [],
            "blocks": [],
            "tables": [],
            "block_count": 0,
            "requirement_row_count": 0,
            "table_count": 0,
            "document_count": 0,
            "scan_elapsed_ms": 0,
            "scan_documents": 0,
            "scan_sheets": 0,
            "error_message": str(exc),
        }
        safe_lookup["error_message"] = safe_lookup.get("error_message") or str(exc)
        return render_template_string(
            HOME_TEMPLATE,
            query=query,
            search_mode=search_mode,
            requirement_form_action=url_for("requirement_browser"),
            results=[],
            document_matches=[],
            table_results=[],
            all_expanded_terms=[],
            selected_terms=[],
            exact_terms=[],
            result_types=["drawings", "text", "tables"],
            summary={"top_categories": [], "top_requirement_ids": []},
            current_page=1,
            total_pages=1,
            total_results=0,
            results_per_page=RESULTS_PER_PAGE,
            req_lookup_query=req_lookup_query,
            req_lookup=safe_lookup,
            doc_count=DocumentRecord.query.count(),
            block_count=RequirementBlock.query.count(),
            table_count=TablePreview.query.count(),
        )


@app.route("/requirement-browser")
def requirement_browser():
    req_lookup_query = request.args.get("req_lookup", "").strip()
    req_lookup = None
    if req_lookup_query:
        try:
            req_lookup = build_requirement_lookup_result(req_lookup_query)
        except Exception as exc:
            db.session.rollback()
            print(f"Requirement Browser error for '{req_lookup_query}': {exc}")
            req_lookup = {
                "found": False,
                "normalized_id": req_lookup_query,
                "requirement_rows": [],
                "blocks": [],
                "tables": [],
                "block_count": 0,
                "requirement_row_count": 0,
                "table_count": 0,
                "document_count": 0,
                "scan_elapsed_ms": 0,
                "scan_documents": 0,
                "scan_sheets": 0,
                "error_message": str(exc),
            }
    safe_lookup = req_lookup or {
        "found": False,
        "normalized_id": req_lookup_query,
        "requirement_rows": [],
        "blocks": [],
        "tables": [],
        "block_count": 0,
        "requirement_row_count": 0,
        "table_count": 0,
        "document_count": 0,
        "scan_elapsed_ms": 0,
        "scan_documents": 0,
        "scan_sheets": 0,
    }
    return render_template_string(
        HOME_TEMPLATE,
        query="",
        search_mode="requirement",
        requirement_form_action=url_for("requirement_browser"),
        results=[],
        document_matches=[],
        table_results=[],
        all_expanded_terms=[],
        selected_terms=[],
        exact_terms=[],
        result_types=["drawings", "text", "tables"],
        summary={"top_categories": [], "top_requirement_ids": []},
        current_page=1,
        total_pages=1,
        total_results=0,
        results_per_page=RESULTS_PER_PAGE,
        req_lookup_query=req_lookup_query,
        req_lookup=safe_lookup,
        doc_count=DocumentRecord.query.count(),
        block_count=RequirementBlock.query.count(),
        table_count=TablePreview.query.count(),
    )


def build_dictionary_result_rows(entries):
    rows = []
    for entry in entries:
        source = entry.source
        content_nl = (entry.content_nl or "").strip()
        content_en = (entry.content_en or "").strip()
        legacy_content = (entry.content or "").strip()
        if not content_nl and not content_en and legacy_content:
            if entry.entry_kind == "abbreviation":
                content_en = legacy_content
            else:
                content_nl = legacy_content
        if entry.entry_kind == "cabinet" and content_nl and not content_en:
            content_en = ensure_dictionary_english_text(content_nl, "", translate_if_missing=True)
        rows.append(
            {
                "term": entry.term,
                "content": legacy_content,
                "content_nl": content_nl,
                "content_en": content_en,
                "kind_label": dictionary_entry_kind_label(entry.entry_kind),
                "source_name": source.name if source else "",
            }
        )
    return rows


@app.route("/dictionary")
def dictionary_lookup():
    query = request.args.get("q", "").strip()
    if "type" in request.args:
        selected_types = [
            value for value in request.args.getlist("type") if value in DICTIONARY_TYPE_IDS
        ]
    else:
        selected_types = list(DICTIONARY_TYPE_IDS)
    entry_kinds = resolve_dictionary_entry_kinds(selected_types)
    entries = (
        search_dictionary_entries(query, entry_kinds=entry_kinds)
        if query
        else []
    )
    return render_template_string(
        DICTIONARY_TEMPLATE,
        query=query,
        results=build_dictionary_result_rows(entries),
        type_options=DICTIONARY_TYPE_OPTIONS,
        selected_types=selected_types,
    )


@app.route("/admin/dictionaries", methods=["GET", "POST"])
def admin_dictionaries():
    if request.method == "POST":
        upload = request.files.get("dictionary_file")
        if not upload or not upload.filename:
            flash("No definition file selected.", "warning")
            return redirect(url_for("admin_dictionaries"))

        original_filename = secure_filename(upload.filename or "")
        ext = file_ext(original_filename)
        if ext not in {"xlsx", "xls"}:
            flash("Definition upload must be an Excel file (.xlsx).", "error")
            return redirect(url_for("admin_dictionaries"))

        temp_fd, temp_path = tempfile.mkstemp(suffix=f".{ext}")
        os.close(temp_fd)
        try:
            upload.save(temp_path)
            imported_sources, skipped_sheets = import_dictionary_excel(
                temp_path, original_filename
            )
            if not imported_sources:
                detail = ""
                if skipped_sheets:
                    detail = f" Sheets skipped: {', '.join(skipped_sheets)}."
                flash(
                    "No definition rows found."
                    + detail
                    + " For Cabinets, use column A = cabinet code (e.g. =__+A001) "
                    "and column B = Dutch description. Filename or sheet name should contain 'Cabinets'.",
                    "error",
                )
            else:
                parts = [
                    f"{item['entry_count']} {dictionary_entry_kind_label(item['entry_kind']).lower()} "
                    f"from {item['source'].sheet_name or item['source'].name}"
                    for item in imported_sources
                ]
                message = "Definitions imported: " + "; ".join(parts) + "."
                if skipped_sheets:
                    message += f" Skipped empty sheets: {', '.join(skipped_sheets)}."
                flash(message, "success")
        except Exception as exc:
            db.session.rollback()
            print(f"Dictionary import error: {exc}")
            flash(f"Definition import failed: {exc}", "error")
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

        return redirect(url_for("admin_dictionaries"))

    sources = DictionarySource.query.order_by(DictionarySource.uploaded_at.desc()).all()
    source_rows = [
        {
            "id": source.id,
            "name": source.name,
            "original_filename": source.original_filename,
            "sheet_name": source.sheet_name,
            "entry_count": source.entry_count,
            "uploaded_at": source.uploaded_at,
            "kind_label": dictionary_entry_kind_label(source.entry_kind),
        }
        for source in sources
    ]
    return render_template_string(ADMIN_DICTIONARIES_TEMPLATE, sources=source_rows, storage_usage=build_storage_usage_summary())


@app.route("/admin/dictionaries/delete/<int:source_id>")
def delete_dictionary_source(source_id):
    source = db.session.get(DictionarySource, source_id)
    if not source:
        flash("Definition file not found.", "error")
        return redirect(url_for("admin_dictionaries"))

    storage.delete_document(source.stored_filename, DICT_FOLDER)
    db.session.delete(source)
    db.session.commit()
    flash(f"Deleted definition file: {source.name}", "success")
    return redirect(url_for("admin_dictionaries"))


@app.route("/healthz")
def healthz():
    storage_status = storage.get_storage_status()
    return jsonify(
        {
            "status": "ok" if storage_status.get("ok", True) else "degraded",
            "environment": APP_ENV,
            "storage_backend": storage.storage_backend_name(),
            "storage": storage_status,
            "data_dir": str(DATA_DIR),
            "database_path": str(DATABASE_PATH),
            "translation": get_translation_status(),
        }
    )


@app.route("/ai-translate", methods=["GET", "POST"])
def ai_translate_tool():
    source_text = ""
    translated_text = ""
    provider_info = get_translation_status()

    if request.method == "POST":
        source_text = request.form.get("source_text", "").strip()
        if source_text:
            translated_text = translate_to_english(source_text)

    return render_template_string(
        AI_TRANSLATE_TEMPLATE,
        source_text=source_text,
        translated_text=translated_text,
        provider_info=provider_info,
    )


@app.route("/api/translate", methods=["POST"])
def api_translate_text():
    payload = request.get_json(silent=True) or {}
    source_text = (payload.get("text") or "").strip()
    target_lang = normalize_translation_lang(payload.get("lang", "en"))
    if not source_text:
        return jsonify({"error": "text is required"}), 400
    if not translation.translation_enabled():
        return jsonify({"error": "translation is disabled"}), 503

    translated = translate_to_language(source_text, target_lang)
    return jsonify(
        {
            "source": source_text,
            "translation": translated,
            "lang": target_lang,
            "provider": get_translation_provider(),
        }
    )


@app.route("/api/document/<int:document_id>/translate-page")
def api_translate_page(document_id):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404
    if not translation.translation_enabled():
        return jsonify({"error": "translation is disabled"}), 503

    page = request.args.get("page", 1, type=int)
    target_lang = normalize_translation_lang(request.args.get("lang", "en"))
    try:
        result = get_or_translate_page(doc, page, target_lang=target_lang)
    except Exception as exc:
        print(f"Translate page API error for doc {document_id} page {page}: {exc}")
        return jsonify(
            {
                "page": page,
                "lang": target_lang,
                "translation": "",
                "error": f"Translation failed: {exc}",
            }
        ), 500

    if result.get("error") and not result.get("translation"):
        return jsonify(result), 404
    return jsonify(result)


@app.route("/api/document/<int:document_id>/translate-all")
def api_translate_all(document_id):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404
    if not translation.translation_enabled():
        return jsonify({"error": "translation is disabled"}), 503

    total_pages = count_document_pages(doc)
    if total_pages <= 0:
        return jsonify({"error": "No pages found for this document."}), 404

    export_limit = MAX_PDF_PAGES if MAX_PDF_PAGES > 0 else total_pages
    pages_to_load = min(total_pages, export_limit)
    results = []
    for page_num in range(1, pages_to_load + 1):
        result = get_or_translate_page(doc, page_num)
        results.append(
            {
                "page": page_num,
                "translation": result.get("translation") or "",
                "provider": result.get("provider") or get_translation_provider(),
                "cached": bool(result.get("cached")),
                "error": result.get("error", ""),
            }
        )

    return jsonify(
        {
            "total_pages": total_pages,
            "loaded_pages": pages_to_load,
            "truncated": pages_to_load < total_pages,
            "pages": results,
        }
    )


@app.route("/document/<int:document_id>/export-translation-pdf")
def export_translation_pdf(document_id):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc or doc.extension.lower() not in TRANSLATABLE_EXTENSIONS:
        abort(404)

    if not translation.translation_enabled():
        flash("Translation is disabled.")
        if doc.extension == "docx":
            return redirect(url_for("docx_viewer", document_id=document_id, page=1))
        if doc.extension == "txt":
            return redirect(url_for("document_view", document_id=document_id, page=1))
        return redirect(url_for("pdf_viewer", document_id=document_id, page=1))

    if not storage.document_exists(doc.stored_filename, DOC_FOLDER) and doc.extension == "pdf":
        flash("Source PDF not found. Re-upload the file and try again.")
        return redirect(url_for("pdf_viewer", document_id=document_id, page=1))

    try:
        target_lang = normalize_translation_lang(request.args.get("lang", "en"))
        pdf_bytes = build_translated_pdf_bytes(doc, target_lang=target_lang)
    except Exception as exc:
        print(f"Translated PDF export error for document {document_id}: {exc}")
        flash(f"Export failed: {exc}")
        if doc.extension == "docx":
            return redirect(url_for("docx_viewer", document_id=document_id, page=1))
        return redirect(url_for("pdf_viewer", document_id=document_id, page=1))

    base_name = secure_filename(
        Path(doc.original_filename).stem or f"document_{document_id}"
    )
    lang_suffix = "english"
    return send_file(
        io.BytesIO(pdf_bytes),
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"{base_name}_{lang_suffix}.pdf",
    )

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")

        user = User.query.filter_by(username=username).first()
        if not user or not user.check_password(password):
            flash("Invalid username or password.")
            return redirect(url_for("login"))

        login_user(user)
        flash("Login successful.")
        return redirect(url_for("home"))

    return render_template_string(LOGIN_TEMPLATE)


@app.route("/logout")
def logout():
    logout_user()
    flash("Logged out.")
    return redirect(url_for("home"))


@app.route("/admin/users", methods=["GET", "POST"])
def admin_users():
    return redirect(url_for("home"))

def build_admin_documents_page_context():
    docs = DocumentRecord.query.order_by(DocumentRecord.uploaded_at.desc()).all()
    stored_files = storage.list_stored_document_filenames(DOC_FOLDER)
    if stored_files is None:
        stored_files = {
            doc.stored_filename
            for doc in docs
            if storage.document_exists(doc.stored_filename, DOC_FOLDER)
        }

    return {
        "doc_rows": [
            {
                "record": doc,
                "file_available": doc.stored_filename in stored_files,
            }
            for doc in docs
        ],
        "doc_count": len(docs),
        "block_count": RequirementBlock.query.count(),
        "table_count": TablePreview.query.count(),
        "storage_usage": build_storage_usage_summary(),
        "allowed_extensions": ", ".join(sorted(ALLOWED_EXTENSIONS)),
        "storage_backend": storage.storage_backend_name(),
        "storage_persistent": storage.object_storage_enabled(),
        "storage_status": storage.get_storage_status(),
    }


def redirect_admin_documents(anchor=""):
    url = url_for("admin_documents")
    if anchor:
        url = f"{url}#{anchor}"
    return redirect(url)


def flash_upload_results(results):
    if not results:
        flash("No files were uploaded.", "warning")
        return
    for category, message in results:
        flash(message, category)


@app.route("/upload", methods=["GET", "POST"])
def upload_files():
    if request.method == "GET":
        return redirect(url_for("admin_documents"))

    try:
        return _handle_upload_post()
    except Exception as exc:
        db.session.rollback()
        print(f"Upload route error: {exc}")
        flash(f"Upload failed: {exc}", "error")
        return redirect_admin_documents("admin-feedback")


def _handle_upload_post():
    storage_status = storage.get_storage_status()
    if storage.object_storage_enabled() and not storage_status.get("ok"):
        flash(
            storage_status.get("hint")
            or storage_status.get("message")
            or "Storage is not configured.",
            "error",
        )
        return redirect_admin_documents("admin-feedback")

    files = request.files.getlist("files")
    if not files or all(f.filename == "" for f in files):
        flash("No files selected.", "warning")
        return redirect_admin_documents("admin-feedback")

    uploaded_count = 0
    skipped_count = 0
    restored_count = 0
    results = []

    for file in files:
        if not file or file.filename == "":
            continue
        if not allowed_file(file.filename):
            results.append(("error", f"File type not allowed: {file.filename}"))
            continue

        original_filename = secure_filename(file.filename or "")
        if not original_filename:
            results.append(("error", "Invalid filename."))
            continue
        ext = file_ext(original_filename)

        temp_fd, temp_path = tempfile.mkstemp(suffix=f".{ext}")
        os.close(temp_fd)
        try:
            file.save(temp_path)
            file_hash = sha256_of_file(temp_path)

            existing = dedupe_lookup(file_hash)
            if existing:
                if not storage.document_exists(existing.stored_filename, DOC_FOLDER):
                    try:
                        restore_document_file_from_temp(existing, temp_path)
                        index_document_record(existing)
                        results.append(
                            (
                                "success",
                                f"Restored missing file and reindexed: {original_filename}",
                            )
                        )
                        restored_count += 1
                    except Exception as exc:
                        db.session.rollback()
                        results.append(
                            (
                                "error",
                                f"Could not restore duplicate file {original_filename}: {exc}",
                            )
                        )
                else:
                    results.append(
                        ("warning", f"Duplicate file skipped: {original_filename}")
                    )
                    skipped_count += 1
                continue

            stored_filename = f"{uuid.uuid4().hex}.{ext}"
            try:
                size_bytes = storage.save_document(
                    stored_filename, temp_path, DOC_FOLDER
                )
                storage.verify_document_stored(stored_filename, DOC_FOLDER)
            except Exception as exc:
                results.append(
                    ("error", f"Upload failed for {original_filename}: {exc}")
                )
                continue

            doc_cls = cast(Any, DocumentRecord)
            doc = doc_cls(
                original_filename=original_filename,
                stored_filename=stored_filename,
                extension=ext,
                file_hash=file_hash,
                size_bytes=size_bytes,
                uploaded_by=current_user.id if current_user.is_authenticated else None,
            )
            db.session.add(doc)
            try:
                db.session.commit()
            except Exception as exc:
                db.session.rollback()
                storage.delete_document(stored_filename, DOC_FOLDER)
                results.append(
                    ("error", f"Database error for {original_filename}: {exc}")
                )
                continue

            try:
                index_document_record(doc)
            except Exception as exc:
                db.session.rollback()
                print(f"Indexing error for {original_filename}: {exc}")
                results.append(
                    (
                        "warning",
                        f"File saved but indexing failed for {original_filename}: {exc}",
                    )
                )

            uploaded_count += 1
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    summary_parts = []
    if uploaded_count:
        summary_parts.append(f"{uploaded_count} uploaded and indexed")
    if restored_count:
        summary_parts.append(f"{restored_count} restored")
    if skipped_count:
        summary_parts.append(f"{skipped_count} duplicate(s) skipped")
    if summary_parts:
        results.insert(0, ("success", "Upload complete: " + ", ".join(summary_parts) + "."))

    flash_upload_results(results)
    return redirect_admin_documents("admin-feedback")

@app.route("/admin/documents")
def admin_documents():
    return render_template_string(DOCS_TEMPLATE, **build_admin_documents_page_context())

@app.route("/reindex/<int:document_id>")
def reindex_document(document_id):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc:
        flash("Document not found.", "error")
        return redirect_admin_documents("admin-feedback")

    try:
        index_document_record(doc)
        block_count = RequirementBlock.query.filter_by(document_id=doc.id).count()
        flash(
            f"Reindexed: {doc.original_filename} ({block_count} requirement blocks)",
            "success",
        )
    except Exception as e:
        db.session.rollback()
        print(f"Reindex error for document {document_id}: {e}")
        flash(f"Reindex error: {str(e)}", "error")

    return redirect_admin_documents("admin-feedback")


@app.route("/reindex-multiple", methods=["POST"])
def bulk_reindex_documents():
    raw_ids = request.form.getlist("document_ids")
    if not raw_ids:
        flash("No documents selected.", "warning")
        return redirect_admin_documents("admin-feedback")

    reindexed_count = 0
    missing_count = 0
    error_count = 0
    error_names = []

    for raw_id in raw_ids:
        try:
            document_id = int(raw_id)
        except ValueError:
            error_count += 1
            continue

        doc = db.session.get(DocumentRecord, document_id)
        if not doc:
            missing_count += 1
            continue

        try:
            index_document_record(doc)
            reindexed_count += 1
        except Exception as exc:
            db.session.rollback()
            error_count += 1
            error_names.append(doc.original_filename)
            print(f"Bulk reindex error for document {document_id}: {exc}")

    summary = f"Bulk reindex completed. Reindexed: {reindexed_count}, Missing: {missing_count}, Errors: {error_count}."
    if error_names:
        preview = ", ".join(error_names[:5])
        if len(error_names) > 5:
            preview += f" (+{len(error_names) - 5} more)"
        summary += f" Failed: {preview}."
    flash(summary, "success" if error_count == 0 else "warning")
    return redirect_admin_documents("admin-feedback")


@app.route("/delete/<int:document_id>")
def delete_document(document_id):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc:
        flash("Document not found.", "error")
        return redirect_admin_documents("admin-feedback")

    storage.delete_document(doc.stored_filename, DOC_FOLDER)

    # DB cascade will remove requirements & tables
    db.session.delete(doc)
    db.session.commit()
    flash(f"Deleted: {doc.original_filename}", "success")
    return redirect_admin_documents("admin-feedback")

@app.route("/requirement/<int:block_id>")
def requirement_detail(block_id):
    block = db.session.get(RequirementBlock, block_id)
    if not block:
        abort(404)
    block_view = {
        "requirement_id": block.requirement_id,
        "title": block.title,
        "section": block.section,
        "major_section": getattr(block, "major_section", "") or "",
        "page": block.page,
        "category": block.category,
        "definition": block.definition,
        "summary": block.summary,
        "full_text": block.full_text,
        "full_text_en": translate_to_english(block.full_text),
        "document": getattr(block, "document", None),
    }
    return render_template_string(
        REQUIREMENT_TEMPLATE,
        block=block_view,
        pdf2image_available=PDF2IMAGE_AVAILABLE
    )

@app.route("/table/<int:document_id>")
def table_preview(document_id):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc:
        abort(404)

    tables = TablePreview.query.filter_by(document_id=document_id).all()
    error_msg = ""
    file_available = storage.document_exists(doc.stored_filename, DOC_FOLDER)
    skip_translation = should_skip_table_translation(doc)

    if (
        file_available
        and doc.extension.lower() in {"csv", "xlsx"}
        and table_preview_needs_regenerate(tables)
    ):
        try:
            tables = regenerate_table_previews(doc)
            if not tables:
                error_msg = "The spreadsheet has no readable sheets or table data."
        except FileNotFoundError:
            file_available = False
        except Exception as exc:
            print(f"Table preview regeneration error for document {document_id}: {exc}")
            error_msg = f"Could not build table preview: {exc}"

    if tables and table_preview_needs_translation(tables):
        refresh_table_translations_from_cache(tables, skip_translation=skip_translation)
        tables = TablePreview.query.filter_by(document_id=document_id).all()

    if not tables and not file_available:
        error_msg = (
            f"Source file '{doc.original_filename}' was not found in storage. "
            "Please go to Documents, delete this entry, upload the Excel file again, "
            "then click Reindex."
        )
    elif not tables and file_available:
        error_msg = error_msg or "No readable table data was found in this spreadsheet."

    table_views = []
    for table_row in tables:
        html_en = "" if skip_translation else render_table_preview_html(table_row, "en")
        table_views.append(
            {
                "preview_title": table_row.preview_title,
                "table_format": table_row.table_format,
                "sheet_name": table_row.sheet_name,
                "html_table": render_table_preview_html(table_row, "nl"),
                "html_table_en": html_en,
            }
        )

    return render_template_string(
        TABLE_TEMPLATE, doc=doc, tables=table_views, error_msg=error_msg
    )


@app.route("/table/<int:document_id>/refresh")
def refresh_table_preview(document_id):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc:
        flash("Document not found.")
        return redirect(url_for("home"))

    if doc.extension.lower() not in {"csv", "xlsx"}:
        flash("This document does not contain a spreadsheet table.")
        return redirect(url_for("home"))

    try:
        tables = regenerate_table_previews(doc)
        if tables:
            flash(f"Table refreshed with translation ({len(tables)} sheet(s)).")
        else:
            flash("No readable table data found in this spreadsheet.")
    except FileNotFoundError:
        flash("Source file not found. Re-upload the spreadsheet and try again.")
    except Exception as exc:
        flash(f"Table refresh failed: {exc}")

    return redirect(url_for("table_preview", document_id=document_id))

@app.route("/document-view/<int:document_id>")
def document_view(document_id):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc:
        abort(404)
    page = request.args.get("page", 1, type=int)
    context = build_document_view_context(doc, page=page)
    if not context:
        abort(404)
    return render_template_string(DOCUMENT_VIEWER_TEMPLATE, **context)


@app.route("/pdf/<int:document_id>")
def pdf_viewer(document_id):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc or doc.extension != "pdf":
        abort(404)
    page = request.args.get("page", 1, type=int)
    context = build_document_view_context(doc, page=page)
    if not context:
        abort(404)
    context["viewer_route"] = "pdf_viewer"
    return render_template_string(DOCUMENT_VIEWER_TEMPLATE, **context)


@app.route("/docx/<int:document_id>")
def docx_viewer(document_id):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc or doc.extension != "docx":
        abort(404)
    page = request.args.get("page", 1, type=int)
    highlight_image_index = request.args.get("image_index", type=int)
    context = build_document_view_context(
        doc, page=page, highlight_image_index=highlight_image_index
    )
    if not context:
        abort(404)
    context["viewer_route"] = "docx_viewer"
    return render_template_string(DOCUMENT_VIEWER_TEMPLATE, **context)


@app.route("/serve-document/<int:document_id>")
def serve_document(document_id):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc or doc.extension.lower() not in TRANSLATABLE_EXTENSIONS:
        abort(404)

    if not storage.document_exists(doc.stored_filename, DOC_FOLDER):
        abort(404)

    file_bytes = storage.read_document_bytes(doc.stored_filename, DOC_FOLDER)
    mime_map = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "txt": "text/plain",
    }
    return send_file(
        io.BytesIO(file_bytes),
        mimetype=mime_map.get(doc.extension.lower(), "application/octet-stream"),
        as_attachment=False,
        download_name=doc.original_filename,
    )


@app.route("/serve-pdf/<int:document_id>")
def serve_pdf(document_id):
    return serve_document(document_id)

@app.route("/page-image/<int:document_id>")
def page_image(document_id):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc or doc.extension != "pdf":
        abort(404)
    page = request.args.get("page", 1, type=int)
    if not storage.document_exists(doc.stored_filename, DOC_FOLDER):
        abort(404)
    with storage.open_document_local_path(doc.stored_filename, DOC_FOLDER) as file_path:
        image_data = get_page_image(file_path, page)
    if image_data is None:
        # fallback: return a 1x1 transparent PNG (or 404)
        from flask import Response
        return Response(status=404)
    return send_file(io.BytesIO(image_data[0]), mimetype=image_data[1])


@app.route("/docx-image/<int:document_id>/<int:image_index>")
def docx_image(document_id, image_index):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc or doc.extension != "docx":
        abort(404)
    if not storage.document_exists(doc.stored_filename, DOC_FOLDER):
        abort(404)
    with storage.open_document_local_path(doc.stored_filename, DOC_FOLDER) as file_path:
        image_data = get_docx_image(file_path, image_index)
    if image_data is None:
        abort(404)
    return send_file(io.BytesIO(image_data[0]), mimetype=image_data[1])

def safe_excel_sheet_name(name, used_names=None) -> str:
    value = re.sub(r"[\\/*?:\[\]]+", "_", str(name or "Sheet")).strip(" .")
    if not value:
        value = "Sheet"
    value = value[:31]
    if used_names is not None:
        base = value
        suffix = 1
        while value in used_names:
            tail = f"_{suffix}"
            value = f"{base[:31 - len(tail)]}{tail}"
            suffix += 1
        used_names.add(value)
    return value


def csv_text_to_dataframe(csv_text):
    if not (csv_text or "").strip():
        return None
    try:
        df = pd.read_csv(io.StringIO(csv_text), dtype=str, keep_default_na=False)
        for bad in ("nan", "NaN", "None", "<NA>"):
            df = df.replace(bad, "")
        return df
    except Exception as exc:
        print(f"CSV table parse error: {exc}")
        return None


def html_table_to_dataframe(html):
    if not (html or "").strip():
        return None
    try:
        frames = pd.read_html(io.StringIO(html))
        if not frames:
            return None
        df = frames[0].fillna("").astype(str)
        for bad in ("nan", "NaN", "None", "<NA>"):
            df = df.replace(bad, "")
        return df
    except Exception as exc:
        print(f"HTML table parse error: {exc}")
        return None


def table_preview_to_dataframe(table_row, lang="nl"):
    lang = (lang or "nl").strip().lower()
    csv_field = {
        "nl": "csv_text",
        "en": "csv_text_en",
        "es": "csv_text_es",
    }.get(lang, "csv_text")
    df = csv_text_to_dataframe(getattr(table_row, csv_field, "") or "")
    if df is not None and not df.empty:
        return df

    html_field = {
        "nl": "html_table",
        "en": "html_table_en",
        "es": "html_table_es",
    }.get(lang, "html_table")
    df = html_table_to_dataframe(getattr(table_row, html_field, "") or "")
    if df is not None and not df.empty:
        return df

    if lang == "nl" or not translation.translation_enabled():
        return df

    nl_df = csv_text_to_dataframe(table_row.csv_text)
    if nl_df is None or nl_df.empty:
        return None
    try:
        return translation.translate_dataframe_values(
            nl_df,
            target_lang=lang,
            cache_get=get_cached_translation,
            cache_set=store_cached_translation,
            max_cells=TRANSLATE_MAX_CELLS,
        )
    except Exception as exc:
        print(f"On-demand table translation failed ({lang}): {exc}")
        return None


def _coerce_excel_cell_value(value):
    if value in (None, "", "nan", "NaN", "None", "<NA>"):
        return ""
    text = str(value).strip()
    if re.fullmatch(r"-?\d+", text):
        return int(text)
    if re.fullmatch(r"-?\d+\.\d+", text):
        return float(text)
    return text


def _row_has_single_banner_cell(row_values):
    filled = [value for value in row_values if str(value).strip()]
    if len(filled) != 1:
        return False
    text = str(filled[0]).strip()
    return len(text) >= 20 or "|" in text or "Drawing" in text or "SPE." in text


_EXCEL_THIN_BORDER = Border(
    left=Side(style="thin", color="000000"),
    right=Side(style="thin", color="000000"),
    top=Side(style="thin", color="000000"),
    bottom=Side(style="thin", color="000000"),
)
_EXCEL_WRAP_ALIGN = Alignment(wrap_text=True, vertical="top")


def _format_excel_table_block(ws, min_row, max_row, min_col, max_col):
    for row in range(min_row, max_row + 1):
        for col in range(min_col, max_col + 1):
            cell = ws.cell(row=row, column=col)
            cell.border = _EXCEL_THIN_BORDER
            cell.alignment = _EXCEL_WRAP_ALIGN


def _auto_fit_excel_columns(ws, min_col, max_col, min_row, max_row):
    for col in range(min_col, max_col + 1):
        max_len = 10
        for row in range(min_row, max_row + 1):
            value = ws.cell(row=row, column=col).value
            if value is not None:
                max_len = max(max_len, len(str(value)))
        ws.column_dimensions[get_column_letter(col)].width = min(max(max_len + 2, 12), 55)


def _append_dataframe_to_sheet(ws, df, start_row, title=None):
    row_cursor = start_row
    if title:
        title_cell = ws.cell(row=row_cursor, column=1, value=title)
        title_cell.font = Font(bold=True, size=12)
        row_cursor += 1

    if df is None or df.empty:
        return row_cursor

    ncols = len(df.columns)
    header_row = row_cursor
    for col_offset, col_name in enumerate(df.columns, start=1):
        header_cell = ws.cell(row=header_row, column=col_offset, value=col_name)
        header_cell.font = Font(bold=True)
        header_cell.alignment = _EXCEL_WRAP_ALIGN
    row_cursor += 1

    data_start = row_cursor
    for row_offset, row_values in enumerate(df.values):
        for col_offset, value in enumerate(row_values, start=1):
            ws.cell(
                row=row_cursor + row_offset,
                column=col_offset,
                value=_coerce_excel_cell_value(value),
            )

    data_end = row_cursor + len(df) - 1
    if ncols > 0:
        for row_idx in range(data_start, data_end + 1):
            row_values = [
                ws.cell(row=row_idx, column=col_idx).value
                for col_idx in range(1, ncols + 1)
            ]
            if _row_has_single_banner_cell(row_values):
                ws.merge_cells(
                    start_row=row_idx,
                    start_column=1,
                    end_row=row_idx,
                    end_column=ncols,
                )
                merged_cell = ws.cell(row=row_idx, column=1)
                merged_cell.alignment = _EXCEL_WRAP_ALIGN

        _format_excel_table_block(ws, header_row, data_end, 1, ncols)
        _auto_fit_excel_columns(ws, 1, ncols, header_row, data_end)

    return data_end + 3


@app.route("/table/<int:document_id>/export-xlsx")
def export_table_xlsx(document_id):
    doc = db.session.get(DocumentRecord, document_id)
    if not doc:
        abort(404)

    tables = TablePreview.query.filter_by(document_id=document_id).all()
    if not tables:
        flash("No table data to export.")
        return redirect(url_for("table_preview", document_id=document_id))

    wb = Workbook()
    ws = wb.active
    ws.title = str(  # pyright: ignore[reportOptionalMemberAccess]
        safe_excel_sheet_name(
            Path(doc.original_filename).stem or f"document_{document_id}"
        )
    )

    row_cursor = 1
    wrote_content = False
    single_table = len(tables) == 1
    section = ""

    for index, table_row in enumerate(tables, start=1):
        if not single_table:
            section = table_row.sheet_name or table_row.preview_title or f"Table {index}"
            section_cell = cast(Any, ws.cell(row=row_cursor, column=1, value=section))  # pyright: ignore[reportOptionalMemberAccess]
            # pyright: ignore[reportOptionalMemberAccess]
            section_cell.font = Font(bold=True, size=13)
            row_cursor += 2

        nl_df = table_preview_to_dataframe(table_row, "nl")
        if nl_df is not None and not nl_df.empty:
            nl_title = "Table 1 - Dutch (Original)" if single_table else f"{section} - Dutch (Original)"
            row_cursor = _append_dataframe_to_sheet(ws, nl_df, row_cursor, nl_title)
            wrote_content = True

        en_df = table_preview_to_dataframe(table_row, "en")
        if en_df is not None and not en_df.empty:
            en_title = "Table 2 - English (Translation)" if single_table else f"{section} - English (Translation)"
            row_cursor = _append_dataframe_to_sheet(ws, en_df, row_cursor, en_title)
            wrote_content = True

    if not wrote_content:
        flash("No table content available to export.")
        return redirect(url_for("table_preview", document_id=document_id))

    xlsx_buffer = io.BytesIO()
    wb.save(xlsx_buffer)
    xlsx_buffer.seek(0)
    base_name = secure_filename(Path(doc.original_filename).stem or f"document_{document_id}")
    return send_file(
        xlsx_buffer,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=f"{base_name}_tables.xlsx",
    )


@app.route("/delete-multiple", methods=["POST"])
def bulk_delete_documents():
    raw_ids = request.form.getlist("document_ids")
    if not raw_ids:
        flash("No documents selected.", "warning")
        return redirect_admin_documents("admin-feedback")

    deleted_count = 0
    missing_count = 0
    error_count = 0

    for raw_id in raw_ids:
        try:
            document_id = int(raw_id)
        except ValueError:
            error_count += 1
            continue

        doc = db.session.get(DocumentRecord, document_id)
        if not doc:
            missing_count += 1
            continue

        try:
            storage.delete_document(doc.stored_filename, DOC_FOLDER)

            db.session.delete(doc)
            db.session.commit()
            deleted_count += 1
        except Exception as e:
            db.session.rollback()
            print(f"Bulk delete error for document {document_id}: {e}")
            error_count += 1

    flash(
        f"Bulk delete completed. Deleted: {deleted_count}, Missing: {missing_count}, Errors: {error_count}.",
        "success" if error_count == 0 else "warning",
    )
    return redirect_admin_documents("admin-feedback")



def ensure_schema():
    from sqlalchemy import inspect, text

    inspector = inspect(db.engine)
    doc_cols = {col["name"] for col in inspector.get_columns("documents")}
    if "page_offsets_json" not in doc_cols:
        db.session.execute(text("ALTER TABLE documents ADD COLUMN page_offsets_json TEXT"))
        db.session.commit()
    if "document_type" not in doc_cols:
        db.session.execute(
            text("ALTER TABLE documents ADD COLUMN document_type VARCHAR(50) DEFAULT 'standard'")
        )
        db.session.commit()

    block_cols = {col["name"] for col in inspector.get_columns("requirement_blocks")}
    if "char_start" not in block_cols:
        db.session.execute(text("ALTER TABLE requirement_blocks ADD COLUMN char_start INTEGER DEFAULT 0"))
        db.session.commit()
    if "major_section" not in block_cols:
        db.session.execute(text("ALTER TABLE requirement_blocks ADD COLUMN major_section VARCHAR(300)"))
        db.session.commit()

    table_cols = {col["name"] for col in inspector.get_columns("table_previews")}
    if "html_table_en" not in table_cols:
        db.session.execute(text("ALTER TABLE table_previews ADD COLUMN html_table_en TEXT"))
        db.session.commit()
    if "html_table_es" not in table_cols:
        db.session.execute(text("ALTER TABLE table_previews ADD COLUMN html_table_es TEXT"))
        db.session.commit()
    if "csv_text_en" not in table_cols:
        db.session.execute(text("ALTER TABLE table_previews ADD COLUMN csv_text_en TEXT"))
        db.session.commit()
    if "csv_text_es" not in table_cols:
        db.session.execute(text("ALTER TABLE table_previews ADD COLUMN csv_text_es TEXT"))
        db.session.commit()

    page_translation_cols = set()
    if "document_page_translations" in inspector.get_table_names():
        page_translation_cols = {
            col["name"] for col in inspector.get_columns("document_page_translations")
        }
    if page_translation_cols and "translated_text_es" not in page_translation_cols:
        db.session.execute(
            text("ALTER TABLE document_page_translations ADD COLUMN translated_text_es TEXT")
        )
        db.session.commit()

    cache_cols = set()
    if "translation_cache" in inspector.get_table_names():
        cache_cols = {col["name"] for col in inspector.get_columns("translation_cache")}
    if cache_cols and "provider" not in cache_cols:
        db.session.execute(text("ALTER TABLE translation_cache ADD COLUMN provider VARCHAR(32)"))
        db.session.commit()

    if "dictionary_sources" in inspector.get_table_names():
        dict_source_cols = {
            col["name"] for col in inspector.get_columns("dictionary_sources")
        }
        if "size_bytes" not in dict_source_cols:
            db.session.execute(
                text("ALTER TABLE dictionary_sources ADD COLUMN size_bytes INTEGER DEFAULT 0")
            )
            db.session.commit()

    if "dictionary_entries" in inspector.get_table_names():
        dict_entry_cols = {
            col["name"] for col in inspector.get_columns("dictionary_entries")
        }
        if "content_nl" not in dict_entry_cols:
            db.session.execute(text("ALTER TABLE dictionary_entries ADD COLUMN content_nl TEXT"))
            db.session.commit()
        if "content_en" not in dict_entry_cols:
            db.session.execute(text("ALTER TABLE dictionary_entries ADD COLUMN content_en TEXT"))
            db.session.commit()
        if "search_keys" not in dict_entry_cols:
            db.session.execute(text("ALTER TABLE dictionary_entries ADD COLUMN search_keys TEXT"))
            db.session.commit()

    if "document_page_translations" not in inspector.get_table_names():
        db.create_all()


# =========================================================
# App initialization
# =========================================================
with app.app_context():
    db.create_all()
    ensure_schema()
    create_default_admin()

def _running_in_notebook():
    try:
        from IPython import get_ipython
        return get_ipython() is not None
    except Exception:
        return False


in_notebook = _running_in_notebook()

if __name__ == "__main__":
    port = int(os.environ.get("PORT", "5000"))
    debug = env_flag("FLASK_DEBUG", default=not IS_PRODUCTION and not in_notebook)
    app.run(
        host="0.0.0.0",
        port=port,
        debug=debug,
        use_reloader=debug,
    )
